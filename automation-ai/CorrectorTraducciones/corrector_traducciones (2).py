"""
Corrector de Traducciones v7.0 — BloomScans
python corrector_traducciones.py

Requiere: Python 3.8+ con tkinter
Opcional:
  pandoc                          → mejor lectura de .docx
  Pillow                          → visor + compresión de imágenes (pip install Pillow)
  pyspellchecker                  → corrector ortográfico (pip install pyspellchecker)
  google-auth-oauthlib + google-api-python-client → Drive
  requests                        → IA Gemini (pip install requests)

MEJORAS v7.0 — IA Gemini (Google AI Studio):
  - Reemplaza Meta/Together.ai por Gemini Vision (más estable, key gratuita)
  - API Key gratis en: aistudio.google.com/apikey
  - Selección de modelo: gemini-2.0-flash / flash-lite / 1.5-flash / 1.5-pro
  - MODO DETECTAR: globos sin traducir, orden, simbología, onomatopeyas, coherencia
  - MODO CORREGIR: genera script corregido completo aplicable con un click
  - Envío de imagen actual O todas las imágenes del capítulo para más contexto
  - Control de calidad JPEG (20–95%) y altura máxima para reducir tokens
  - Compresión con PIL si está disponible, fallback sin PIL

MEJORAS v6.0 (IA visual):
  - Panel IA integrado con Meta Llama Vision
MEJORAS v4.0 (sobre v3.0):
  1.  Auto-tilde con preview y confirmación antes de aplicar
  4.  Corrector ortográfico integrado
  5.  Detección de mezcla de sistemas de simbología (BL/+18 vs +15)
  7.  Sincronización bidireccional visor ↔ editor
  8b. Panel de botones de símbolos clickeables
MEJORAS v3.0:
  2.  Validación de numeración de tiras
  3.  Exportar reporte de errores a .txt
  6.  Líneas muy largas configurables
  8.  Visor de manwha (local/ZIP/Drive)
  9.  Coherencia de nombres propios
  10. Google Drive: abrir .docx, subir corregido, cargar raws
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import re, os, threading, json, tempfile, base64, io
from pathlib import Path

# IA META — requests opcional
try:
    import requests as _requests
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

# ═══════════════════════════════════════════════════════════════════════════════
#  GEMINI (Google AI Studio) — configuración
# ═══════════════════════════════════════════════════════════════════════════════
GEMINI_KEY_FILE = Path("gemini_api_key.txt")

# Modelos disponibles (de mejor a más económico)
GEMINI_MODELOS = {
    "gemini-2.0-flash":        "Gemini 2.0 Flash  (rápido, muy económico ★★★)",
    "gemini-2.0-flash-lite":   "Gemini 2.0 Flash Lite  (ultra económico ★★★★)",
    "gemini-1.5-flash":        "Gemini 1.5 Flash  (equilibrado ★★★)",
    "gemini-1.5-flash-8b":     "Gemini 1.5 Flash-8B  (más económico ★★★★)",
    "gemini-1.5-pro":          "Gemini 1.5 Pro  (máxima calidad ★★)",
}
GEMINI_MODELO_DEFAULT = "gemini-2.0-flash"

# URL base de la API
GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta/models"

_SIMBOLOS_CONTEXTO = """
Simbología BloomScans:
G: = Grito  |  T: = Nervioso/Tembloroso/Sollozo  |  P: = Pensamiento
DF: = Diálogo fuera de globo  |  N: = Narración/Cuadro  |  (): = Llamadas/Otros
*: = Terror  |  NT: = Nota de traductor
D2: = Diálogo 2  |  P2: = Pensamiento 2  |  T2: = Tembloroso 2
(()): = Diálogo bonito  |  {}: = Nube  |  <>: = Carta  |  /: = Celular  |  []: = Cuadro juego
"""

def _cargar_gemini_key() -> str:
    if GEMINI_KEY_FILE.exists():
        return GEMINI_KEY_FILE.read_text(encoding="utf-8").strip()
    return ""

def _guardar_gemini_key(key: str):
    GEMINI_KEY_FILE.write_text(key.strip(), encoding="utf-8")


def _procesar_imagen(fuente, calidad: int = 60, max_alto: int = 1200) -> tuple:
    """
    Convierte bytes / path str / PIL.Image a (base64_str, mime_type).
    Aplica reducción de calidad y tamaño para ahorrar tokens de Gemini.
    calidad: 1-95 (JPEG quality). 60 = bueno para OCR, bajo costo.
    max_alto: altura máxima en px. Las páginas de manhwa suelen ser muy altas.
    """
    # Intentar redimensionar y comprimir con PIL
    _pil_ok = False
    try:
        from PIL import Image as _Img
        _pil_ok = True
    except ImportError:
        pass

    if _pil_ok:
        try:
            if isinstance(fuente, bytes):
                img = _Img.open(io.BytesIO(fuente))
            elif isinstance(fuente, str):
                img = _Img.open(fuente)
            else:
                img = fuente  # ya es PIL.Image

            # Convertir a RGB (elimina canal alpha, reduce tamaño)
            if img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            # Redimensionar si es demasiado alta
            ow, oh = img.size
            if oh > max_alto:
                ratio = max_alto / oh
                nw = max(1, int(ow * ratio))
                img = img.resize((nw, max_alto), _Img.LANCZOS)

            # Comprimir como JPEG
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=calidad, optimize=True)
            return base64.b64encode(buf.getvalue()).decode(), "image/jpeg"
        except Exception:
            pass  # fallback a modo sin PIL

    # Fallback sin PIL: devolver raw en base64
    if isinstance(fuente, bytes):
        raw = fuente
    elif isinstance(fuente, str):
        with open(fuente, "rb") as f:
            raw = f.read()
    else:
        buf = io.BytesIO(); fuente.save(buf, format="PNG"); raw = buf.getvalue()

    mime = "image/jpeg"
    if raw[:8] == b'\x89PNG\r\n\x1a\n': mime = "image/png"
    return base64.b64encode(raw).decode(), mime


def _construir_instruccion(modo: str, tira_nombre: str, todas_tiras: bool = False) -> str:
    """Genera el prompt según el modo."""
    ctx_img = "todas las imágenes adjuntas del capítulo" if todas_tiras else "la imagen adjunta"

    GUIA_GLOBOS = """
GUÍA DE TIPOS DE GLOBO Y SU SIMBOLOGÍA:
• Globo redondo/ovalado normal con cola → diálogo normal → D2: (o G: si hay líneas de tensión)
• Globo con bordes en zigzag/rayos/explosivo → grito fuerte → G:
• Globo con borde tembloroso/ondulado/irregular → nervios, temblor, sollozo → T:
• Nube (círculos pequeños como burbujas) → pensamiento → P:
• Caja/rectángulo de narración → narración/cuadro → N:
• Texto fuera de cualquier globo pero dentro de la viñeta → diálogo fuera de globo → DF:
• Globo con borde suave/redondeado doble → diálogo bonito/especial → (()): 
• Globo en forma de nube esponjosa → nube → {}:
• Rectángulo con borde decorativo → carta/mensaje escrito → <>:
• Globo con icono de teléfono o líneas de señal → celular → /:
• Rectángulo estilo pantalla/videojuego → cuadro juego → []:
• Llamada/ring de teléfono / sonido externo entre paréntesis → ():
• Efectos de sonido/onomatopeyas DENTRO de globo → corresponde al tipo de globo
• Efectos de sonido/onomatopeyas FUERA de globo y que son puramente decorativos → NO van en el script
"""

    if modo == "detectar":
        return f"""Sos un corrector de traducciones de manhwa/manga para BloomScans.

{_SIMBOLOS_CONTEXTO}
{GUIA_GLOBOS}

Se te proporciona: {ctx_img} y el script de traducción al español para esa(s) tira(s).

Analizá CADA globo visible en la imagen y comparalo con el script. Reportá lo siguiente:

═══ ERRORES A DETECTAR ═══

[GLOBO FALTANTE] Un globo con texto visible en la imagen NO aparece en el script.
Formato: [GLOBO FALTANTE] Tira X — posición/descripción del globo: "texto original si podés leerlo"

[ORDEN INCORRECTO] El orden de líneas en el script no sigue el orden de lectura visual
(izquierda→derecha, arriba→abajo por panel/viñeta).
Formato: [ORDEN INCORRECTO] Tira X — la línea "..." debería ir antes/después de "..."

[SIMBOLOGÍA] La simbología usada no corresponde al tipo de globo visible en la imagen.
Mirá la forma del globo: ¿es redondo normal? ¿tiene bordes en zigzag (grito)? ¿es nube de burbujas (pensamiento)? ¿tiene borde tembloroso (nervios)? ¿es cuadro de narración?
Formato: [SIMBOLOGÍA] Tira X, línea "..." — usó SIMBOLO_ACTUAL: pero el globo visible es de tipo TIPO → debería ser SIMBOLO_CORRECTO:

[ONOMATOPEYA EXTRA] Hay en el script un efecto de sonido/onomatopeya que en la imagen
está escrito FUERA de todo globo de forma puramente decorativa (no es diálogo ni narración).
Formato: [ONOMATOPEYA EXTRA] Tira X — "texto" está fuera de globo en la imagen

═══ REGLAS ═══
✓ Sí reportá problemas de simbología aunque no estés 100% seguro — indicalo con "(probable)"
✓ Sí reportá globos faltantes aunque el texto sea pequeño o difícil de leer
✓ NO corrijas ortografía, tildes ni redacción — solo estructura y simbología
✓ NO reportes líneas como GLOBO EXTRA si tienen un globo asociado en la imagen

Si no encontrás errores respondé:
✓ Sin errores visuales detectados.
"""
    else:  # corregir
        return f"""Sos un corrector de traducciones de manhwa/manga para BloomScans.

{_SIMBOLOS_CONTEXTO}
{GUIA_GLOBOS}

Se te proporciona: {ctx_img} y el script de traducción al español.

CORREGÍ el script basándote en lo que ves en las imágenes. Aplicá estos cambios:

1. GLOBOS FALTANTES: Agregá una línea por cada globo que existe en la imagen pero no en el script.
   Usá la simbología correcta según la forma del globo (ver guía arriba).
   Para el texto usá "???" si no podés leerlo con claridad.

2. ORDEN: Reordenás las líneas para que coincidan con el orden visual de lectura de los globos.

3. SIMBOLOGÍA: Corregís el símbolo si no corresponde al tipo de globo visible.
   Ej: si el globo tiene bordes en zigzag → G:, si es nube de burbujas → P:, etc.

4. ONOMATOPEYAS DECORATIVAS: Eliminás líneas que sean efectos de sonido sueltos fuera de globos.

⚠️ NO modifiques el texto de las traducciones existentes.
⚠️ NO corrijas ortografía ni tildes.
⚠️ Mantenés el formato: ## Tira N → SIMBOLO: texto

Respondé ÚNICAMENTE con el script corregido, sin explicaciones.
"""


def llamar_gemini(api_key: str, modo: str,
                  texto_script: str,
                  imagenes: list,          # lista de fuentes (bytes/path/PIL)
                  tira_nombre: str = "",
                  modelo: str = GEMINI_MODELO_DEFAULT,
                  calidad: int = 60,
                  max_alto: int = 1200,
                  timeout: int = 120) -> str:
    """
    Llama a la API de Gemini (Google AI Studio).
    imagenes: lista de fuentes de imagen — puede ser 1 (tira actual) o todas.
    calidad: calidad JPEG 1-95 para comprimir antes de enviar.
    max_alto: altura máxima en píxeles para redimensionar.
    """
    if not REQUESTS_OK:
        raise RuntimeError("Instalá requests: pip install requests")
    if not api_key:
        raise RuntimeError("Ingresá tu API Key de Google AI Studio.")

    todas = len(imagenes) > 1
    instruccion = _construir_instruccion(modo, tira_nombre, todas)

    # Primero el texto del script, LUEGO las imágenes
    # (Gemini procesa mejor cuando el texto de referencia está antes de las imágenes)
    parts = [{"text": instruccion + "\n\n--- SCRIPT A REVISAR ---\n" + texto_script
              + "\n--- FIN DEL SCRIPT ---\n\nAhora analizá las imágenes adjuntas:"}]

    for i, fuente in enumerate(imagenes):
        try:
            b64, mime = _procesar_imagen(fuente, calidad, max_alto)
            parts.append({"inline_data": {"mime_type": mime, "data": b64}})
        except Exception as ex:
            parts.append({"text": f"[Imagen {i+1} no disponible: {ex}]"})

    url = f"{GEMINI_API_BASE}/{modelo}:generateContent?key={api_key}"
    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {
            "temperature": 0.3,   # equilibrado: preciso pero no demasiado tímido
            "maxOutputTokens": 2048,
        }
    }

    resp = _requests.post(url,
                          headers={"Content-Type": "application/json"},
                          json=payload, timeout=timeout)
    if resp.status_code != 200:
        data = resp.json() if resp.content else {}
        msg = data.get("error", {}).get("message", resp.text[:300])
        raise RuntimeError(f"Error Gemini ({resp.status_code}): {msg}")

    data = resp.json()
    texto = (data.get("candidates", [{}])[0]
                 .get("content", {})
                 .get("parts", [{}])[0]
                 .get("text", ""))
    if not texto.strip():
        # Intentar extraer de safety feedback
        finish = (data.get("candidates", [{}])[0].get("finishReason", ""))
        if finish == "SAFETY":
            raise RuntimeError("Gemini bloqueó la respuesta por filtros de seguridad. "
                               "Probá con otro modelo o reducí el contenido sensible.")
        raise RuntimeError("Gemini devolvió una respuesta vacía.")
    return texto.strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  PIL opcional
# ═══════════════════════════════════════════════════════════════════════════════
try:
    from PIL import Image, ImageTk
    PIL_OK = True
except Exception:
    PIL_OK = False

# ── Corrector ortográfico (MEJORA 4) ──────────────────────────────────────────
try:
    from spellchecker import SpellChecker as _SC
    _spell_es = _SC(language="es")
    _VOCAB_EXTRA = {
        "eunsol","jinhwan","jisu","umi","yeo","ra","senpai","idol",
        "manhwa","webtoon","taekwondo","peto","mesada","paparazzi",
        "jajaja","jaja","jeje","badum","kyaaa","waah","aaah","bah","uf",
        "mhm","hmm","shh","ops","d2","p2","t2","nt","df",
    }
    _spell_es.word_frequency.load_words(_VOCAB_EXTRA)
    SPELL_OK = True
except Exception:
    SPELL_OK = False
    _spell_es = None

# ═══════════════════════════════════════════════════════════════════════════════
#  PALETA
# ═══════════════════════════════════════════════════════════════════════════════
C = {
    "bg": "#080b12", "bg2": "#0d1120", "bg3": "#111827",
    "bg4": "#182030", "bg5": "#1e2840", "border": "#1e2d47", "border2": "#253248",
    "teal": "#00cfb5", "teal2": "#009e8a", "teal3": "#006d5e",
    "cyan": "#1adeff",
    "text": "#ccd9ee", "text2": "#7a9cbf", "dim": "#354d6a", "bright": "#ffffff",
    "err": "#ef4444", "ok": "#22c55e", "warn": "#f59e0b", "purple": "#a78bfa",
    "err_punt_bg": "#2d1a1a", "err_punt_fg": "#ff8080",
    "err_may_bg":  "#1a1d2d", "err_may_fg":  "#8080ff",
    "err_sim_bg":  "#2d2a1a", "err_sim_fg":  "#ffcc60",
    "err_til_bg":  "#1a2a1a", "err_til_fg":  "#80ff80",
    "err_dup_bg":  "#2a1a2d", "err_dup_fg":  "#cc80ff",
    "err_long_bg": "#1a1a2d", "err_long_fg": "#80c0ff",
    "err_ort_bg":  "#2a1525", "err_ort_fg":  "#ff80cc",
    "err_mix_bg":  "#2a2010", "err_mix_fg":  "#ffaa40",
    "search_bg":   "#2d2800", "search_fg":   "#ffe066",
}

# ═══════════════════════════════════════════════════════════════════════════════
#  SIMBOLOGÍA
# ═══════════════════════════════════════════════════════════════════════════════
SIMBOLOS = {
    "G:":    ("Grito",                            "BL/+18"),
    "T:":    ("Nervioso / Tembloroso / Sollozo",  "BL/+18"),
    "P:":    ("Pensamiento",                      "BL/+18"),
    "DF:":   ("Diálogo fuera de globo",           "BL/+18"),
    "N:":    ("Narración / Cuadro",               "BL/+18"),
    "():":   ("Llamadas / Otros",                 "BL/+18"),
    "*:":    ("Terror",                           "BL/+18"),
    "NT:":   ("Nota de traductor",                "BL/+18"),
    "D2:":   ("Diálogo 2",                        "+15"),
    "P2:":   ("Pensamiento 2",                    "+15"),
    "T2:":   ("Nervioso / Tembloroso / Sollozo 2","+15"),
    "(()): ":("Diálogo bonito",                   "+15"),
    "{}:":   ("Nube",                             "+15"),
    "<>:":   ("Carta",                            "+15"),
    "/:":    ("Celular",                          "+15"),
    "[]:":   ("Cuadro juego",                     "+15"),
}
PREFIJOS = set(s.strip() for s in SIMBOLOS)

SIMBOLOS_TYPOS = {
    "{}": "{}:", "{ }:": "{}:", "{ }": "{}:",
    "<> :": "<>:", "< >:": "<>:", "[] :": "[]:", "[ ]:": "[]:",
    "NT :": "NT:", "DF :": "DF:", "D2 :": "D2:", "P2 :": "P2:", "T2 :": "T2:",
    "d2:": "D2:", "p2:": "P2:", "t2:": "T2:", "df:": "DF:", "nt:": "NT:",
    "n:": "N:", "g:": "G:", "p:": "P:", "t:": "T:",
}

TILDES = [
    ("tambien","también"),("todavia","todavía"),("asi","así"),
    ("alli","allí"),("aqui","aquí"),("ahi","ahí"),
    ("mas","más"),("jamas","jamás"),("ademas","además"),
    ("algun","algún"),("ningun","ningún"),
    ("entrare","entraré"),("llegare","llegaré"),
    ("saldra","saldrá"),("saldras","saldrás"),("saldran","saldrán"),
    ("traera","traerá"),("traeras","traerás"),("traeran","traerán"),
    ("hara","hará"),("haras","harás"),("haran","harán"),
    ("podra","podrá"),("podras","podrás"),("podran","podrán"),
    ("vendra","vendrá"),("vendras","vendrás"),("vendran","vendrán"),
    ("tendra","tendrá"),("tendras","tendrás"),("tendran","tendrán"),
    ("sera","será"),("seras","serás"),("seran","serán"),
    ("dira","dirá"),("diras","dirás"),("diran","dirán"),
    ("dara","dará"),("daras","darás"),("daran","darán"),
    ("volvera","volverá"),("volveras","volverás"),("volveran","volverán"),
    ("pasara","pasará"),("pasaras","pasarás"),("pasaran","pasarán"),
    ("quedara","quedará"),("quedaras","quedarás"),("quedaran","quedarán"),
    ("estara","estará"),("estaras","estarás"),("estaran","estarán"),
    ("habra","habrá"),("habras","habrás"),("habran","habrán"),
    ("lograra","logrará"),("lograras","lograrás"),("lograran","lograrán"),
    ("pedira","pedirá"),("pediras","pedirás"),("pediran","pedirán"),
    ("vivira","vivirá"),("viviras","vivirás"),("viviran","vivirán"),
]

MAX_CHARS_LINEA = 200

# ── MEJORA 5: Mezcla de sistemas de simbología ───────────────────────────────
# BL/+18: P, T, G, DF, N, (), *, NT
# +15:    D2, P2, T2, (()), {}, <>, /, []
# Si coexisten P y P2 (o T y T2) en el mismo archivo → probablemente error
_PARES_CONFLICTO = [
    ("P:",  "P2:",  "Pensamiento"),
    ("T:",  "T2:",  "Nervioso/Sollozo"),
]

def detectar_mezcla_sistemas(texto: str) -> list:
    """Devuelve lista de (num_linea, 'mezcla', mensaje)."""
    lineas = texto.splitlines()
    presentes = {}   # sym -> primera línea donde aparece
    for i, linea in enumerate(lineas):
        l = linea.strip()
        if not l or es_encabezado_tira(l): continue
        p, _ = extraer_prefijo(l)
        if p and p not in presentes:
            presentes[p] = i + 1

    problemas = []
    for s_bl, s_15, nombre in _PARES_CONFLICTO:
        if s_bl in presentes and s_15 in presentes:
            # Reportar en la línea del símbolo que aparece segundo (el conflicto)
            n = max(presentes[s_bl], presentes[s_15])
            problemas.append((n, "mezcla",
                f"Mezcla de sistemas: '{s_bl}' y '{s_15}' para {nombre} — elegí uno"))
    return problemas


# ── MEJORA 4: Ortografía ─────────────────────────────────────────────────────
_IGNORAR_SPELL = {
    "jajaja","jaja","ja","jeje","jejeje","mhm","mm","ng","uf","ugh","ay","bah",
    "wow","waah","aah","aaah","kyaaa","badum","yay","ops","ups","eh","ah",
    "oh","uh","hm","hmm","shh","sss","jiji","bwah","ugh",
    "eunsol","jinhwan","jisu","umi","yeo","ra","senpai","idol",
    "manhwa","webtoon","taekwondo","peto","mesada","paparazzi",
    "vámonos","váyanse","oigan","oye","mira","anda","vaya",
}
# ── Español neutro: palabras peninsulares a evitar ───────────────────────────
_PALABRAS_PENINSULAR = {
    "vosotros": "ustedes",
    "vosotras": "ustedes",
    "erais":    "eran",
    "estabais": "estaban",
    "ibais":    "iban",
    "habíais":  "habían",
    "seríais":  "serían",
    "tendríais":"tendrían",
    "querríais":"querrían",
    "habréis":  "habrán",
    "seréis":   "serán",
    "iréis":    "irán",
    "tenéis":   "tienen",
    "queréis":  "quieren",
    "sois":     "son",
    "hacéis":   "hacen",
    "sabéis":   "saben",
    "podéis":   "pueden",
    "decís":    "dicen",
    "veis":     "ven",
    "dais":     "dan",
    "vais":     "van",
    "estáis":   "están",
    "habéis":   "han",
    "fuisteis": "fueron",
    "disteis":  "dieron",
}

# ── Onomatopeyas de movimiento / sonido cortas: ignorar check de mayúsculas ──
_PALABRAS_ONOMATOPEYA_DF = {
    "paso","tap","clap","bam","bang","crash","zas","pum","toc","clic",
    "click","splash","swoosh","whoosh","crack","snap","pop","boom",
    "pánico","temblor","susurro","murmullo","silencio","ring","bip",
    "sprint","corre","huye","avanza","retrocede","salta","cae","gira",
    "ruido","golpe","cruje","ruge","llora","ríe","grita","suspi",
}



def revisar_ortografia_linea(contenido: str, nombres_propios: set = None) -> list:
    """
    Devuelve lista de (palabra_original, sugerencia) para typos evidentes.
    Solo si SPELL_OK es True.
    """
    if not SPELL_OK:
        return []
    ignorar = _IGNORAR_SPELL.copy()
    if nombres_propios:
        ignorar |= {n.lower() for n in nombres_propios}

    # Solo palabras de 4+ letras en español (con tildes)
    palabras = re.findall(r"[a-záéíóúüñA-ZÁÉÍÓÚÜÑ]{4,}", contenido)
    resultado = []
    for p in palabras:
        pl = p.lower()
        if pl in ignorar: continue
        if p[0].isupper() and len(p) <= 8: continue   # probable nombre propio
        misspelled = _spell_es.unknown([pl])
        if misspelled:
            candidatos = _spell_es.candidates(pl) or set()
            if candidatos:
                mejor = min(candidatos, key=lambda x: abs(len(x) - len(pl)))
                if mejor != pl:
                    resultado.append((p, mejor))
                    break   # 1 por línea para no saturar
    return resultado


# ── MEJORA 1: Auto-tilde con listas de cambios ───────────────────────────────
def calcular_cambios_tilde(texto: str) -> list:
    """
    Calcula todos los reemplazos de tilde posibles.
    Devuelve lista de (num_linea, palabra_sin_tilde, palabra_con_tilde, pos_en_linea).
    """
    cambios = []
    for i, linea in enumerate(texto.splitlines()):
        l = linea.strip()
        if not l or es_encabezado_tira(l): continue
        _, contenido = extraer_prefijo(l)
        cl = contenido.lower()
        for sin, con in TILDES:
            m = re.search(r"\b" + re.escape(sin) + r"\b", cl)
            if m:
                original = contenido[m.start():m.end()]
                if original.lower() == sin:
                    cambios.append((i + 1, original, con, m.start()))
                    break   # 1 por línea
    return cambios


def aplicar_cambios_tilde(texto: str, seleccionados: set) -> tuple:
    """
    Aplica solo los cambios de tilde cuyo (num_linea, sin) esté en seleccionados.
    Devuelve (texto_nuevo, cantidad_aplicada).
    """
    lineas = texto.splitlines()
    resultado = []
    aplicados = 0
    for i, linea in enumerate(lineas):
        num = i + 1
        nueva = linea
        prefijo_m = re.match(r"^(\s*[A-Z][A-Z0-9]*:\s*|\s*\([^)]*\):\s*|"
                             r"\s*[{}<>\[\]*/]+:\s*|\s*)", linea)
        offset = len(prefijo_m.group()) if prefijo_m else 0
        contenido = linea[offset:]
        for sin, con in TILDES:
            if (num, sin) in seleccionados:
                patron = r"\b" + re.escape(sin) + r"\b"
                nueva_c, n = re.subn(patron, con, contenido, flags=re.IGNORECASE)
                if n:
                    nueva = linea[:offset] + nueva_c
                    aplicados += n
                    break
        resultado.append(nueva)
    return "\n".join(resultado), aplicados

# Google Drive — credentials embebidos (mismo proyecto que BloomStitch)
GDRIVE_CLIENT_SECRET = {
    "installed": {
        "client_id":     "YOUR_GOOGLE_CLIENT_ID.apps.googleusercontent.com.apps.googleusercontent.com",
        "project_id":    "stitch-488305",
        "auth_uri":      "https://accounts.google.com/o/oauth2/auth",
        "token_uri":     "https://oauth2.googleapis.com/token",
        "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
        "client_secret": "YOUR_GOOGLE_CLIENT_SECRET",
        "redirect_uris": ["http://localhost"]
    }
}
GDRIVE_SCOPES    = ["https://www.googleapis.com/auth/drive"]
GDRIVE_TOKEN_FILE = Path("corrector_gdrive_token.json")

# ═══════════════════════════════════════════════════════════════════════════════
#  LÓGICA DE ANÁLISIS
# ═══════════════════════════════════════════════════════════════════════════════
def extraer_prefijo(linea: str):
    l = linea.strip()
    for sym in sorted(PREFIJOS, key=len, reverse=True):
        if l.startswith(sym + " ") or l == sym:
            return sym, l[len(sym):].lstrip()
    return None, l



def analizar_linea(linea: str, nombres_propios: set = None):
    """Devuelve lista de (tipo, mensaje, col) donde col es posicion 0-based en la linea."""
    errores = []
    l = linea.strip()
    if not l or es_encabezado_tira(l):
        return errores

    line_offset = len(linea) - len(linea.lstrip())

    # Simbolo con typo
    primer = l.split()[0] if l.split() else ""
    for typo, correcto in SIMBOLOS_TYPOS.items():
        if primer.lower() == typo.lower():
            if primer not in PREFIJOS:
                errores.append(("simbolo", f"Simbolo mal escrito '{primer}' -> '{correcto}'", line_offset))
                break
    if not any(t == "simbolo" for t, *_ in errores):
        m_pref = re.match(r"^([a-z][a-z0-9]*:|[<>\[\]{}()+*/]+:)\s", l)
        if m_pref:
            cand = m_pref.group(1)
            if cand.upper() in PREFIJOS and cand not in PREFIJOS:
                errores.append(("simbolo", f"Simbolo en minusculas '{cand}' -> '{cand.upper()}'", line_offset))

    prefijo, contenido = extraer_prefijo(l)
    if not contenido:
        return errores

    prefijo_len = len(prefijo) + 1 if prefijo else 0
    col_base = line_offset + prefijo_len

    # ¿?
    ab = ci = 0
    for idx_ch, ch in enumerate(contenido):
        if ch == "¿": ab += 1
        if ch == "?":
            ci += 1
            if ci > ab:
                errores.append(("puntuacion", "Falta ¿ de apertura para ?", col_base + idx_ch))
                break
    if contenido.count("¿") > contenido.count("?"):
        pos = contenido.rfind("¿")
        errores.append(("puntuacion", "Falta ? de cierre para ¿", col_base + pos))

    # ¡!
    ab = ci = 0
    for idx_ch, ch in enumerate(contenido):
        if ch == "¡": ab += 1
        if ch == "!":
            ci += 1
            if ci > ab:
                errores.append(("puntuacion", "Falta ¡ de apertura para !", col_base + idx_ch))
                break
    if contenido.count("¡") > contenido.count("!"):
        pos = contenido.rfind("¡")
        errores.append(("puntuacion", "Falta ! de cierre para ¡", col_base + pos))

    # ── Mayuscula al inicio ──────────────────────────────────────────────────
    stripped_i = contenido.lstrip("¿¡\"'`...—-")
    off_i = len(contenido) - len(stripped_i)
    if stripped_i and stripped_i[0].isalpha() and stripped_i[0].islower():
        errores.append(("mayuscula", f"Minuscula al inicio: '{contenido[:22]}...'", col_base + off_i))

    # ── REGLA NUEVA 1: Texto en MAYÚSCULAS ───────────────────────────────────
    # Detecta cuando la mayor parte del contenido está en mayúsculas sostenidas.
    # Excepción: onomatopeyas/movimiento cortas en DF:/G:/N:
    palabras_alpha = [p for p in re.findall(r"[A-Za-záéíóúüñÁÉÍÓÚÜÑ]{2,}", contenido) if p.isalpha()]
    if palabras_alpha:
        n_mayus = sum(1 for p in palabras_alpha if p == p.upper())
        es_onoma_df = (
            prefijo == "DF:" and
            len(palabras_alpha) <= 3 and
            all(p.lower() in _PALABRAS_ONOMATOPEYA_DF or len(p) <= 5 for p in palabras_alpha)
        )
        if n_mayus >= 2 and (n_mayus / len(palabras_alpha)) >= 0.6 and not es_onoma_df:
            sugerencia = re.sub(
                r"\b([A-ZÁÉÍÓÚÜÑ]{2,})\b",
                lambda m2: m2.group(1)[0] + m2.group(1)[1:].lower(),
                contenido
            )
            errores.append(("mayuscula",
                f"Texto en MAYÚSCULAS → convertir a normal: '{sugerencia[:35]}...'",
                col_base))

    # ── REGLA NUEVA 2: Español neutro (palabras peninsulares) ────────────────
    cont_lower_pen = contenido.lower()
    for palabra_pen, alternativa in _PALABRAS_PENINSULAR.items():
        if re.search(r"\b" + re.escape(palabra_pen) + r"\b", cont_lower_pen):
            pos_pen = cont_lower_pen.find(palabra_pen)
            errores.append(("mayuscula",
                f"Español no neutro: '{palabra_pen}' → usar '{alternativa}'",
                col_base + pos_pen))
            break

    # ── REGLA NUEVA 3: Minúscula tras ... dentro de ¡...! o ¿...? ────────────
    # Ej: ¡Vamos a... darle! → debería ser ¡Vamos a... Darle!
    for mp in re.finditer(r"(\.{3}|…)[ ]?([a-záéíóúüñ])", contenido):
        antes = contenido[:mp.start()]
        abre_excl = antes.count("¡") - antes.count("!")
        abre_inter = antes.count("¿") - antes.count("?")
        if abre_excl > 0 or abre_inter > 0:
            letra = mp.group(2)
            errores.append(("mayuscula",
                f"Minúscula tras '...' dentro de exclamación: "
                f"'...{letra}' → '...{letra.upper()}'",
                col_base + mp.start()))

    # Mayusculas internas: DESACTIVADO

    # Tilde faltante
    cont_lower = contenido.lower()
    for sin, con in TILDES:
        m = re.search(r"\b" + re.escape(sin) + r"\b", cont_lower)
        if m:
            orig = contenido[m.start():m.end()]
            if orig.lower() == sin:
                errores.append(("tilde", f"Tilde faltante: '{orig}' -> '{con}'", col_base + m.start()))
                break

    # Guiones multiples
    m_g = re.search(r"\w(-{2,})", contenido)
    if m_g:
        errores.append(("estructura", "Guion de corte multiple (---): usar solo (-)", col_base + m_g.start(1)))

    # Linea termina con coma
    if contenido.rstrip("~\\").endswith(","):
        errores.append(("puntuacion", "Linea termina con coma ,", col_base + len(contenido) - 1))

    # Linea muy larga
    if len(contenido) > MAX_CHARS_LINEA:
        errores.append(("largo", f"Linea muy larga ({len(contenido)} chars > {MAX_CHARS_LINEA})", col_base + MAX_CHARS_LINEA))

    # Ortografia
    ort = revisar_ortografia_linea(contenido, nombres_propios)
    for orig_w, sug in ort:
        pos_w = contenido.find(orig_w)
        errores.append(("ortografia", f"Posible typo: '{orig_w}' -> ??'{sug}'??", col_base + max(0, pos_w)))

    # Espacio antes de puntuacion
    m_esp = re.search(r" [,;:!?]", contenido)
    if m_esp:
        errores.append(("puntuacion", "Espacio antes de signo de puntuacion", col_base + m_esp.start()))

    # Puntos dobles
    m_pts = re.search(r"\.\.(?!\.)", re.sub(r"\.{3,}", "···", contenido))
    if m_pts:
        errores.append(("puntuacion", "Puntos dobles .. (puntos suspensivos?)", col_base + m_pts.start()))

    # Comas dobles
    m_cc = re.search(r",,", contenido)
    if m_cc:
        errores.append(("puntuacion", "Comas dobles ,,", col_base + m_cc.start()))

    return errores


def es_encabezado_tira(linea: str) -> bool:
    """
    Detecta cualquier variante de encabezado de tira:
      ## Tira 2 | ##Tira 2 | \\## Tira 2 | ### ##Tira 1 | ##tira2 | etc.
    """
    l = re.sub(r'^[#\s]+(?=##)', '', linea.strip())
    return bool(re.match(r'^\\?##', l))


def numero_tira_de_linea(linea: str):
    """Extrae el número de tira de un encabezado. Devuelve int o None."""
    m = re.search(r'\d+', linea)
    return int(m.group()) if m else None


def analizar_duplicados(lineas):
    dups = {}
    vistos = []
    for i, linea in enumerate(lineas):
        l = linea.strip()
        if es_encabezado_tira(l):
            vistos.clear(); continue
        if not l: continue
        _, contenido = extraer_prefijo(l)
        norm = contenido.lower().strip()
        if norm and norm in vistos:
            dups[i + 1] = f"Línea duplicada en esta tira: '{contenido[:30]}'"
        else:
            vistos.append(norm)
    return dups


# MEJORA 2: Validación de tiras ───────────────────────────────────────────────
def analizar_tiras(lineas):
    """Detecta tiras faltantes, repetidas o mal numeradas."""
    errores_tira = []
    nums = []
    for i, linea in enumerate(lineas):
        if es_encabezado_tira(linea):
            n = numero_tira_de_linea(linea)
            if n is not None:
                nums.append((i + 1, n))
    for idx, (num_linea, n) in enumerate(nums):
        if idx > 0:
            prev = nums[idx - 1][1]
            if n == prev:
                errores_tira.append((num_linea, f"Tira {n} repetida"))
            elif n != prev + 1:
                errores_tira.append((num_linea,
                    f"Salto de tira: va de Tira {prev} a Tira {n} (¿falta Tira {prev+1}?)"))
    return errores_tira


# MEJORA 9: Coherencia de nombres propios ─────────────────────────────────────
def detectar_nombres_propios(texto: str) -> set:
    """Extrae palabras que aparecen con mayúscula en múltiples contextos."""
    contador = {}
    for linea in texto.splitlines():
        _, contenido = extraer_prefijo(linea.strip())
        for palabra in contenido.split():
            p = re.sub(r"[¿¡?!.,;:\"'~\-*\\]", "", palabra)
            if len(p) > 3 and p[0].isupper():
                base = p.lower()
                contador[base] = contador.get(base, 0) + 1
    return {w.capitalize() for w, c in contador.items() if c >= 2}


def analizar_nombres(texto: str, nombres: set) -> list:
    """
    Detecta variantes de escritura de un mismo nombre.
    Ej: 'Eunsol' vs 'EunsoL' vs 'Eunsól'
    """
    errores = []
    # Agrupar por nombre base (primera letra + len similar)
    variantes = {}
    for nombre in nombres:
        clave = (nombre[0].lower(), len(nombre))
        variantes.setdefault(clave, set()).add(nombre)

    # Buscar en el texto formas similares que difieran en 1-2 chars
    for linea in texto.splitlines():
        num = texto[:texto.find(linea)].count("\n") + 1 if linea in texto else 0
        _, contenido = extraer_prefijo(linea.strip())
        for palabra in contenido.split():
            p = re.sub(r"[¿¡?!.,;:\"'~\-*\\]", "", palabra)
            if len(p) < 4 or not p[0].isupper(): continue
            for nombre_ref in nombres:
                if p != nombre_ref and p.lower() != nombre_ref.lower():
                    # Distancia de edición simple: si difieren en solo 1 char
                    if len(p) == len(nombre_ref):
                        diffs = sum(a != b for a, b in zip(p.lower(), nombre_ref.lower()))
                        if diffs == 1:
                            errores.append((num, "nombre",
                                f"Posible error en nombre: '{p}' vs '{nombre_ref}'"))
    return errores


def estadisticas(texto):
    stats = {s: 0 for s in PREFIJOS}
    stats.update({"SIN_SIMBOLO": 0, "TIRAS": 0, "TOTAL_LINEAS": 0})
    for linea in texto.splitlines():
        l = linea.strip()
        if not l: continue
        if es_encabezado_tira(l):
            stats["TIRAS"] += 1; continue
        stats["TOTAL_LINEAS"] += 1
        p, _ = extraer_prefijo(l)
        if p: stats[p] = stats.get(p, 0) + 1
        else:  stats["SIN_SIMBOLO"] += 1
    return stats


def _capitalizar_mayusculas_sostenidas(contenido: str) -> str:
    """Convierte TEXTO EN MAYÚSCULAS SOSTENIDAS → Texto en minúsculas.
    Solo aplica si hay 2+ palabras en caps. Preserva palabras cortas (onomatopeyas)."""
    palabras = re.findall(r"[A-Za-záéíóúüñÁÉÍÓÚÜÑ]{2,}", contenido)
    if not palabras:
        return contenido
    n_mayus = sum(1 for p in palabras if p.isalpha() and p == p.upper())
    if n_mayus < 2 or n_mayus / len(palabras) < 0.6:
        return contenido
    def _cap_word(m2):
        w = m2.group(0)
        if not w.isalpha() or len(w) <= 2:
            return w
        return w[0].upper() + w[1:].lower()
    return re.sub(r"[A-ZÁÉÍÓÚÜÑ]{2,}", _cap_word, contenido)


def _mayuscula_tras_puntos_suspensivos(contenido: str) -> str:
    """¡Vamos a... darle! → ¡Vamos a... Darle!
    Sube a mayúscula la letra tras ... cuando hay ¡ o ¿ sin cerrar antes."""
    def _reemplazar(m):
        antes = contenido[:m.start()]
        abre_excl = antes.count("¡") - antes.count("!")
        abre_inter = antes.count("¿") - antes.count("?")
        if abre_excl > 0 or abre_inter > 0:
            return m.group(1) + m.group(2) + m.group(3).upper()
        return m.group(0)
    return re.sub(r"(\.{3}|…)(\s?)([a-záéíóúüñ])", _reemplazar, contenido)


def corregir_texto(texto):
    lineas = texto.splitlines()
    resultado = []; cambios = 0
    for linea in lineas:
        l, orig = linea, linea
        l = re.sub(r" ([,;:!?])", r"\1", l)
        l = l.replace(",,", ",")
        l = re.sub(r"(?<!\.)\.[{2}(?!\.", "...", l)
        l = re.sub(r"(\w)-{2,}", r"\1-", l)
        m = re.match(r"^([A-Z][A-Z0-9]*:|\([^)]*\):|[{}<>\[\]*/]+:)\s*", l)
        if m:
            prefijo_part = l[:m.end()]
            resto = l[m.end():]
            resto = _mayuscula_tras_puntos_suspensivos(resto)
            resto = _capitalizar_mayusculas_sostenidas(resto)
            idx = next((i for i, c in enumerate(resto) if c.isalpha()), None)
            if idx is not None and resto[idx].islower():
                resto = resto[:idx] + resto[idx].upper() + resto[idx+1:]
            l = prefijo_part + resto
        else:
            stripped = l.lstrip()
            if stripped and not es_encabezado_tira(stripped):
                lead = l[:len(l)-len(stripped)]
                stripped = _mayuscula_tras_puntos_suspensivos(stripped)
                stripped = _capitalizar_mayusculas_sostenidas(stripped)
                idx = None
                for i, c in enumerate(stripped):
                    if c.isalpha(): idx = i; break
                    elif c not in "¿¡\"'": break
                if idx is not None and stripped[idx].islower():
                    stripped = stripped[:idx] + stripped[idx].upper() + stripped[idx+1:]
                l = lead + stripped
        if l != orig: cambios += 1
        resultado.append(l)
    return "\n".join(resultado), cambios


def exportar_reporte(errores_por_linea: dict, errores_tira: list, texto: str, path: str):
    lineas = texto.splitlines()
    with open(path, "w", encoding="utf-8") as f:
        f.write("═" * 60 + "\n")
        f.write("  REPORTE DE ERRORES — BloomScans Corrector v3.0\n")
        f.write("═" * 60 + "\n\n")

        # Errores de estructura de tiras
        if errores_tira:
            f.write("── ERRORES DE ESTRUCTURA DE TIRAS ──\n\n")
            for num_linea, msg in errores_tira:
                f.write(f"  L{num_linea:4d}  ⚠  {msg}\n")
            f.write("\n")

        # Errores por línea
        total = sum(len(v) for v in errores_por_linea.values())
        f.write(f"── ERRORES EN TRADUCCIONES ({total} total) ──\n\n")

        ICONOS = {"puntuacion":"⚠","mayuscula":"Aa","simbolo":"?",
                  "tilde":"´","estructura":"⚙","duplicado":"⊡",
                  "largo":"↔","nombre":"👤","tira":"#"}
        for num_linea in sorted(errores_por_linea):
            errores = errores_por_linea[num_linea]
            linea_txt = lineas[num_linea - 1] if num_linea <= len(lineas) else ""
            f.write(f"  L{num_linea:4d}  {linea_txt[:60]}\n")
            for err_t in errores:
                tipo = err_t[0]; msg = err_t[1]; col = err_t[2] if len(err_t) > 2 else 0
                icono = ICONOS.get(tipo, "•")
                col_str = f" [col {col+1}]" if col > 0 else ""
                f.write(f"         {icono}  [{tipo.upper()}]{col_str} {msg}\n")
            f.write("\n")

        f.write("═" * 60 + "\n")
        f.write(f"  Total: {total} error(es) en {len(errores_por_linea)} línea(s)\n")
        f.write("═" * 60 + "\n")


# ═══════════════════════════════════════════════════════════════════════════════
#  GOOGLE DRIVE
# ═══════════════════════════════════════════════════════════════════════════════
def _ensure_gdrive():
    needed = []
    try: import google.oauth2.credentials
    except ImportError: needed.append("google-auth-oauthlib")
    try: import googleapiclient.discovery
    except ImportError: needed.append("google-api-python-client")
    if needed:
        import subprocess, sys
        subprocess.run([sys.executable, "-m", "pip", "install", "--quiet"] + needed,
                       capture_output=True, timeout=120)

class DriveClient:
    def __init__(self):
        _ensure_gdrive()
        self.service = None
        self.creds   = None

    def conectar(self) -> bool:
        try:
            from google.oauth2.credentials import Credentials
            from google_auth_oauthlib.flow import InstalledAppFlow
            from google.auth.transport.requests import Request as GRequest
            from googleapiclient.discovery import build as gdrive_build

            creds = None
            if GDRIVE_TOKEN_FILE.exists():
                creds = Credentials.from_authorized_user_file(
                    str(GDRIVE_TOKEN_FILE), GDRIVE_SCOPES)
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(GRequest())
                else:
                    tf = tempfile.NamedTemporaryFile(
                        mode="w", suffix=".json", delete=False, encoding="utf-8")
                    json.dump(GDRIVE_CLIENT_SECRET, tf); tf.close()
                    flow = InstalledAppFlow.from_client_secrets_file(
                        tf.name, GDRIVE_SCOPES)
                    creds = flow.run_local_server(port=0, open_browser=True)
                    try: Path(tf.name).unlink()
                    except: pass
                GDRIVE_TOKEN_FILE.write_text(creds.to_json(), encoding="utf-8")
            self.creds   = creds
            self.service = gdrive_build("drive", "v3", credentials=creds,
                                        cache_discovery=False)
            return True
        except Exception as ex:
            raise RuntimeError(str(ex))

    def desconectar(self):
        try:
            if GDRIVE_TOKEN_FILE.exists(): GDRIVE_TOKEN_FILE.unlink()
        except: pass
        self.service = self.creds = None

    def listar_docx(self, carpeta_id: str = None):
        q = "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' and trashed=false"
        if carpeta_id:
            q += f" and '{carpeta_id}' in parents"
        res = self.service.files().list(q=q, fields="files(id,name,modifiedTime)",
                                        orderBy="modifiedTime desc",
                                        pageSize=50).execute()
        return res.get("files", [])

    def listar_imagenes(self, carpeta_id: str):
        q = (f"'{carpeta_id}' in parents and trashed=false and "
             "(mimeType contains 'image/')")
        res = self.service.files().list(q=q, fields="files(id,name)",
                                        orderBy="name", pageSize=200).execute()
        return res.get("files", [])

    def descargar_bytes(self, file_id: str) -> bytes:
        import io
        from googleapiclient.http import MediaIoBaseDownload
        req = self.service.files().get_media(fileId=file_id)
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        return buf.getvalue()

    def subir_archivo(self, path: str, carpeta_id: str = None, nombre: str = None):
        from googleapiclient.http import MediaFileUpload
        nombre = nombre or Path(path).name
        meta = {"name": nombre}
        if carpeta_id:
            meta["parents"] = [carpeta_id]
        media = MediaFileUpload(path, resumable=True)
        f = self.service.files().create(body=meta, media_body=media,
                                        fields="id,name").execute()
        return f

    def buscar_carpeta(self, nombre: str):
        q = f"name='{nombre}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
        res = self.service.files().list(q=q, fields="files(id,name)",
                                        pageSize=10).execute()
        return res.get("files", [])


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPERS UI
# ═══════════════════════════════════════════════════════════════════════════════
def make_btn(parent, text, cmd, bg=None, fg=None, hover_bg=None, hover_fg=None,
             small=False, **kw):
    _bg = bg or C["teal"]; _fg = fg or C["bg"]
    _hbg = hover_bg or C["cyan"]; _hfg = hover_fg or _fg
    _font = ("Segoe UI", 8) if small else ("Segoe UI", 9, "bold")
    _py = 3 if small else 7; _px = kw.pop("padx", 8 if small else 14)
    b = tk.Button(parent, text=text, command=cmd, bg=_bg, fg=_fg, font=_font,
                  relief="flat", cursor="hand2", padx=_px, pady=_py,
                  activebackground=_hbg, activeforeground=_hfg, bd=0, **kw)
    b.bind("<Enter>", lambda e: b.config(bg=_hbg, fg=_hfg))
    b.bind("<Leave>", lambda e: b.config(bg=_bg, fg=_fg))
    return b

def sep(parent, **kw):
    tk.Frame(parent, bg=C["border"], height=1, **kw).pack(fill="x")

def scroll_text(parent, **kw):
    outer = tk.Frame(parent, bg=C["bg"])
    sb = ttk.Scrollbar(outer, orient="vertical", style="Vertical.TScrollbar")
    t = tk.Text(outer, bg=C["bg2"], fg=C["dim"], font=("Segoe UI", 8),
                wrap="word", relief="flat", bd=0, padx=8, pady=6,
                state="disabled", yscrollcommand=sb.set, **kw)
    sb.config(command=t.yview)
    sb.pack(side="right", fill="y"); t.pack(side="left", fill="both", expand=True)
    return outer, t


# ═══════════════════════════════════════════════════════════════════════════════
#  VISOR DE MANWHA (MEJORA 8)
# ═══════════════════════════════════════════════════════════════════════════════
class VisorManwha(tk.Toplevel):
    """
    Ventana flotante con las imágenes del manwha.
    Fuentes: carpeta/ZIP local, o imágenes descargadas de Drive.
    Se puede sincronizar con la tira activa del editor.
    """
    def __init__(self, parent, on_close=None, on_navigate=None):
        super().__init__(parent)
        self.title("🖼 Visor de Manwha")
        self.configure(bg=C["bg"])
        self.geometry("520x780")
        self.minsize(300, 400)
        self._on_close = on_close
        self._on_navigate = on_navigate   # MEJORA 7: callback bidireccional
        self.protocol("WM_DELETE_WINDOW", self._cerrar)

        self._imagenes = []   # list of (nombre_str, bytes_o_path_str)
        self._idx = 0
        self._tk_img = None
        self._zoom = 1.0

        self._build()

    def _build(self):
        # Barra superior
        bar = tk.Frame(self, bg=C["bg2"]); bar.pack(fill="x")
        tk.Label(bar, text="🖼  Visor de Manwha", bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=10, pady=7)

        bf = tk.Frame(bar, bg=C["bg2"]); bf.pack(side="right", padx=6, pady=5)
        make_btn(bf, "📁 Abrir carpeta", self._abrir_carpeta,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(side="left", padx=2)
        make_btn(bf, "🗜 Abrir ZIP", self._abrir_zip,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(side="left", padx=2)

        sep(self)

        # Navegación
        nav = tk.Frame(self, bg=C["bg3"]); nav.pack(fill="x")
        make_btn(nav, "◀◀", lambda: self._ir(0),
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="left", padx=4, pady=4)
        make_btn(nav, "◀", self._anterior,
                 bg=C["bg3"], fg=C["text"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="left", padx=2, pady=4)
        self._lbl_idx = tk.Label(nav, text="—", bg=C["bg3"], fg=C["text2"],
                                 font=("Segoe UI", 9))
        self._lbl_idx.pack(side="left", expand=True)
        make_btn(nav, "▶", self._siguiente,
                 bg=C["bg3"], fg=C["text"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="right", padx=2, pady=4)
        make_btn(nav, "▶▶", lambda: self._ir(len(self._imagenes) - 1),
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="right", padx=4, pady=4)

        # Zoom
        zoom_bar = tk.Frame(self, bg=C["bg2"]); zoom_bar.pack(fill="x")
        tk.Label(zoom_bar, text="Zoom:", bg=C["bg2"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(side="left", padx=8)
        for z, lbl in [(0.5,"50%"),(0.75,"75%"),(1.0,"100%"),(1.5,"150%"),(2.0,"200%")]:
            zz = z
            make_btn(zoom_bar, lbl, lambda v=zz: self._set_zoom(v),
                     bg=C["bg2"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                     small=True).pack(side="left", padx=1, pady=3)

        sep(self)

        # Canvas con scrollbars
        canvas_frame = tk.Frame(self, bg=C["bg"]); canvas_frame.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(canvas_frame, orient="vertical",   style="Vertical.TScrollbar")
        hsb = ttk.Scrollbar(canvas_frame, orient="horizontal", style="Horizontal.TScrollbar")
        self._canvas = tk.Canvas(canvas_frame, bg=C["bg2"], highlightthickness=0,
                                 yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self._canvas.yview)
        hsb.config(command=self._canvas.xview)
        vsb.pack(side="right", fill="y"); hsb.pack(side="bottom", fill="x")
        self._canvas.pack(side="left", fill="both", expand=True)
        self._canvas.bind("<MouseWheel>",
                          lambda e: self._canvas.yview_scroll(int(-1*(e.delta/120)),"units"))
        self._canvas.bind("<Button-4>",
                          lambda e: self._canvas.yview_scroll(-1,"units"))
        self._canvas.bind("<Button-5>",
                          lambda e: self._canvas.yview_scroll(1,"units"))

        # Label sin PIL
        if not PIL_OK:
            self._canvas.create_text(260, 100, text="Instala Pillow para ver imágenes:\npip install Pillow",
                                     fill=C["warn"], font=("Segoe UI", 10), anchor="center")

    def _abrir_carpeta(self):
        carpeta = filedialog.askdirectory(title="Seleccionar carpeta con imágenes del manwha")
        if not carpeta: return
        exts = {".jpg",".jpeg",".png",".webp",".bmp",".gif"}
        archivos = sorted(
            f for f in Path(carpeta).iterdir()
            if f.suffix.lower() in exts
        )
        self._imagenes = [(f.name, str(f)) for f in archivos]
        self._idx = 0
        self._mostrar()

    def _abrir_zip(self):
        path = filedialog.askopenfilename(
            title="Abrir ZIP con imágenes del manwha",
            filetypes=[("ZIP", "*.zip"), ("Todos", "*.*")])
        if not path: return
        import zipfile, io
        exts = {".jpg",".jpeg",".png",".webp",".bmp"}
        imgs = []
        with zipfile.ZipFile(path) as z:
            for nombre in sorted(z.namelist()):
                if Path(nombre).suffix.lower() in exts:
                    data = z.read(nombre)
                    imgs.append((Path(nombre).name, data))
        self._imagenes = imgs
        self._idx = 0
        self._mostrar()

    def cargar_desde_drive(self, imagenes: list):
        """Recibe [(nombre, bytes), ...] ya descargados."""
        self._imagenes = imagenes
        self._idx = 0
        self._mostrar()

    def ir_a_tira(self, n: int):
        """Navega a la imagen de la tira N (0-based)."""
        if 0 <= n < len(self._imagenes):
            self._ir(n)

    def _ir(self, idx: int):
        self._idx = max(0, min(idx, len(self._imagenes) - 1))
        self._mostrar()
        # MEJORA 7: notificar al editor para que haga scroll a la tira
        if self._on_navigate:
            try:
                self._on_navigate(self._idx)
            except Exception:
                pass

    def _anterior(self):
        self._ir(self._idx - 1)

    def _siguiente(self):
        self._ir(self._idx + 1)

    def _set_zoom(self, z: float):
        self._zoom = z
        self._mostrar()

    def _mostrar(self):
        if not self._imagenes or not PIL_OK:
            return
        nombre, fuente = self._imagenes[self._idx]
        self._lbl_idx.config(
            text=f"{nombre}  ({self._idx + 1} / {len(self._imagenes)})")

        import io
        if isinstance(fuente, bytes):
            img = Image.open(io.BytesIO(fuente))
        else:
            img = Image.open(fuente)

        w, h = img.size
        nw, nh = int(w * self._zoom), int(h * self._zoom)
        img = img.resize((nw, nh), Image.LANCZOS)
        self._tk_img = ImageTk.PhotoImage(img)

        self._canvas.delete("all")
        self._canvas.config(scrollregion=(0, 0, nw, nh))
        self._canvas.create_image(0, 0, anchor="nw", image=self._tk_img)

    def _cerrar(self):
        if self._on_close: self._on_close()
        self.destroy()


# ═══════════════════════════════════════════════════════════════════════════════
#  PANEL DRIVE (MEJORA 10)
# ═══════════════════════════════════════════════════════════════════════════════
class PanelDrive(tk.Toplevel):
    """Ventana para operaciones con Google Drive."""

    def __init__(self, parent, on_abrir_texto, on_imagenes_drive, visor_ref):
        super().__init__(parent)
        self.title("☁ Google Drive")
        self.configure(bg=C["bg"])
        self.geometry("540x580")
        self._on_abrir_texto    = on_abrir_texto
        self._on_imagenes_drive = on_imagenes_drive
        self._visor_ref         = visor_ref
        self._drive = DriveClient()
        self._conectado = False
        self._build()

    def _build(self):
        bar = tk.Frame(self, bg=C["bg2"]); bar.pack(fill="x")
        tk.Label(bar, text="☁  Google Drive", bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 11, "bold")).pack(side="left", padx=14, pady=10)
        self._dot = tk.Label(bar, text="● Desconectado", bg=C["bg2"], fg=C["err"],
                             font=("Segoe UI", 8))
        self._dot.pack(side="right", padx=14)
        sep(self)

        inner = tk.Frame(self, bg=C["bg"]); inner.pack(fill="both", expand=True, padx=14, pady=10)

        # Auth
        auth_card = self._card(inner, "🔑  Autenticación")
        make_btn(auth_card, "Conectar con Google", self._conectar,
                 bg=C["teal"], fg=C["bg"], hover_bg=C["cyan"],
                 small=True).pack(side="left")
        make_btn(auth_card, "Desconectar", self._desconectar,
                 bg=C["bg4"], fg=C["err"], hover_bg=C["bg5"],
                 small=True).pack(side="left", padx=8)

        # Abrir .docx
        doc_card = self._card(inner, "📄  Abrir traducción desde Drive")
        tk.Label(doc_card, text="Carpeta o ID de Drive (opcional):",
                 bg=C["bg3"], fg=C["text2"], font=("Segoe UI", 8)).pack(anchor="w")
        self._var_carpeta_doc = tk.StringVar()
        tk.Entry(doc_card, textvariable=self._var_carpeta_doc,
                 bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                 relief="flat", font=("Segoe UI", 9)).pack(fill="x", pady=(2,6), ipady=3)
        make_btn(doc_card, "📋 Ver .docx disponibles", self._listar_docx,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(anchor="w")
        list_frame = tk.Frame(doc_card, bg=C["bg3"]); list_frame.pack(fill="x", pady=(4,0))
        dsb = ttk.Scrollbar(list_frame, orient="vertical", style="Vertical.TScrollbar")
        self._docx_list = tk.Listbox(list_frame, bg=C["bg2"], fg=C["text"],
                                     selectbackground=C["teal3"],
                                     selectforeground=C["bright"],
                                     font=("Segoe UI", 8), relief="flat", bd=0,
                                     height=5, activestyle="none",
                                     yscrollcommand=dsb.set)
        dsb.config(command=self._docx_list.yview)
        dsb.pack(side="right", fill="y"); self._docx_list.pack(side="left", fill="both", expand=True)
        make_btn(doc_card, "⬇ Abrir seleccionado", self._abrir_docx_drive,
                 bg=C["teal3"], fg=C["bright"], hover_bg=C["teal2"],
                 small=True).pack(anchor="w", pady=(4,0))

        # Imágenes manwha
        img_card = self._card(inner, "🖼  Cargar imágenes del manwha desde Drive")
        tk.Label(img_card, text="ID o nombre de la carpeta de imágenes:",
                 bg=C["bg3"], fg=C["text2"], font=("Segoe UI", 8)).pack(anchor="w")
        self._var_carpeta_img = tk.StringVar()
        tk.Entry(img_card, textvariable=self._var_carpeta_img,
                 bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                 relief="flat", font=("Segoe UI", 9)).pack(fill="x", pady=(2,6), ipady=3)
        make_btn(img_card, "🖼 Cargar en visor", self._cargar_imagenes_drive,
                 bg=C["teal3"], fg=C["bright"], hover_bg=C["teal2"],
                 small=True).pack(anchor="w")
        self._lbl_img_status = tk.Label(img_card, text="", bg=C["bg3"], fg=C["dim"],
                                        font=("Segoe UI", 8))
        self._lbl_img_status.pack(anchor="w", pady=(4,0))

        # Subir corregido
        sub_card = self._card(inner, "⬆  Subir archivo corregido a Drive")
        tk.Label(sub_card, text="ID de carpeta destino (opcional):",
                 bg=C["bg3"], fg=C["text2"], font=("Segoe UI", 8)).pack(anchor="w")
        self._var_carpeta_sub = tk.StringVar()
        tk.Entry(sub_card, textvariable=self._var_carpeta_sub,
                 bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                 relief="flat", font=("Segoe UI", 9)).pack(fill="x", pady=(2,6), ipady=3)
        make_btn(sub_card, "⬆ Subir archivo guardado", self._subir,
                 bg=C["teal"], fg=C["bg"], hover_bg=C["cyan"],
                 small=True).pack(anchor="w")
        self._lbl_sub_status = tk.Label(sub_card, text="", bg=C["bg3"], fg=C["dim"],
                                        font=("Segoe UI", 8))
        self._lbl_sub_status.pack(anchor="w", pady=(4,0))

        self._lbl_status = tk.Label(self, text="", bg=C["bg"], fg=C["dim"],
                                    font=("Segoe UI", 8))
        self._lbl_status.pack(pady=4)

    def _card(self, parent, titulo):
        f = tk.Frame(parent, bg=C["bg3"],
                     highlightbackground=C["border2"], highlightthickness=1)
        f.pack(fill="x", pady=(0,8))
        tk.Label(f, text=titulo, bg=C["bg3"], fg=C["teal"],
                 font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=10, pady=(6,2))
        inner = tk.Frame(f, bg=C["bg3"])
        inner.pack(fill="x", padx=10, pady=(0,8))
        return inner

    def _set_status(self, msg, color=None):
        self._lbl_status.config(text=msg, fg=color or C["dim"])

    def _conectar(self):
        self._set_status("🔄 Conectando…", C["warn"])
        def _worker():
            try:
                self._drive.conectar()
                self._conectado = True
                self.after(0, lambda: (
                    self._dot.config(text="● Conectado", fg=C["ok"]),
                    self._set_status("✅ Conectado a Google Drive", C["ok"])
                ))
            except Exception as ex:
                self.after(0, lambda m=str(ex): (
                    self._dot.config(text="● Error", fg=C["err"]),
                    self._set_status(f"❌ {m}", C["err"])
                ))
        threading.Thread(target=_worker, daemon=True).start()

    def _desconectar(self):
        self._drive.desconectar()
        self._conectado = False
        self._dot.config(text="● Desconectado", fg=C["err"])
        self._set_status("Sesión cerrada", C["dim"])

    def _requiere_conexion(self) -> bool:
        if not self._conectado:
            messagebox.showwarning("Sin conexión", "Primero conectate a Google Drive.")
            return False
        return True

    def _listar_docx(self):
        if not self._requiere_conexion(): return
        self._set_status("🔄 Buscando archivos…", C["warn"])
        carpeta = self._var_carpeta_doc.get().strip() or None
        def _worker():
            try:
                archivos = self._drive.listar_docx(carpeta)
                def _update():
                    self._docx_list.delete(0, "end")
                    for f in archivos:
                        self._docx_list.insert("end", f["name"])
                        self._docx_list.itemconfig("end", fg=C["text"])
                    self._docx_list._archivos = archivos  # guardar metadatos
                    self._set_status(f"{len(archivos)} archivo(s) encontrado(s)", C["ok"])
                self.after(0, _update)
            except Exception as ex:
                self.after(0, lambda m=str(ex): self._set_status(f"❌ {m}", C["err"]))
        threading.Thread(target=_worker, daemon=True).start()

    def _abrir_docx_drive(self):
        if not self._requiere_conexion(): return
        sel = self._docx_list.curselection()
        if not sel:
            messagebox.showinfo("Seleccioná", "Seleccioná un archivo de la lista.")
            return
        archivos = getattr(self._docx_list, "_archivos", [])
        if not archivos or sel[0] >= len(archivos): return
        archivo = archivos[sel[0]]
        self._set_status(f"🔄 Descargando {archivo['name']}…", C["warn"])
        def _worker():
            try:
                data = self._drive.descargar_bytes(archivo["id"])
                # Guardar en temp y abrir
                tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                tmp.write(data); tmp.close()
                self.after(0, lambda p=tmp.name, n=archivo["name"]: (
                    self._on_abrir_texto(p, n),
                    self._set_status(f"✅ {n} abierto", C["ok"])
                ))
            except Exception as ex:
                self.after(0, lambda m=str(ex): self._set_status(f"❌ {m}", C["err"]))
        threading.Thread(target=_worker, daemon=True).start()

    def _cargar_imagenes_drive(self):
        if not self._requiere_conexion(): return
        carpeta_input = self._var_carpeta_img.get().strip()
        if not carpeta_input:
            messagebox.showinfo("Carpeta", "Ingresá el nombre o ID de la carpeta de imágenes.")
            return
        self._lbl_img_status.config(text="🔄 Buscando…", fg=C["warn"])
        def _worker():
            try:
                # Intentar como ID directo, si falla buscar por nombre
                carpeta_id = carpeta_input
                if not re.match(r"^[a-zA-Z0-9_-]{25,}$", carpeta_input):
                    resultados = self._drive.buscar_carpeta(carpeta_input)
                    if not resultados:
                        raise RuntimeError(f"No se encontró la carpeta '{carpeta_input}'")
                    carpeta_id = resultados[0]["id"]

                archivos = self._drive.listar_imagenes(carpeta_id)
                if not archivos:
                    raise RuntimeError("No hay imágenes en esa carpeta")

                self.after(0, lambda: self._lbl_img_status.config(
                    text=f"⬇ Descargando {len(archivos)} imagen(es)…", fg=C["warn"]))

                imgs = []
                for i, f in enumerate(archivos):
                    data = self._drive.descargar_bytes(f["id"])
                    imgs.append((f["name"], data))
                    self.after(0, lambda c=i+1, t=len(archivos):
                               self._lbl_img_status.config(
                                   text=f"⬇ {c}/{t}…", fg=C["warn"]))

                def _done():
                    self._lbl_img_status.config(
                        text=f"✅ {len(imgs)} imagen(es) cargada(s)", fg=C["ok"])
                    self._on_imagenes_drive(imgs)
                self.after(0, _done)
            except Exception as ex:
                self.after(0, lambda m=str(ex):
                           self._lbl_img_status.config(text=f"❌ {m}", fg=C["err"]))
        threading.Thread(target=_worker, daemon=True).start()

    def _subir(self):
        if not self._requiere_conexion(): return
        path = filedialog.askopenfilename(
            title="Seleccioná el archivo corregido",
            filetypes=[("Texto", "*.txt"), ("Markdown", "*.md"), ("Todos", "*.*")])
        if not path: return
        carpeta_id = self._var_carpeta_sub.get().strip() or None
        self._lbl_sub_status.config(text="🔄 Subiendo…", fg=C["warn"])
        def _worker():
            try:
                f = self._drive.subir_archivo(path, carpeta_id)
                self.after(0, lambda n=f.get("name","?"):
                           self._lbl_sub_status.config(
                               text=f"✅ Subido: {n}", fg=C["ok"]))
            except Exception as ex:
                self.after(0, lambda m=str(ex):
                           self._lbl_sub_status.config(text=f"❌ {m}", fg=C["err"]))
        threading.Thread(target=_worker, daemon=True).start()


# ═══════════════════════════════════════════════════════════════════════════════
#  APP PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════
class CorrectorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("✏ Corrector de Traducciones v7.0 — BloomScans")
        self.root.configure(bg=C["bg"])
        self.root.minsize(960, 600)
        self.root.geometry("1200x720")
        self._archivo_path    = ""
        self._texto_original  = ""
        self._errores_por_linea = {}
        self._errores_tira    = []
        self._busqueda_indices = []
        self._busqueda_actual  = 0
        self._nombres_propios  = set()
        self._visor = None        # VisorManwha
        self._panel_drive = None  # PanelDrive
        # Vinculaciones párrafo ↔ región imagen
        self._vinculaciones   = {}
        self._vincular_activo = False
        self._vincular_linea  = None
        self._captura_activa  = False
        self._captura_start   = None
        self._captura_rect_id = None
        self._captura_nlin    = 1
        # Labels — se asignan en _build_* pero los pre-inicializamos para
        # evitar AttributeError si _actualizar_vista corre antes
        self._lbl_count   = None
        self._lbl_pos     = None
        self._lbl_estado  = None
        self._lbl_archivo = None
        self.err_list     = None
        self._build_styles()
        self._build_ui()

    # ── Estilos ───────────────────────────────────────────────────────────────
    def _build_styles(self):
        s = ttk.Style(); s.theme_use("clam")
        for o in ("Vertical", "Horizontal"):
            s.configure(f"{o}.TScrollbar", gripcount=0, background=C["bg4"],
                        troughcolor=C["bg2"], bordercolor=C["bg2"],
                        arrowcolor=C["dim"], relief="flat")
        s.configure("TNotebook", background=C["bg"], borderwidth=0)
        s.configure("TNotebook.Tab", background=C["bg3"], foreground=C["text2"],
                    padding=[10, 4], font=("Segoe UI", 8))
        s.map("TNotebook.Tab",
              background=[("selected", C["bg2"])],
              foreground=[("selected", C["teal"])])

    # ── UI ────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        self._build_topbar()
        self._build_infobar()
        sep(self.root)
        self._build_buscador()
        sep(self.root)
        self._build_main_panel()
        self._build_statusbar()
        # Todo construido — ahora sí poblamos el editor con el texto de bienvenida
        self._set_welcome()

    def _build_topbar(self):
        bar = tk.Frame(self.root, bg=C["bg2"]); bar.pack(fill="x")
        tk.Label(bar, text="✏  Corrector de Traducciones",
                 bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 12, "bold")).pack(side="left", padx=16, pady=10)
        tk.Label(bar, text="v7.0  ·  BloomScans",
                 bg=C["bg2"], fg=C["dim"], font=("Segoe UI", 8)).pack(side="left")
        bf = tk.Frame(bar, bg=C["bg2"]); bf.pack(side="right", padx=10, pady=7)
        btns = [
            ("🤖  IA Gemini",       self._abrir_panel_ia,         C["purple"], C["bright"], C["bg5"],   C["purple"]),
            ("📂  Abrir",         self._abrir,                  C["bg4"],   C["text"],   C["bg5"],   C["bright"]),
            ("🔍  Analizar",      self._analizar,               C["bg4"],   C["text"],   C["bg5"],   C["bright"]),
            ("✨  Auto-corregir", self._auto_corregir,          C["teal3"], C["bright"], C["teal2"], C["bright"]),
            ("´  Auto-tilde",    self._auto_tilde,              C["teal3"], C["bright"], C["teal2"], C["bright"]),
            ("📊  Stats",         self._mostrar_stats,          C["bg4"],   C["cyan"],   C["bg5"],   C["cyan"]),
            ("🔀  Comparar",      self._comparar,               C["bg4"],   C["purple"], C["bg5"],   C["purple"]),
            ("📋  Reporte",       self._exportar_reporte,       C["bg4"],   C["warn"],   C["bg5"],   C["warn"]),
            ("📄  Word",          self._exportar_word_errores,  C["bg4"],   "#ff60a0",   C["bg5"],   "#ff80b0"),
            ("📝  G.Docs",        self._exportar_gdocs_errores, C["bg4"],   "#4dc0a0",   C["bg5"],   "#6ddfc0"),
            ("🖼  Visor",         self._abrir_visor,            C["bg4"],   C["teal"],   C["bg5"],   C["teal"]),
            ("☁  Drive",         self._abrir_drive,            C["bg4"],   C["cyan"],   C["bg5"],   C["cyan"]),
            ("💾  Guardar",       self._guardar,                C["teal"],  C["bg"],     C["cyan"],  C["bg"]),
        ]
        for i, (txt, cmd, bg, fg, hbg, hfg) in enumerate(btns):
            b = make_btn(bf, txt, cmd, bg=bg, fg=fg, hover_bg=hbg, hover_fg=hfg, small=True)
            b.pack(side="left", padx=2)
            if i == 0:
                self._btn_ia = b  # guardar referencia al botón IA
        sep(bar)

    def _build_infobar(self):
        bar = tk.Frame(self.root, bg=C["bg"]); bar.pack(fill="x", padx=12, pady=(5,3))
        self._lbl_archivo = tk.Label(bar, text="Sin archivo abierto",
                                     bg=C["bg"], fg=C["dim"], font=("Segoe UI", 9))
        self._lbl_archivo.pack(side="left")
        self._lbl_estado = tk.Label(bar, text="", bg=C["bg"], fg=C["ok"],
                                    font=("Segoe UI", 9))
        self._lbl_estado.pack(side="right")

    def _build_buscador(self):
        bar = tk.Frame(self.root, bg=C["bg2"]); bar.pack(fill="x")
        tk.Label(bar, text="🔎", bg=C["bg2"], fg=C["dim"],
                 font=("Segoe UI", 10)).pack(side="left", padx=(10,4), pady=4)
        self._buscar_var = tk.StringVar()
        self._buscar_entry = tk.Entry(bar, textvariable=self._buscar_var,
                                      bg=C["bg4"], fg=C["text"],
                                      insertbackground=C["teal"],
                                      relief="flat", font=("Segoe UI", 9), bd=0)
        self._buscar_entry.pack(side="left", fill="x", expand=True, pady=4, ipady=3)
        self._buscar_entry.bind("<Return>",       self._siguiente_busqueda)
        self._buscar_entry.bind("<Shift-Return>", self._anterior_busqueda)
        self._buscar_var.trace_add("write", lambda *_: self._buscar())
        self._lbl_buscar_res = tk.Label(bar, text="", bg=C["bg2"], fg=C["dim"],
                                        font=("Segoe UI", 8))
        self._lbl_buscar_res.pack(side="left", padx=8)
        make_btn(bar, "✕", self._limpiar_busqueda,
                 bg=C["bg2"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["err"],
                 small=True).pack(side="left", padx=6)
        # Config longitud máxima de línea
        tk.Label(bar, text=" | Largo máx:", bg=C["bg2"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(side="left")
        self._var_maxlen = tk.IntVar(value=MAX_CHARS_LINEA)
        sp = tk.Spinbox(bar, from_=40, to=300, textvariable=self._var_maxlen,
                        width=4, bg=C["bg4"], fg=C["text"],
                        insertbackground=C["teal"], relief="flat",
                        buttonbackground=C["bg5"], font=("Segoe UI", 8))
        sp.pack(side="left", padx=4, pady=4)

    def _build_main_panel(self):
        # Horizontal: [editor+errores  |  visor de imagen  |  panel IA (colapsable)]
        self._h_paned = tk.PanedWindow(self.root, orient="horizontal",
                                       bg=C["border"], sashwidth=6,
                                       sashrelief="flat", handlesize=0)
        self._h_paned.pack(fill="both", expand=True)

        # Left side: vertical split [editor | errores+simbolos]
        left = tk.Frame(self._h_paned, bg=C["bg"])
        self._v_paned = tk.PanedWindow(left, orient="vertical",
                                       bg=C["border"], sashwidth=5,
                                       sashrelief="flat", handlesize=0)
        self._v_paned.pack(fill="both", expand=True)
        self._v_paned.add(self._build_editor(self._v_paned),
                          minsize=200, stretch="always")
        self._v_paned.add(self._build_sidebar(self._v_paned),
                          minsize=120, stretch="never")
        self._h_paned.add(left, minsize=380, stretch="always")

        # Center: integrated image viewer
        self._h_paned.add(self._build_visor_integrado(self._h_paned),
                          minsize=0, stretch="never")

        # Right: IA panel (starts hidden — toggled with the 🤖 button)
        self._panel_ia = PanelGeminiIA(
            self._h_paned,
            get_texto_fn=lambda: self.editor.get("1.0", "end-1c"),
            get_script_tira_fn=self._ia_get_script_tira,
            get_visor_fn=self._ia_get_imagen_actual,
            get_todas_imagenes_fn=self._ia_get_todas_imagenes,
            get_tira_fn=self._ia_get_tira_actual,
            aplicar_correccion_fn=self._ia_aplicar_correccion,
            get_nombres_propios_fn=lambda: self._nombres_propios,
        )
        self._ia_visible = False  # starts hidden

    def _build_editor(self, parent):
        frame = tk.Frame(parent, bg=C["bg"])
        tk.Label(frame, text="  Texto de traducción",
                 bg=C["bg3"], fg=C["text2"],
                 font=("Segoe UI", 8, "bold"), anchor="w").pack(fill="x")

        # MEJORA 8b: Panel de símbolos clickeables ────────────────────────────
        sym_bar = tk.Frame(frame, bg=C["bg3"])
        sym_bar.pack(fill="x")
        tk.Label(sym_bar, text=" Insertar:", bg=C["bg3"], fg=C["dim"],
                 font=("Segoe UI", 7)).pack(side="left", padx=(6,2), pady=3)

        # Símbolos frecuentes con su texto a insertar
        _SYM_BTNS = [
            ("D2",  "D2: "), ("P2",  "P2: "), ("T2",  "T2: "),
            ("G",   "G: "),  ("P",   "P: "),  ("T",   "T: "),
            ("N",   "N: "),  ("DF",  "DF: "), ("()",  "(): "),
            ("NT",  "NT: "), ("{}",  "{}: "), ("<>",  "<>: "),
            ("/",   "/: "),  ("[]",  "[]: "), ("##",  "\n\\## Tira "),
        ]
        for label, insertar in _SYM_BTNS:
            ins = insertar  # capture
            b = tk.Button(sym_bar, text=label, font=("Consolas", 7, "bold"),
                          bg=C["bg4"], fg=C["cyan"],
                          activebackground=C["teal3"], activeforeground=C["bright"],
                          relief="flat", bd=0, padx=5, pady=2, cursor="hand2",
                          command=lambda s=ins: self._insertar_simbolo(s))
            b.bind("<Enter>", lambda e, w=b: w.config(bg=C["teal3"], fg=C["bright"]))
            b.bind("<Leave>", lambda e, w=b: w.config(bg=C["bg4"],   fg=C["cyan"]))
            b.pack(side="left", padx=1, pady=3)

        # Botón marcar error manual
        make_btn(sym_bar, "⚑ Error manual", self._marcar_error_manual,
                 bg=C["bg4"], fg=C["err"], hover_bg=C["bg5"], hover_fg=C["err"],
                 small=True).pack(side="right", padx=4, pady=2)
        make_btn(sym_bar, "📸 Captura→Comentario", self._captura_a_comentario,
                 bg=C["bg4"], fg=C["warn"], hover_bg=C["bg5"], hover_fg=C["warn"],
                 small=True).pack(side="right", padx=2, pady=2)

        sep(frame)
        inner = tk.Frame(frame, bg=C["bg"]); inner.pack(fill="both", expand=True)
        sbv = ttk.Scrollbar(inner, orient="vertical",   style="Vertical.TScrollbar")
        sbh = ttk.Scrollbar(inner, orient="horizontal", style="Horizontal.TScrollbar")

        # Números de línea (canvas)
        self._ln_canvas = tk.Canvas(inner, bg=C["bg3"], width=44,
                                    highlightthickness=0, bd=0)
        self._ln_canvas.pack(side="left", fill="y")

        self.editor = tk.Text(
            inner, bg=C["bg2"], fg=C["text"], insertbackground=C["teal"],
            selectbackground=C["teal3"], selectforeground=C["bright"],
            font=("Consolas", 10), wrap="none", relief="flat", bd=0,
            padx=10, pady=8, yscrollcommand=sbv.set, xscrollcommand=sbh.set, undo=True)
        sbv.config(command=self._editor_yview_sync)
        sbh.config(command=self.editor.xview)
        sbv.pack(side="right", fill="y"); sbh.pack(side="bottom", fill="x")
        self.editor.pack(side="left", fill="both", expand=True)

        # Tags
        self.editor.tag_config("encabezado", foreground=C["teal"],
                               font=("Consolas", 10, "bold"))
        self.editor.tag_config("prefijo", foreground=C["cyan"])
        for tipo, bg_k, fg_k in [
            ("puntuacion","err_punt_bg","err_punt_fg"),
            ("mayuscula", "err_may_bg", "err_may_fg"),
            ("simbolo",   "err_sim_bg", "err_sim_fg"),
            ("tilde",     "err_til_bg", "err_til_fg"),
            ("estructura","err_dup_bg", "err_dup_fg"),
            ("duplicado", "err_dup_bg", "err_dup_fg"),
            ("largo",     "err_long_bg","err_long_fg"),
            ("nombre",    "err_may_bg", "err_may_fg"),
            ("tira",      "err_sim_bg", "err_sim_fg"),
            ("ortografia","err_ort_bg", "err_ort_fg"),
            ("mezcla",    "err_mix_bg", "err_mix_fg"),
        ]:
            self.editor.tag_config(f"err_{tipo}",
                                   background=C[bg_k], foreground=C[fg_k],
                                   underline=True)
        self.editor.tag_config("err_manual",
                               background="#2d1020", foreground="#ff60a0",
                               underline=True)
        self.editor.tag_config("err_comentario",
                               background="#0d1a2a", foreground="#60c0ff",
                               underline=False, font=("Consolas", 10, "italic"))
        self.editor.tag_config("busqueda",
                               background=C["search_bg"], foreground=C["search_fg"])
        self.editor.tag_config("busqueda_actual",
                               background="#ffbb00", foreground="#000000")
        self.editor.tag_config("err_linea_activa",
                               background="#2a1f00", foreground=C["bright"],
                               relief="raised", borderwidth=1)
        # err_linea_activa debe estar por encima de los tags de error
        self.editor.tag_raise("err_linea_activa")
        self.editor.tag_config("vinculado",
                               underline=True, foreground=C["cyan"])
        self.editor.bind("<ButtonRelease-1>", self._on_editor_click)
        self.editor.bind("<KeyRelease>",      self._on_key_update)
        self.editor.bind("<ButtonRelease>",   self._actualizar_posicion)
        self.editor.bind("<Configure>",       lambda e: self._actualizar_numeros_linea())
        # Inicializar almacén de errores manuales y comentarios
        self._errores_manuales  = {}   # {num_linea: "descripción"}
        self._comentarios_img   = {}   # {num_linea: {"texto": str, "img_b64": str|None}}
        self._errores_ignorados = {}   # {num_linea: {tipo, ...}}  errores suprimidos por el usuario
        return frame

    def _editor_yview_sync(self, *args):
        """Sincroniza el scroll del editor con los números de línea."""
        self.editor.yview(*args)
        self._actualizar_numeros_linea()

    def _actualizar_numeros_linea(self, _=None):
        """Dibuja los números de línea en el canvas lateral."""
        if not hasattr(self, "_ln_canvas") or not hasattr(self, "editor"):
            return
        c = self._ln_canvas
        c.delete("all")
        # Obtener rango visible
        try:
            first_line = int(self.editor.index("@0,0").split(".")[0])
            last_line  = int(self.editor.index(f"@0,{self.editor.winfo_height()}").split(".")[0])
        except Exception:
            return
        # Altura de la fuente
        try:
            fh = self.editor.dlineinfo(f"{first_line}.0")
        except Exception:
            fh = None
        for n in range(first_line, last_line + 2):
            try:
                info = self.editor.dlineinfo(f"{n}.0")
                if info is None:
                    break
                y = info[1]  # y pixel dentro del widget texto
                # Colorear diferente si tiene error manual o comentario
                col = C["dim"]
                if hasattr(self, "_errores_manuales") and n in self._errores_manuales:
                    col = "#ff60a0"
                elif hasattr(self, "_comentarios_img") and n in self._comentarios_img:
                    col = "#60c0ff"
                elif hasattr(self, "_errores_por_linea") and n in self._errores_por_linea:
                    col = C["warn"]
                c.create_text(38, y + 9, text=str(n), anchor="e",
                              fill=col, font=("Consolas", 9))
            except Exception:
                break

    def _build_sidebar(self, parent):
        frame = tk.Frame(parent, bg=C["bg"])
        nb = ttk.Notebook(frame); nb.pack(fill="both", expand=True)

        # Tab Errores
        tab_err = tk.Frame(nb, bg=C["bg"]); nb.add(tab_err, text="⚠ Errores")
        tk.Label(tab_err, text="  Clic → ir · Clic derecho → ignorar error",
                 bg=C["bg3"], fg=C["dim"], font=("Segoe UI", 7), anchor="w").pack(fill="x")
        eo = tk.Frame(tab_err, bg=C["bg"]); eo.pack(fill="both", expand=True)
        esb = ttk.Scrollbar(eo, orient="vertical", style="Vertical.TScrollbar")
        self.err_list = tk.Listbox(eo, bg=C["bg2"], fg=C["text"],
                                   selectbackground=C["teal3"],
                                   selectforeground=C["bright"],
                                   font=("Segoe UI", 8), relief="flat", bd=0,
                                   activestyle="none", yscrollcommand=esb.set)
        esb.config(command=self.err_list.yview)
        esb.pack(side="right", fill="y"); self.err_list.pack(side="left", fill="both", expand=True)
        self.err_list.bind("<<ListboxSelect>>", self._ir_a_error)
        self.err_list.bind("<Button-3>", self._err_list_context_menu)

        # ── Leyenda colapsable ─────────────────────────────────────────────────
        self._legend_visible = False   # empieza oculta
        legend_header = tk.Frame(tab_err, bg=C["bg3"]); legend_header.pack(fill="x", pady=(2,0))
        self._btn_legend = tk.Button(
            legend_header, text="▶ Referencia de colores",
            font=("Segoe UI", 7), bg=C["bg3"], fg=C["dim"],
            relief="flat", bd=0, cursor="hand2", anchor="w",
            activebackground=C["bg4"], activeforeground=C["text"],
            command=self._toggle_legend)
        self._btn_legend.pack(fill="x", padx=6, pady=2)

        self._legend_frame = tk.Frame(tab_err, bg=C["bg3"])
        # No se hace .pack() aqui — empieza oculto
        for color, label in [
            (C["err_punt_fg"], "Puntuacion / Signos"),
            (C["err_may_fg"],  "Mayuscula / Nombre"),
            (C["err_sim_fg"],  "Simbolo / Tira"),
            (C["err_til_fg"],  "Tilde faltante"),
            (C["err_dup_fg"],  "Estructura / Duplicado"),
            (C["err_long_fg"], "Linea muy larga"),
            (C["err_ort_fg"],  "Ortografia (typo)"),
            (C["err_mix_fg"],  "Mezcla de sistemas"),
        ]:
            row = tk.Frame(self._legend_frame, bg=C["bg3"]); row.pack(fill="x", padx=8, pady=1)
            tk.Label(row, text="■", bg=C["bg3"], fg=color,
                     font=("Segoe UI", 9)).pack(side="left")
            tk.Label(row, text=f"  {label}", bg=C["bg3"], fg=C["dim"],
                     font=("Segoe UI", 8)).pack(side="left")

        # Tab Símbolos
        tab_sym = tk.Frame(nb, bg=C["bg"]); nb.add(tab_sym, text="📖 Símbolos")
        o, t = scroll_text(tab_sym); o.pack(fill="both", expand=True)
        t.tag_config("key", foreground=C["cyan"], font=("Consolas", 8, "bold"))
        t.tag_config("cat", foreground=C["teal"], font=("Segoe UI", 8, "bold"))
        t.config(state="normal")
        cats = {}
        for sym, (desc, cat) in SIMBOLOS.items():
            cats.setdefault(cat, []).append((sym, desc))
        for cat, items in cats.items():
            t.insert("end", f"\n── {cat} ──\n", ("cat",))
            for sym, desc in items:
                t.insert("end", sym.strip(), ("key",))
                t.insert("end", f"   {desc}\n")
        t.insert("end", "\n── Especiales ──\n", ("cat",))
        t.insert("end", "(vacío)", ("key",)); t.insert("end", "   Diálogo / Gemidos\n")
        t.insert("end", "## N", ("key",)); t.insert("end", "   Salto de página\n")
        t.config(state="disabled")

        # Tab Tildes
        tab_til = tk.Frame(nb, bg=C["bg"]); nb.add(tab_til, text="´ Tildes")
        o2, t2 = scroll_text(tab_til); o2.pack(fill="both", expand=True)
        t2.tag_config("mal",  foreground="#ff8080", font=("Consolas", 8))
        t2.tag_config("bien", foreground="#80ff80", font=("Consolas", 8))
        t2.tag_config("arr",  foreground=C["border2"])
        t2.config(state="normal")
        t2.insert("end", "Palabras sin tilde detectadas:\n\n")
        seen = set()
        for sin, con in TILDES:
            if sin not in seen:
                seen.add(sin); t2.insert("end", sin, ("mal",))
                t2.insert("end", "  →  ", ("arr",)); t2.insert("end", con + "\n", ("bien",))
        t2.config(state="disabled")

        return frame


    def _build_visor_integrado(self, parent):
        """Panel del visor de imagen integrado al layout principal."""
        frame = tk.Frame(parent, bg=C["bg"])
        self._visor_frame = frame   # para mostrar/ocultar

        # Barra superior del visor
        vbar = tk.Frame(frame, bg=C["bg2"]); vbar.pack(fill="x")
        tk.Label(vbar, text="🖼  Imagen", bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 9, "bold")).pack(side="left", padx=8, pady=5)

        # Botones de fuente
        bf = tk.Frame(vbar, bg=C["bg2"]); bf.pack(side="left", padx=4)
        make_btn(bf, "📁", self._visor_abrir_carpeta,
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="left", padx=1)
        make_btn(bf, "🗜", self._visor_abrir_zip,
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="left", padx=1)
        make_btn(bf, "☁", self._abrir_drive,
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                 small=True).pack(side="left", padx=1)

        # Botón vincular párrafo ↔ región
        self._vincular_activo  = False
        self._vincular_linea   = None   # num de línea del editor a vincular
        self._vinculaciones    = {}     # {num_linea: (x1,y1,x2,y2,img_idx)}
        self._rect_drag_start  = None
        self._rect_drag_id     = None
        self._btn_vincular = make_btn(
            bf, "📌 Vincular",
            self._toggle_modo_vincular,
            bg=C["bg3"], fg=C["warn"], hover_bg=C["bg4"], hover_fg=C["warn"],
            small=True)
        self._btn_vincular.pack(side="left", padx=(8,1))
        self._lbl_vincular_hint = tk.Label(
            bf, text="", bg=C["bg2"], fg=C["warn"], font=("Segoe UI", 7))
        self._lbl_vincular_hint.pack(side="left", padx=4)

        # Navegación
        nav = tk.Frame(vbar, bg=C["bg2"]); nav.pack(side="right", padx=6)
        make_btn(nav, "◀◀", lambda: self._visor_ir(0),
                 bg=C["bg2"], fg=C["dim"], hover_bg=C["bg3"], hover_fg=C["teal"],
                 small=True).pack(side="left")
        make_btn(nav, "◀", self._visor_anterior,
                 bg=C["bg2"], fg=C["text"], hover_bg=C["bg3"], hover_fg=C["teal"],
                 small=True).pack(side="left")
        self._lbl_visor_idx = tk.Label(nav, text="Sin imágenes", bg=C["bg2"],
                                       fg=C["text2"], font=("Segoe UI", 8))
        self._lbl_visor_idx.pack(side="left", padx=6)
        make_btn(nav, "▶", self._visor_siguiente,
                 bg=C["bg2"], fg=C["text"], hover_bg=C["bg3"], hover_fg=C["teal"],
                 small=True).pack(side="left")
        make_btn(nav, "▶▶", lambda: self._visor_ir(len(self._visor_imagenes)-1),
                 bg=C["bg2"], fg=C["dim"], hover_bg=C["bg3"], hover_fg=C["teal"],
                 small=True).pack(side="left")

        sep(frame)

        # Barra de zoom
        zbar = tk.Frame(frame, bg=C["bg3"]); zbar.pack(fill="x")
        tk.Label(zbar, text="Zoom:", bg=C["bg3"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(side="left", padx=(8,4), pady=3)
        self._visor_zoom = 1.0
        for z, lbl in [(0.25,"25%"),(0.5,"50%"),(0.75,"75%"),
                       (1.0,"100%"),(1.5,"150%"),(2.0,"200%"),(3.0,"300%")]:
            zv = z
            make_btn(zbar, lbl, lambda v=zv: self._visor_set_zoom(v),
                     bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["teal"],
                     small=True).pack(side="left", padx=1, pady=2)
        make_btn(zbar, "+", self._visor_zoom_in,
                 bg=C["bg3"], fg=C["teal"], hover_bg=C["bg4"], hover_fg=C["cyan"],
                 small=True).pack(side="right", padx=2, pady=2)
        make_btn(zbar, "−", self._visor_zoom_out,
                 bg=C["bg3"], fg=C["teal"], hover_bg=C["bg4"], hover_fg=C["cyan"],
                 small=True).pack(side="right", padx=1, pady=2)
        self._lbl_zoom_pct = tk.Label(zbar, text="100%", bg=C["bg3"], fg=C["teal"],
                                      font=("Segoe UI", 8, "bold"))
        self._lbl_zoom_pct.pack(side="right", padx=6)

        sep(frame)

        # Canvas con scrollbars
        cv_frame = tk.Frame(frame, bg=C["bg"]); cv_frame.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(cv_frame, orient="vertical",   style="Vertical.TScrollbar")
        hsb = ttk.Scrollbar(cv_frame, orient="horizontal", style="Horizontal.TScrollbar")
        self._visor_canvas = tk.Canvas(cv_frame, bg=C["bg2"], highlightthickness=0,
                                       cursor="crosshair",
                                       yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        vsb.config(command=self._visor_canvas.yview)
        hsb.config(command=self._visor_canvas.xview)
        vsb.pack(side="right", fill="y"); hsb.pack(side="bottom", fill="x")
        self._visor_canvas.pack(side="left", fill="both", expand=True)

        # Rueda sola = scroll vertical; Ctrl+rueda = zoom
        self._visor_canvas.bind("<MouseWheel>",
            lambda e: (self._visor_scroll_zoom(e.delta)
                       if (e.state & 0x4) else
                       self._visor_canvas.yview_scroll(int(-1*(e.delta/120)), "units")))
        self._visor_canvas.bind("<Button-4>",
            lambda e: (self._visor_scroll_zoom(120)
                       if (e.state & 0x4) else
                       self._visor_canvas.yview_scroll(-1, "units")))
        self._visor_canvas.bind("<Button-5>",
            lambda e: (self._visor_scroll_zoom(-120)
                       if (e.state & 0x4) else
                       self._visor_canvas.yview_scroll(1, "units")))
        # Ctrl+rueda explícito (Windows duplica el evento, no importa)
        self._visor_canvas.bind("<Control-MouseWheel>",
            lambda e: self._visor_scroll_zoom(e.delta))
        # Drag para vincular región o capturar comentario
        self._visor_canvas.bind("<ButtonPress-1>",   self._visor_drag_or_capture_start)
        self._visor_canvas.bind("<B1-Motion>",        self._visor_drag_or_capture_move)
        self._visor_canvas.bind("<ButtonRelease-1>",  self._visor_drag_or_capture_end)

        # Mensaje inicial
        # Mensaje inicial
        self._visor_canvas.create_text(160, 120,
            text="Sin imagenes\n\nAbri imagenes con carpeta, ZIP o Drive",
            fill=C["dim"], font=("Segoe UI", 10), anchor="center",
            justify="center", tags="placeholder")

        # Estado interno del visor
        self._visor_imagenes = []   # [(nombre, bytes_o_path), ...]
        self._visor_idx      = 0
        self._visor_tk_img   = None

        return frame

    # ── Métodos del visor integrado ────────────────────────────────────────────
    def _visor_abrir_carpeta(self):
        if not PIL_OK:
            self._pedir_pillow(); return
        carpeta = filedialog.askdirectory(title="Carpeta con imágenes del manwha")
        if not carpeta: return
        exts = {".jpg",".jpeg",".png",".webp",".bmp",".gif"}
        archivos = sorted(f for f in Path(carpeta).iterdir()
                          if f.suffix.lower() in exts)
        self._visor_imagenes = [(f.name, str(f)) for f in archivos]
        self._visor_idx = 0
        self._visor_mostrar()

    def _visor_abrir_zip(self):
        if not PIL_OK:
            self._pedir_pillow(); return
        path = filedialog.askopenfilename(
            title="ZIP con imágenes",
            filetypes=[("ZIP","*.zip"),("Todos","*.*")])
        if not path: return
        import zipfile
        exts = {".jpg",".jpeg",".png",".webp",".bmp"}
        imgs = []
        with zipfile.ZipFile(path) as z:
            for nombre in sorted(z.namelist()):
                if Path(nombre).suffix.lower() in exts:
                    imgs.append((Path(nombre).name, z.read(nombre)))
        self._visor_imagenes = imgs
        self._visor_idx = 0
        self._visor_mostrar()

    def cargar_imagenes_en_visor(self, imgs):
        """Llamado desde Drive: [(nombre, bytes)]."""
        self._visor_imagenes = imgs
        self._visor_idx = 0
        self._visor_mostrar()

    def _pedir_pillow(self):
        if messagebox.askyesno("Pillow necesario",
                               "El visor necesita Pillow.\n¿Instalar ahora?"):
            import subprocess, sys
            subprocess.run([sys.executable,"-m","pip","install","Pillow"],
                           capture_output=True)
            messagebox.showinfo("Pillow","Instalado. Reiniciá el corrector.")

    def _visor_ir(self, idx):
        self._visor_idx = max(0, min(idx, len(self._visor_imagenes)-1))
        self._visor_mostrar()

    def _visor_anterior(self):
        self._visor_ir(self._visor_idx - 1)

    def _visor_siguiente(self):
        self._visor_ir(self._visor_idx + 1)

    def _visor_set_zoom(self, z):
        self._visor_zoom = z
        self._lbl_zoom_pct.config(text=f"{int(z*100)}%")
        self._visor_mostrar(skip_reload=True)

    def _visor_zoom_in(self):
        self._visor_set_zoom(min(self._visor_zoom * 1.25, 5.0))

    def _visor_zoom_out(self):
        self._visor_set_zoom(max(self._visor_zoom / 1.25, 0.1))

    def _visor_scroll_zoom(self, delta):
        """Ctrl+rueda o rueda directa = zoom."""
        if delta > 0:
            self._visor_zoom_in()
        else:
            self._visor_zoom_out()

    def _visor_mostrar(self, skip_reload=False):
        if not self._visor_imagenes or not PIL_OK:
            return
        nombre, fuente = self._visor_imagenes[self._visor_idx]
        self._lbl_visor_idx.config(
            text=f"{nombre}  ({self._visor_idx+1}/{len(self._visor_imagenes)})")
        self._lbl_zoom_pct.config(text=f"{int(self._visor_zoom*100)}%")

        import io
        if not skip_reload or not hasattr(self, "_visor_img_orig"):
            if isinstance(fuente, bytes):
                self._visor_img_orig = Image.open(io.BytesIO(fuente))
            else:
                self._visor_img_orig = Image.open(fuente)

        img = self._visor_img_orig
        w, h = img.size
        nw, nh = max(1,int(w*self._visor_zoom)), max(1,int(h*self._visor_zoom))
        img_res = img.resize((nw, nh), Image.LANCZOS)
        self._visor_tk_img = ImageTk.PhotoImage(img_res)
        self._visor_canvas.delete("all")
        self._visor_canvas.config(scrollregion=(0,0,nw,nh))
        self._visor_canvas.create_image(0, 0, anchor="nw", image=self._visor_tk_img)
        # Volver al tope al cambiar de imagen (solo si NO es solo zoom)
        if not skip_reload:
            self._visor_canvas.yview_moveto(0)
            self._visor_canvas.xview_moveto(0)

        # Sync con editor al cambiar de imagen
        self._visor_navego(self._visor_idx)

    # ── Vinculación párrafo ↔ región de imagen ────────────────────────────────
    def _toggle_modo_vincular(self):
        """Activa/desactiva el modo de vincular línea del editor con zona de imagen."""
        if not self._vincular_activo:
            # Averiguar qué línea está seleccionada en el editor
            try:
                pos = self.editor.index("insert")
                nlin = int(pos.split(".")[0])
            except Exception:
                nlin = None
            if not nlin:
                return
            self._vincular_activo = True
            self._vincular_linea  = nlin
            self._btn_vincular.config(bg=C["warn"], fg=C["bg"],
                                      text="📌 Cancelar")
            self._lbl_vincular_hint.config(
                text=f"Arrastrá la zona del texto en la imagen (L{nlin})")
            self._visor_canvas.config(cursor="crosshair")
        else:
            self._vincular_activo = False
            self._vincular_linea  = None
            self._btn_vincular.config(bg=C["bg3"], fg=C["warn"],
                                      text="📌 Vincular")
            self._lbl_vincular_hint.config(text="")
            self._visor_canvas.config(cursor="crosshair")
            if self._rect_drag_id:
                self._visor_canvas.delete(self._rect_drag_id)
                self._rect_drag_id = None

    def _canvas_to_img_coords(self, cx, cy):
        """Convierte coordenadas de canvas (ya ajustadas por scroll) a coords de imagen sin zoom."""
        z = self._visor_zoom if self._visor_zoom > 0 else 1.0
        return cx / z, cy / z

    def _visor_drag_or_capture_start(self, event):
        if getattr(self, "_captura_activa", False):
            self._visor_captura_start(event)
        else:
            self._visor_drag_start(event)

    def _visor_drag_or_capture_move(self, event):
        if getattr(self, "_captura_activa", False):
            self._visor_captura_move(event)
        else:
            self._visor_drag_move(event)

    def _visor_drag_or_capture_end(self, event):
        if getattr(self, "_captura_activa", False):
            self._visor_captura_end(event)
        else:
            self._visor_drag_end(event)

    def _visor_drag_start(self, event):
        if not self._vincular_activo:
            return
        cx = self._visor_canvas.canvasx(event.x)
        cy = self._visor_canvas.canvasy(event.y)
        self._rect_drag_start = (cx, cy)
        if self._rect_drag_id:
            self._visor_canvas.delete(self._rect_drag_id)
            self._rect_drag_id = None

    def _visor_drag_move(self, event):
        if not self._vincular_activo or not self._rect_drag_start:
            return
        x0, y0 = self._rect_drag_start
        cx = self._visor_canvas.canvasx(event.x)
        cy = self._visor_canvas.canvasy(event.y)
        if self._rect_drag_id:
            self._visor_canvas.delete(self._rect_drag_id)
        self._rect_drag_id = self._visor_canvas.create_rectangle(
            x0, y0, cx, cy,
            outline=C["warn"], width=2, dash=(4, 2), tags="drag_rect")

    def _visor_drag_end(self, event):
        if not self._vincular_activo or not self._rect_drag_start:
            return
        x0, y0 = self._rect_drag_start
        x1 = self._visor_canvas.canvasx(event.x)
        y1 = self._visor_canvas.canvasy(event.y)
        if abs(x1 - x0) < 8 or abs(y1 - y0) < 8:
            if self._rect_drag_id:
                self._visor_canvas.delete(self._rect_drag_id)
                self._rect_drag_id = None
            return
        # Guardar en coordenadas de imagen original (sin zoom)
        ix0, iy0 = self._canvas_to_img_coords(x0, y0)
        ix1, iy1 = self._canvas_to_img_coords(x1, y1)
        nlin = self._vincular_linea
        img_idx = self._visor_idx
        self._vinculaciones[nlin] = (
            min(ix0,ix1), min(iy0,iy1),
            max(ix0,ix1), max(iy0,iy1),
            img_idx)
        # Marcar la línea del editor visualmente
        self.editor.tag_add("vinculado", f"{nlin}.0", f"{nlin}.end")
        # Desactivar modo vincular
        self._toggle_modo_vincular()
        # Mostrar la región guardada
        self._mostrar_region_vinculada(nlin)

    def _mostrar_region_vinculada(self, nlin: int):
        """Dibuja el rectángulo de la vinculación guardada sobre la imagen actual."""
        self._visor_canvas.delete("vinc_rect")
        if nlin not in self._vinculaciones:
            return
        ix0, iy0, ix1, iy1, img_idx = self._vinculaciones[nlin]
        if img_idx != self._visor_idx:
            # La vinculación es de otra imagen — navegar a ella
            self._visor_ir(img_idx)
            return
        z = self._visor_zoom
        cx0, cy0 = ix0 * z, iy0 * z
        cx1, cy1 = ix1 * z, iy1 * z
        self._visor_canvas.create_rectangle(
            cx0, cy0, cx1, cy1,
            outline=C["cyan"], width=2, tags="vinc_rect")
        self._visor_canvas.create_rectangle(
            cx0-1, cy0-1, cx1+1, cy1+1,
            outline=C["bg"], width=1, dash=(3,3), tags="vinc_rect")
        self._visor_canvas.see = lambda *a: None   # no-op guard
        # Scroll para que la región sea visible
        try:
            sr = self._visor_canvas.cget("scrollregion").split()
            if sr:
                tw, th = float(sr[2]), float(sr[3])
                if tw > 0: self._visor_canvas.xview_moveto(max(0, cx0/tw - 0.1))
                if th > 0: self._visor_canvas.yview_moveto(max(0, cy0/th - 0.1))
        except Exception:
            pass

    def _build_statusbar(self):
        sep(self.root)
        bar = tk.Frame(self.root, bg=C["bg2"]); bar.pack(fill="x", side="bottom")
        self._lbl_count = tk.Label(bar, text="", bg=C["bg2"], fg=C["dim"],
                                   font=("Segoe UI", 8))
        self._lbl_count.pack(side="left", padx=12, pady=4)
        self._lbl_pos = tk.Label(bar, text="", bg=C["bg2"], fg=C["dim"],
                                 font=("Segoe UI", 8))
        self._lbl_pos.pack(side="right", padx=12, pady=4)

    def _actualizar_posicion(self, _=None):
        pos = self.editor.index("insert")
        ln, col = pos.split(".")
        self._lbl_pos.config(text=f"Línea {ln}  Col {int(col)+1}")
        self._actualizar_numeros_linea()

    def _on_key_update(self, _=None):
        self._actualizar_posicion()
        self._actualizar_numeros_linea()

    # ── Sincronización visor + vinculaciones ─────────────────────────────────
    def _on_editor_click(self, event=None):
        """Clic en el editor: sincroniza visor con tira, y muestra región vinculada."""
        self._actualizar_posicion()
        try:
            pos  = self.editor.index("insert")
            nlin = int(pos.split(".")[0])
        except Exception:
            return
        # Mostrar región vinculada si existe para esta línea
        if hasattr(self, "_vinculaciones") and nlin in self._vinculaciones:
            self._mostrar_region_vinculada(nlin)
        # Sync visor si es una línea de tira
        if self._visor_imagenes:
            try:
                linea = self.editor.get(f"{nlin}.0", f"{nlin}.end").strip()
                m = re.search(r"\d+", linea)
                if es_encabezado_tira(linea) and m:
                    self._visor_ir(int(m.group()) - 1)
            except Exception:
                pass

    def _sync_visor_on_click(self, _=None):
        """Compat: redirige a _on_editor_click."""
        self._on_editor_click()



    # ═══════════════════════════════════════════════════════════════════════════
    #  ACCIONES
    # ═══════════════════════════════════════════════════════════════════════════
    # ── MEJORA 8b: Insertar símbolo desde panel ───────────────────────────────
    def _insertar_simbolo(self, texto: str):
        """Inserta el prefijo de símbolo en la posición actual del cursor."""
        try:
            pos = self.editor.index("insert")
            # Si la posición actual no es inicio de línea, ir al inicio de la siguiente
            col = int(pos.split(".")[1])
            if col != 0:
                self.editor.mark_set("insert", f"{pos} lineend")
                self.editor.insert("insert", "\n")
            self.editor.insert("insert", texto)
            self.editor.see("insert")
            self.editor.focus()
        except Exception:
            pass

    # ── MEJORA 1: Auto-tilde con preview y confirmación ───────────────────────
    def _auto_tilde(self):
        texto = self.editor.get("1.0", "end-1c")
        cambios = calcular_cambios_tilde(texto)

        if not cambios:
            messagebox.showinfo("Auto-tilde", "No se encontraron tildes faltantes.")
            return

        # Ventana de confirmación con lista de cambios
        win = tk.Toplevel(self.root)
        win.title("´ Auto-tilde — Confirmar cambios")
        win.configure(bg=C["bg"])
        win.geometry("560x480")
        win.grab_set()   # modal

        tk.Label(win, text="´  Tildes faltantes encontradas",
                 bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 11, "bold")).pack(fill="x", pady=10, padx=0)
        sep(win)
        tk.Label(win,
                 text="  Marcá los cambios que querés aplicar (todos marcados por defecto):",
                 bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 8)).pack(anchor="w", padx=12, pady=(8,4))

        # Lista scrolleable con checkboxes
        outer = tk.Frame(win, bg=C["bg"]); outer.pack(fill="both", expand=True, padx=12)
        canvas = tk.Canvas(outer, bg=C["bg2"], highlightthickness=0)
        sb = ttk.Scrollbar(outer, orient="vertical", style="Vertical.TScrollbar",
                           command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y"); canvas.pack(side="left", fill="both", expand=True)
        lista_frame = tk.Frame(canvas, bg=C["bg2"])
        canvas.create_window((0, 0), window=lista_frame, anchor="nw")

        vars_chk = {}
        for num_linea, sin, con, pos_en_linea in cambios:
            var = tk.BooleanVar(value=True)
            vars_chk[(num_linea, sin)] = var
            row = tk.Frame(lista_frame, bg=C["bg2"])
            row.pack(fill="x", padx=8, pady=2)
            tk.Checkbutton(row, variable=var, bg=C["bg2"],
                           activebackground=C["bg3"],
                           selectcolor=C["teal3"],
                           fg=C["text"], font=("Segoe UI", 8)).pack(side="left")
            tk.Label(row, text=f"L{num_linea:4d}", bg=C["bg2"], fg=C["dim"],
                     font=("Consolas", 8)).pack(side="left", padx=(2,8))
            tk.Label(row, text=sin, bg=C["bg2"], fg="#ff8080",
                     font=("Consolas", 9, "bold")).pack(side="left")
            tk.Label(row, text="  →  ", bg=C["bg2"], fg=C["dim"],
                     font=("Segoe UI", 8)).pack(side="left")
            tk.Label(row, text=con, bg=C["bg2"], fg="#80ff80",
                     font=("Consolas", 9, "bold")).pack(side="left")

        lista_frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))
        canvas.bind("<MouseWheel>",
                    lambda e: canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

        # Botones seleccionar todo / ninguno
        btn_row = tk.Frame(win, bg=C["bg"]); btn_row.pack(fill="x", padx=12, pady=4)
        def seleccionar_todos():
            for v in vars_chk.values(): v.set(True)
        def deseleccionar_todos():
            for v in vars_chk.values(): v.set(False)
        make_btn(btn_row, "Marcar todos", seleccionar_todos,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "Desmarcar todos", deseleccionar_todos,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(side="left")

        sep(win)
        action_row = tk.Frame(win, bg=C["bg"]); action_row.pack(fill="x", padx=12, pady=8)

        def aplicar():
            seleccionados = {k for k, v in vars_chk.items() if v.get()}
            if not seleccionados:
                win.destroy(); return
            if not self._texto_original:
                self._texto_original = texto
            nuevo, n = aplicar_cambios_tilde(texto, seleccionados)
            self.editor.delete("1.0", "end")
            self.editor.insert("1.0", nuevo)
            self._actualizar_vista(nuevo)
            self._lbl_estado.config(text=f"´  {n} tilde(s) aplicada(s)", fg=C["ok"])
            win.destroy()

        make_btn(action_row, f"Aplicar seleccionados ({len(cambios)})", aplicar,
                 bg=C["teal"], fg=C["bg"], hover_bg=C["cyan"],
                 small=True).pack(side="left", padx=4)
        make_btn(action_row, "Cancelar", win.destroy,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["err"],
                 small=True).pack(side="left")

    def _abrir(self, path=None, nombre=None):
        if not path:
            path = filedialog.askopenfilename(
                title="Abrir archivo de traducción",
                filetypes=[("Word","*.docx"),("Texto","*.txt"),
                           ("Markdown","*.md"),("Todos","*.*")])
        if not path: return
        self._archivo_path = path
        try:
            texto = self._leer(path)
            self._texto_original = texto
            self._nombres_propios = detectar_nombres_propios(texto)
            self.editor.delete("1.0", "end")
            self.editor.insert("1.0", texto)
            self._actualizar_vista(texto)
            display = nombre or os.path.basename(path)
            self._lbl_archivo.config(text=f"📄  {display}", fg=C["teal"])
        except Exception as ex:
            messagebox.showerror("Error al abrir", str(ex))

    def _leer(self, path):
        if path.lower().endswith(".docx"):
            try:
                import subprocess
                r = subprocess.run(["pandoc", path, "-t", "markdown"],
                                   capture_output=True, text=True, timeout=30)
                if r.returncode == 0: return r.stdout
            except FileNotFoundError:
                pass
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(path) as z:
                xml = z.read("word/document.xml")
            root = ET.fromstring(xml)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            return "\n".join(
                "".join(n.text or "" for n in p.findall(".//w:t", ns))
                for p in root.findall(".//w:p", ns))
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _analizar(self):
        self._actualizar_vista(self.editor.get("1.0", "end-1c"))

    def _auto_corregir(self):
        texto = self.editor.get("1.0", "end-1c")
        if not self._texto_original: self._texto_original = texto
        nuevo, cambios = corregir_texto(texto)
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", nuevo)
        self._actualizar_vista(nuevo)
        self._lbl_estado.config(
            text=f"✓  {cambios} corrección(es)" if cambios else "Sin cambios automáticos",
            fg=C["ok"] if cambios else C["dim"])

    def _guardar(self):
        texto = self.editor.get("1.0", "end-1c")
        save = self._archivo_path
        if not save or save.lower().endswith(".docx"):
            save = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[("Texto","*.txt"),("Markdown","*.md")],
                initialfile="traduccion_corregida.txt")
        if not save: return
        with open(save, "w", encoding="utf-8") as f: f.write(texto)
        self._lbl_estado.config(text=f"💾  {os.path.basename(save)}", fg=C["ok"])

    # MEJORA 3: Exportar reporte ──────────────────────────────────────────────
    def _exportar_reporte(self):
        if not self._errores_por_linea and not self._errores_tira:
            messagebox.showinfo("Reporte", "No hay errores. Primero hacé click en Analizar.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Texto","*.txt")],
            initialfile="reporte_errores.txt")
        if not path: return
        texto = self.editor.get("1.0", "end-1c")
        exportar_reporte(self._errores_por_linea, self._errores_tira, texto, path)
        self._lbl_estado.config(text=f"📋  Reporte guardado: {os.path.basename(path)}", fg=C["ok"])
        if messagebox.askyesno("Reporte", "¿Abrir el reporte ahora?"):
            os.startfile(path) if os.name == "nt" else os.system(f"xdg-open '{path}'")

    # Estadísticas ─────────────────────────────────────────────────────────────
    def _mostrar_stats(self):
        texto = self.editor.get("1.0", "end-1c")
        stats = estadisticas(texto)
        win = tk.Toplevel(self.root); win.title("📊 Estadísticas")
        win.configure(bg=C["bg"]); win.geometry("420x560")
        tk.Label(win, text="📊  Estadísticas", bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 11, "bold")).pack(fill="x", pady=10)
        sep(win)
        outer = tk.Frame(win, bg=C["bg"]); outer.pack(fill="both", expand=True, padx=16, pady=10)
        res = tk.Frame(outer, bg=C["bg3"],
                       highlightbackground=C["border2"], highlightthickness=1)
        res.pack(fill="x", pady=(0,10))
        inner = tk.Frame(res, bg=C["bg3"]); inner.pack(fill="x", padx=12, pady=8)
        for label, val in [("Tiras", stats["TIRAS"]),
                           ("Líneas", stats["TOTAL_LINEAS"]),
                           ("Sin símbolo", stats["SIN_SIMBOLO"]),
                           ("Nombres detectados", len(self._nombres_propios))]:
            row = tk.Frame(inner, bg=C["bg3"]); row.pack(fill="x", pady=1)
            tk.Label(row, text=label, bg=C["bg3"], fg=C["text2"],
                     font=("Segoe UI", 9)).pack(side="left")
            tk.Label(row, text=str(val), bg=C["bg3"], fg=C["teal"],
                     font=("Segoe UI", 9, "bold")).pack(side="right")
        tk.Label(outer, text="Líneas por símbolo", bg=C["bg"], fg=C["dim"],
                 font=("Segoe UI", 8, "bold")).pack(anchor="w", pady=(0,4))
        max_val = max((stats.get(s, 0) for s in PREFIJOS), default=1)
        for sym in sorted(PREFIJOS):
            count = stats.get(sym, 0)
            if not count: continue
            row = tk.Frame(outer, bg=C["bg"]); row.pack(fill="x", pady=1)
            tk.Label(row, text=sym, bg=C["bg"], fg=C["cyan"],
                     font=("Consolas", 9, "bold"), width=8,
                     anchor="w").pack(side="left")
            bw = max(4, int(160 * count / max(max_val, 1)))
            tk.Frame(row, bg=C["teal3"], width=bw, height=14).pack(side="left", padx=4)
            tk.Label(row, text=str(count), bg=C["bg"], fg=C["text2"],
                     font=("Segoe UI", 8)).pack(side="left")
        if self._nombres_propios:
            tk.Label(outer, text=f"Nombres propios detectados: {', '.join(sorted(self._nombres_propios)[:12])}",
                     bg=C["bg"], fg=C["dim"],
                     font=("Segoe UI", 7), wraplength=380, justify="left").pack(anchor="w", pady=(6,0))
        make_btn(win, "Cerrar", win.destroy, bg=C["bg4"], fg=C["text"],
                 hover_bg=C["bg5"], hover_fg=C["bright"], small=True).pack(pady=10)

    # Comparar ─────────────────────────────────────────────────────────────────
    def _comparar(self):
        actual = self.editor.get("1.0", "end-1c")
        orig = self._texto_original
        if not orig or orig == actual:
            messagebox.showinfo("Comparar",
                "No hay versión original.\nAbrí un archivo primero, luego aplicá correcciones.")
            return
        win = tk.Toplevel(self.root)
        win.title("🔀 Original vs. Corregido")
        win.configure(bg=C["bg"]); win.geometry("1100x650")
        tk.Label(win, text="🔀  Líneas cambiadas en amarillo",
                 bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 11, "bold")).pack(fill="x", pady=8, padx=16, anchor="w")
        sep(win)
        paned = tk.PanedWindow(win, orient="horizontal", bg=C["border"], sashwidth=4)
        paned.pack(fill="both", expand=True)
        lineas_orig = orig.splitlines(); lineas_corr = actual.splitlines()
        diff = {i+1 for i,(a,b) in enumerate(zip(lineas_orig, lineas_corr)) if a != b}
        def panel(titulo, contenido, color):
            f = tk.Frame(paned, bg=C["bg"])
            tk.Label(f, text=f"  {titulo}", bg=C["bg3"], fg=color,
                     font=("Segoe UI", 9, "bold"), anchor="w").pack(fill="x")
            inner = tk.Frame(f, bg=C["bg"]); inner.pack(fill="both", expand=True)
            sbv = ttk.Scrollbar(inner, orient="vertical",   style="Vertical.TScrollbar")
            sbh = ttk.Scrollbar(inner, orient="horizontal", style="Horizontal.TScrollbar")
            t = tk.Text(inner, bg=C["bg2"], fg=C["text"], font=("Consolas", 9),
                        wrap="none", relief="flat", bd=0, padx=8, pady=6,
                        yscrollcommand=sbv.set, xscrollcommand=sbh.set, state="disabled")
            sbv.config(command=t.yview); sbh.config(command=t.xview)
            sbv.pack(side="right", fill="y"); sbh.pack(side="bottom", fill="x")
            t.pack(side="left", fill="both", expand=True)
            t.config(state="normal"); t.insert("1.0", contenido)
            t.tag_config("diff", background="#2d2200", foreground="#ffdd88")
            for n in diff: t.tag_add("diff", f"{n}.0", f"{n}.end")
            t.config(state="disabled"); return f
        paned.add(panel("📄 Original", orig,   C["text2"]), minsize=300, stretch="always")
        paned.add(panel("✅ Corregido", actual, C["ok"]),    minsize=300, stretch="always")

    # MEJORA 8: Visor ─────────────────────────────────────────────────────────
    def _abrir_visor(self):
        """El visor ahora está integrado al layout — foco en el canvas."""
        if not PIL_OK:
            self._pedir_pillow(); return
        if self._visor_imagenes:
            self._visor_canvas.focus_set()
        else:
            self._visor_abrir_carpeta()

    def _visor_navego(self, idx: int):
        """Bidireccional: cuando visor cambia de imagen, editor hace scroll a ## Tira N."""
        tira_num = idx + 1
        texto = self.editor.get("1.0", "end-1c")
        for i, linea in enumerate(texto.splitlines()):
            if self._es_encabezado_tira(linea):
                if self._numero_tira(linea) == tira_num:
                    self.editor.see(f"{i+1}.0")
                    self.editor.mark_set("insert", f"{i+1}.0")
                    self._actualizar_posicion()
                    return

    # MEJORA 10: Drive ────────────────────────────────────────────────────────
    def _abrir_drive(self):
        if self._panel_drive and self._panel_drive.winfo_exists():
            self._panel_drive.lift(); return
        self._panel_drive = PanelDrive(
            self.root,
            on_abrir_texto=self._abrir,
            on_imagenes_drive=self._on_imagenes_drive,
            visor_ref=self._visor
        )

    def _on_imagenes_drive(self, imgs):
        """Recibe imágenes descargadas de Drive y las carga en el visor integrado."""
        self.cargar_imagenes_en_visor(imgs)

    # ── Buscador ──────────────────────────────────────────────────────────────
    def _buscar(self, *_):
        query = self._buscar_var.get()
        self.editor.tag_remove("busqueda",       "1.0", "end")
        self.editor.tag_remove("busqueda_actual", "1.0", "end")
        self._busqueda_indices = []
        if not query or len(query) < 2:
            self._lbl_buscar_res.config(text=""); return
        start = "1.0"
        while True:
            pos = self.editor.search(query, start, nocase=True, stopindex="end")
            if not pos: break
            end = f"{pos}+{len(query)}c"
            self.editor.tag_add("busqueda", pos, end)
            self._busqueda_indices.append(pos); start = end
        total = len(self._busqueda_indices)
        if total:
            self._busqueda_actual = 0; self._resaltar_actual()
            self._lbl_buscar_res.config(text=f"{total} resultado(s)", fg=C["search_fg"])
        else:
            self._lbl_buscar_res.config(text="Sin resultados", fg=C["err"])

    def _resaltar_actual(self):
        self.editor.tag_remove("busqueda_actual", "1.0", "end")
        if not self._busqueda_indices: return
        pos = self._busqueda_indices[self._busqueda_actual]
        self.editor.tag_add("busqueda_actual", pos, f"{pos}+{len(self._buscar_var.get())}c")
        self.editor.see(pos)

    def _siguiente_busqueda(self, _=None):
        if not self._busqueda_indices: return
        self._busqueda_actual = (self._busqueda_actual + 1) % len(self._busqueda_indices)
        self._resaltar_actual()

    def _anterior_busqueda(self, _=None):
        if not self._busqueda_indices: return
        self._busqueda_actual = (self._busqueda_actual - 1) % len(self._busqueda_indices)
        self._resaltar_actual()

    def _limpiar_busqueda(self):
        self._buscar_var.set("")
        self.editor.tag_remove("busqueda",       "1.0", "end")
        self.editor.tag_remove("busqueda_actual", "1.0", "end")
        self._busqueda_indices = []; self._lbl_buscar_res.config(text="")

    # ── Vista / resaltado ─────────────────────────────────────────────────────
    def _actualizar_vista(self, texto: str):
        # Todavía no se terminó de construir la UI
        if self.err_list is None or self._lbl_count is None:
            return
        max_len = self._var_maxlen.get() if hasattr(self, "_var_maxlen") else MAX_CHARS_LINEA
        for tag in ("encabezado","prefijo","err_puntuacion","err_mayuscula",
                    "err_simbolo","err_tilde","err_estructura","err_duplicado",
                    "err_largo","err_nombre","err_tira","err_ortografia","err_mezcla",
                    "err_linea_activa","err_manual","err_comentario"):
            self.editor.tag_remove(tag, "1.0", "end")
        self.err_list.delete(0, "end"); self._errores_por_linea.clear()

        lineas = texto.splitlines()
        total_err = 0

        # ── Detectar qué tira corresponde a cada línea ────────────────────────
        tira_de_linea = {}   # {num_linea: num_tira}
        tira_actual = 0
        for i, linea in enumerate(lineas):
            num = i + 1
            if es_encabezado_tira(linea.strip()):
                t = numero_tira_de_linea(linea)
                if t is not None:
                    tira_actual = t
            tira_de_linea[num] = tira_actual

        # Recolectar todos los errores con su tira
        errores_por_tira = {}  # {tira: [(num_linea, tipo, msg), ...]}

        # MEJORA 2: Tiras
        self._errores_tira = analizar_tiras(lineas)
        for num_linea, msg in self._errores_tira:
            self.editor.tag_add("err_tira", f"{num_linea}.0", f"{num_linea}.end")
            t = tira_de_linea.get(num_linea, 0)
            errores_por_tira.setdefault(t, []).append((num_linea, "tira", msg))
            total_err += 1

        # MEJORA 5: Mezcla de sistemas
        for num_linea, tipo, msg in detectar_mezcla_sistemas(texto):
            self.editor.tag_add("err_mezcla", f"{num_linea}.0", f"{num_linea}.end")
            t = tira_de_linea.get(num_linea, 0)
            errores_por_tira.setdefault(t, []).append((num_linea, "mezcla", msg))
            total_err += 1

        # MEJORA 9: Nombres
        errores_nombres = analizar_nombres(texto, self._nombres_propios)
        dups = analizar_duplicados(lineas)

        ICONOS  = {"puntuacion":"⚠","mayuscula":"Aa","simbolo":"?",
                   "tilde":"´","estructura":"⚙","duplicado":"⊡",
                   "largo":"↔","nombre":"👤","tira":"#",
                   "ortografia":"abc","mezcla":"⚡",
                   "manual":"⚑","comentario":"📸"}
        COLORES = {"puntuacion":C["err_punt_fg"],"mayuscula":C["err_may_fg"],
                   "simbolo":C["err_sim_fg"],"tilde":C["err_til_fg"],
                   "estructura":C["err_dup_fg"],"duplicado":C["err_dup_fg"],
                   "largo":C["err_long_fg"],"nombre":C["err_may_fg"],
                   "tira":C["err_sim_fg"],"ortografia":C["err_ort_fg"],
                   "mezcla":C["err_mix_fg"],"manual":"#ff60a0","comentario":"#60c0ff"}

        for i, linea in enumerate(lineas):
            num = i + 1
            stripped = linea.strip()
            if es_encabezado_tira(stripped):
                self.editor.tag_add("encabezado", f"{num}.0", f"{num}.end"); continue
            prefijo, _ = extraer_prefijo(stripped)
            if prefijo:
                col = len(linea) - len(linea.lstrip())
                self.editor.tag_add("prefijo", f"{num}.{col}",
                                    f"{num}.{col+len(prefijo)}")
            errores = analizar_linea(linea, self._nombres_propios)
            if num in dups:
                errores.append(("duplicado", dups[num], 0))

            # ── Filtrar errores ignorados por el usuario ───────────────────
            ignorados = getattr(self, "_errores_ignorados", {}).get(num, set())
            if ignorados:
                if "*" in ignorados:
                    errores = []
                else:
                    errores = [(t, m, *r) for t, m, *r in errores if t not in ignorados]

            if errores:
                self._errores_por_linea[num] = errores
                total_err += len(errores)
            t = tira_de_linea.get(num, 0)
            for tipo, msg, *_ in errores:
                self.editor.tag_add(f"err_{tipo}", f"{num}.0", f"{num}.end")
                errores_por_tira.setdefault(t, []).append((num, tipo, msg))

        # Nombres
        for num_linea, tipo, msg in errores_nombres:
            ignorados = getattr(self, "_errores_ignorados", {}).get(num_linea, set())
            if "*" in ignorados or tipo in ignorados:
                continue
            self._errores_por_linea.setdefault(num_linea, []).append((tipo, msg, 0))
            self.editor.tag_add(f"err_{tipo}", f"{num_linea}.0", f"{num_linea}.end")
            t = tira_de_linea.get(num_linea, 0)
            errores_por_tira.setdefault(t, []).append((num_linea, tipo, msg))
            total_err += 1

        # Errores manuales
        for num_linea, desc in getattr(self, "_errores_manuales", {}).items():
            self.editor.tag_add("err_manual", f"{num_linea}.0", f"{num_linea}.end")
            t = tira_de_linea.get(num_linea, 0)
            errores_por_tira.setdefault(t, []).append((num_linea, "manual", desc))
            total_err += 1

        # Comentarios de imagen
        for num_linea, dat in getattr(self, "_comentarios_img", {}).items():
            self.editor.tag_add("err_comentario", f"{num_linea}.0", f"{num_linea}.end")
            t = tira_de_linea.get(num_linea, 0)
            errores_por_tira.setdefault(t, []).append((num_linea, "comentario",
                                                        dat.get("texto","(sin descripción)")))

        # ── Poblar la lista de errores agrupada por tira ──────────────────────
        for tira_n in sorted(errores_por_tira.keys()):
            items = errores_por_tira[tira_n]
            # Encabezado de tira
            lbl = f"── Tira {tira_n} ──" if tira_n > 0 else "── General ──"
            self.err_list.insert("end", f"  {lbl}")
            self.err_list.itemconfig("end", fg=C["teal"])
            for num_linea, tipo, msg in sorted(items, key=lambda x: x[0]):
                ico = ICONOS.get(tipo, "•")
                col = COLORES.get(tipo, C["text"])
                self.err_list.insert("end", f"    L{num_linea}  {ico}  {msg}")
                self.err_list.itemconfig("end", fg=col)

        n_lineas = len(self._errores_por_linea)
        n_tiras  = sum(1 for l in lineas if es_encabezado_tira(l))
        if self.err_list:
            if total_err == 0 and not getattr(self, "_comentarios_img", {}):
                self.err_list.insert("end", "  ✓  Sin errores encontrados")
                self.err_list.itemconfig("end", fg=C["ok"])
        if self._lbl_estado:
            if total_err == 0:
                self._lbl_estado.config(text="✓  Sin errores detectados", fg=C["ok"])
            else:
                self._lbl_estado.config(
                    text=f"⚠  {total_err} error(es) en {n_lineas} línea(s)", fg=C["warn"])
        if self._lbl_count:
            self._lbl_count.config(
            text=f"{len(lineas)} líneas  ·  {n_tiras} tiras  ·  {total_err} errores")
        # Actualizar números de línea
        self.editor.after_idle(self._actualizar_numeros_linea)

    def _toggle_legend(self):
        self._legend_visible = not self._legend_visible
        if self._legend_visible:
            self._legend_frame.pack(fill="x", pady=(0,4))
            self._btn_legend.config(text="▼ Referencia de colores")
        else:
            self._legend_frame.pack_forget()
            self._btn_legend.config(text="▶ Referencia de colores")

    def _ir_a_error(self, _=None):
        sel = self.err_list.curselection()
        if not sel: return
        m = re.match(r"\s*L(\d+)\s", self.err_list.get(sel[0]))
        if not m: return
        nlin = m.group(1)
        self.editor.see(f"{nlin}.0")
        self.editor.mark_set("insert", f"{nlin}.0")
        self.editor.focus()
        self._actualizar_posicion()
        # Flash: resalta la línea 3 veces con color brillante
        self._flash_linea(int(nlin), 0)


    # ── Menú contextual lista de errores ─────────────────────────────────────
    def _err_list_context_menu(self, event):
        """Clic derecho en la lista de errores: muestra opciones para ignorar."""
        idx = self.err_list.nearest(event.y)
        if idx < 0:
            return
        self.err_list.selection_clear(0, "end")
        self.err_list.selection_set(idx)
        item_text = self.err_list.get(idx)

        m2 = re.match(r"\s*L\s*(\d+)", item_text)
        if not m2:
            return
        nlin = int(m2.group(1))

        menu = tk.Menu(self.root, tearoff=0,
                       bg=C["bg3"], fg=C["text"],
                       activebackground=C["teal3"],
                       activeforeground=C["bright"],
                       relief="flat", bd=0)

        menu.add_command(
            label=f"\u2715  Ignorar este error (L{nlin})",
            command=lambda n=nlin, t=item_text: self._ignorar_error(n, t, solo_tipo=True))

        menu.add_command(
            label=f"\u2715\u2715  Ignorar TODOS los errores de L{nlin}",
            command=lambda n=nlin: self._ignorar_error(n, "", solo_tipo=False))

        menu.add_separator()

        ignorados = self._errores_ignorados.get(nlin, set())
        if ignorados:
            menu.add_command(
                label=f"\u21ba  Restaurar errores de L{nlin}",
                command=lambda n=nlin: self._restaurar_errores(n))
        else:
            menu.add_command(label="\u21ba  Restaurar errores de L...", state="disabled")

        menu.add_separator()
        menu.add_command(
            label="\u21ba\u21ba  Restaurar TODOS los ignorados",
            command=self._restaurar_todos_errores)

        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _ignorar_error(self, nlin: int, item_text: str, solo_tipo: bool):
        if solo_tipo:
            LABEL_TO_TIPO = {
                "PUNTUACI": "puntuacion", "MAYUSC": "mayuscula",
                "SIMBOL":   "simbolo",    "TILDE":  "tilde",
                "STRUCTURA":"estructura", "DUPLICAD":"duplicado",
                "LARGO":    "largo",      "NOMBRE": "nombre",
                "TIRA":     "tira",       "ORTOGRAF":"ortografia",
                "MEZCLA":   "mezcla",     "MANUAL": "manual",
                "COMENTAR": "comentario",
            }
            tipo_detectado = None
            upper = item_text.upper()
            for key, tipo in LABEL_TO_TIPO.items():
                if key in upper:
                    tipo_detectado = tipo
                    break
            if tipo_detectado:
                self._errores_ignorados.setdefault(nlin, set()).add(tipo_detectado)
            else:
                self._errores_ignorados[nlin] = {"*"}
        else:
            self._errores_ignorados[nlin] = {"*"}

        texto = self.editor.get("1.0", "end-1c")
        self._actualizar_vista(texto)

    def _restaurar_errores(self, nlin: int):
        self._errores_ignorados.pop(nlin, None)
        texto = self.editor.get("1.0", "end-1c")
        self._actualizar_vista(texto)

    def _restaurar_todos_errores(self):
        self._errores_ignorados.clear()
        texto = self.editor.get("1.0", "end-1c")
        self._actualizar_vista(texto)

    def _flash_linea(self, nlin: int, paso: int):
        """Parpadea el fondo de la línea `nlin` para que sea obvio dónde está."""
        tag = "err_linea_activa"
        if paso % 2 == 0:
            self.editor.tag_add(tag, f"{nlin}.0", f"{nlin}.end")
        else:
            self.editor.tag_remove(tag, "1.0", "end")
        if paso < 5:   # 3 parpadeos completos
            self.editor.after(180, lambda: self._flash_linea(nlin, paso + 1))
        else:
            # Queda resaltado de forma permanente hasta el próximo click
            self.editor.tag_add(tag, f"{nlin}.0", f"{nlin}.end")

    # ── Error manual ──────────────────────────────────────────────────────────
    def _marcar_error_manual(self):
        """Abre diálogo para marcar un error manual en la línea seleccionada."""
        try:
            pos = self.editor.index("insert")
            nlin = int(pos.split(".")[0])
        except Exception:
            nlin = 1

        win = tk.Toplevel(self.root)
        win.title("⚑ Marcar error manual")
        win.configure(bg=C["bg"])
        win.geometry("480x240")
        win.grab_set()
        win.resizable(False, False)

        linea_actual = self.editor.get(f"{nlin}.0", f"{nlin}.end").strip()[:60]

        tk.Label(win, text="⚑  Marcar error manual en línea",
                 bg=C["bg2"], fg="#ff60a0", font=("Segoe UI", 10, "bold")).pack(fill="x", pady=8)

        row1 = tk.Frame(win, bg=C["bg"]); row1.pack(fill="x", padx=14, pady=4)
        tk.Label(row1, text="Número de línea:", bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 9)).pack(side="left")
        var_nlin = tk.IntVar(value=nlin)
        tk.Spinbox(row1, from_=1, to=9999, textvariable=var_nlin, width=6,
                   bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                   relief="flat", font=("Segoe UI", 9),
                   buttonbackground=C["bg5"]).pack(side="left", padx=8)
        lbl_preview = tk.Label(row1, text=f"→ {linea_actual}", bg=C["bg"],
                               fg=C["dim"], font=("Consolas", 8))
        lbl_preview.pack(side="left", padx=4)

        def _update_preview(*_):
            try:
                n = var_nlin.get()
                txt = self.editor.get(f"{n}.0", f"{n}.end").strip()[:60]
                lbl_preview.config(text=f"→ {txt}")
            except Exception:
                pass
        var_nlin.trace_add("write", _update_preview)

        tk.Label(win, text="Descripción del error:", bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 9)).pack(anchor="w", padx=14)
        var_desc = tk.StringVar()
        tk.Entry(win, textvariable=var_desc, bg=C["bg4"], fg=C["text"],
                 insertbackground=C["teal"], relief="flat",
                 font=("Segoe UI", 9)).pack(fill="x", padx=14, pady=4, ipady=4)

        def _aplicar():
            n = var_nlin.get()
            desc = var_desc.get().strip() or "Error marcado manualmente"
            self._errores_manuales[n] = desc
            win.destroy()
            texto = self.editor.get("1.0", "end-1c")
            self._actualizar_vista(texto)

        def _quitar():
            n = var_nlin.get()
            self._errores_manuales.pop(n, None)
            win.destroy()
            texto = self.editor.get("1.0", "end-1c")
            self._actualizar_vista(texto)

        btn_row = tk.Frame(win, bg=C["bg"]); btn_row.pack(fill="x", padx=14, pady=8)
        make_btn(btn_row, "⚑ Marcar error", _aplicar,
                 bg="#ff60a0", fg=C["bg"], hover_bg="#ff80b0",
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "✕ Quitar marca", _quitar,
                 bg=C["bg4"], fg=C["err"], hover_bg=C["bg5"],
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "Cancelar", win.destroy,
                 bg=C["bg4"], fg=C["dim"], hover_bg=C["bg5"],
                 small=True).pack(side="left")

    # ── Captura de imagen → comentario ────────────────────────────────────────
    def _captura_a_comentario(self):
        """Activa modo captura de región de la imagen actual para crear un comentario."""
        if not self._visor_imagenes:
            messagebox.showwarning("Sin imagen",
                "Cargá las imágenes del manwha en el visor primero.")
            return
        if not PIL_OK:
            messagebox.showwarning("Pillow",
                "Necesitás Pillow para capturar regiones.\npip install Pillow")
            return

        # Obtener línea actual del editor
        try:
            pos = self.editor.index("insert")
            nlin_defecto = int(pos.split(".")[0])
        except Exception:
            nlin_defecto = 1

        # Activar modo de selección de región en el visor
        self._captura_nlin = nlin_defecto
        self._captura_activa = True
        self._captura_start = None
        self._captura_rect_id = None

        # Cambiar cursor y mostrar hint
        self._visor_canvas.config(cursor="crosshair")
        messagebox.showinfo("Captura de región",
            f"Arrastrá un rectángulo sobre la imagen para capturar esa zona.\n"
            f"Se adjuntará como comentario en la línea {nlin_defecto}.\n\n"
            "Podés cambiar el número de línea en el diálogo que se abre después.")

    def _visor_captura_start(self, event):
        if not getattr(self, "_captura_activa", False):
            return False
        cx = self._visor_canvas.canvasx(event.x)
        cy = self._visor_canvas.canvasy(event.y)
        self._captura_start = (cx, cy)
        if self._captura_rect_id:
            self._visor_canvas.delete(self._captura_rect_id)
            self._captura_rect_id = None
        return True

    def _visor_captura_move(self, event):
        if not getattr(self, "_captura_activa", False) or not self._captura_start:
            return False
        x0, y0 = self._captura_start
        cx = self._visor_canvas.canvasx(event.x)
        cy = self._visor_canvas.canvasy(event.y)
        if self._captura_rect_id:
            self._visor_canvas.delete(self._captura_rect_id)
        self._captura_rect_id = self._visor_canvas.create_rectangle(
            x0, y0, cx, cy,
            outline="#60c0ff", width=2, dash=(4,2), tags="captura_rect")
        return True

    def _visor_captura_end(self, event):
        if not getattr(self, "_captura_activa", False) or not self._captura_start:
            return False
        x0, y0 = self._captura_start
        x1 = self._visor_canvas.canvasx(event.x)
        y1 = self._visor_canvas.canvasy(event.y)
        self._captura_activa = False
        self._visor_canvas.config(cursor="crosshair")
        if self._captura_rect_id:
            self._visor_canvas.delete(self._captura_rect_id)
            self._captura_rect_id = None
        if abs(x1-x0) < 8 or abs(y1-y0) < 8:
            return True

        # Recortar imagen — x0,y0,x1,y1 ya son coords de canvas (con scroll y zoom)
        img_b64 = None
        try:
            _, fuente = self._visor_imagenes[self._visor_idx]
            if isinstance(fuente, bytes):
                img_orig = Image.open(io.BytesIO(fuente))
            else:
                img_orig = Image.open(fuente)
            z = self._visor_zoom
            # Convertir de coords canvas a coords imagen original
            ix0 = int(min(x0,x1) / z); iy0 = int(min(y0,y1) / z)
            ix1 = int(max(x0,x1) / z); iy1 = int(max(y0,y1) / z)
            ix0 = max(0, ix0); iy0 = max(0, iy0)
            ix1 = min(img_orig.width, ix1); iy1 = min(img_orig.height, iy1)
            crop = img_orig.crop((ix0, iy0, ix1, iy1))
            max_w = 480
            if crop.width > max_w:
                r = max_w / crop.width
                crop = crop.resize((max_w, int(crop.height*r)), Image.LANCZOS)
            buf = io.BytesIO()
            crop.save(buf, format="PNG")
            img_b64 = base64.b64encode(buf.getvalue()).decode()
        except Exception:
            img_b64 = None

        # Diálogo para poner descripción y número de línea
        win = tk.Toplevel(self.root)
        win.title("📸 Adjuntar comentario de imagen")
        win.configure(bg=C["bg"])
        win.geometry("500x420")
        win.grab_set()

        tk.Label(win, text="📸  Comentario de imagen",
                 bg=C["bg2"], fg="#60c0ff",
                 font=("Segoe UI", 10, "bold")).pack(fill="x", pady=8)

        # Preview de la captura
        if img_b64 and PIL_OK:
            try:
                raw = base64.b64decode(img_b64)
                preview_img = Image.open(io.BytesIO(raw))
                max_w2 = 460; max_h2 = 120
                r2 = min(max_w2/preview_img.width, max_h2/preview_img.height, 1.0)
                if r2 < 1.0:
                    preview_img = preview_img.resize(
                        (int(preview_img.width*r2), int(preview_img.height*r2)),
                        Image.LANCZOS)
                _tk_preview = ImageTk.PhotoImage(preview_img)
                lbl_img = tk.Label(win, image=_tk_preview, bg=C["bg"])
                lbl_img.image = _tk_preview
                lbl_img.pack(pady=4)
            except Exception:
                pass

        row1 = tk.Frame(win, bg=C["bg"]); row1.pack(fill="x", padx=14, pady=4)
        tk.Label(row1, text="Línea del error:", bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 9)).pack(side="left")
        var_nlin2 = tk.IntVar(value=getattr(self, "_captura_nlin", 1))
        tk.Spinbox(row1, from_=1, to=9999, textvariable=var_nlin2, width=6,
                   bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                   relief="flat", font=("Segoe UI", 9),
                   buttonbackground=C["bg5"]).pack(side="left", padx=8)

        tk.Label(win, text="Descripción del comentario:", bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 9)).pack(anchor="w", padx=14)
        txt_desc = tk.Text(win, bg=C["bg4"], fg=C["text"], insertbackground=C["teal"],
                           relief="flat", font=("Segoe UI", 9), height=3, wrap="word")
        txt_desc.pack(fill="x", padx=14, pady=4)

        def _guardar_comentario():
            n = var_nlin2.get()
            desc = txt_desc.get("1.0", "end-1c").strip() or "(sin descripción)"
            self._comentarios_img[n] = {"texto": desc, "img_b64": img_b64}
            win.destroy()
            texto = self.editor.get("1.0", "end-1c")
            self._actualizar_vista(texto)

        btn_row = tk.Frame(win, bg=C["bg"]); btn_row.pack(fill="x", padx=14, pady=8)
        make_btn(btn_row, "📸 Guardar comentario", _guardar_comentario,
                 bg="#60c0ff", fg=C["bg"], hover_bg="#80d0ff",
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "Cancelar", win.destroy,
                 bg=C["bg4"], fg=C["dim"], hover_bg=C["bg5"],
                 small=True).pack(side="left")

        return True

    # ── Exportar Word con errores coloreados ──────────────────────────────────

    # ── Diálogo de feedback antes de exportar ────────────────────────────────
    def _dialogo_feedback(self) -> dict | None:
        """
        Muestra una ventana modal con 3 campos opcionales de feedback:
          - Mensaje general (texto libre, ej: "Muy bien, tené en cuenta...")
          - Aprobado (checkbox + texto opcional)
          - Nota adicional
        Devuelve dict con los valores, o None si el usuario canceló.
        """
        resultado = {"cancelado": True}

        win = tk.Toplevel(self.root)
        win.title("📝 Feedback del corrector")
        win.configure(bg=C["bg"])
        win.geometry("520x520")
        win.grab_set()
        win.resizable(False, False)

        # ── Encabezado ────────────────────────────────────────────────────────
        tk.Label(win, text="📝  Feedback del corrector",
                 bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 11, "bold")).pack(fill="x", pady=8, padx=0)
        tk.Label(win, text="  Todos los campos son opcionales",
                 bg=C["bg2"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(fill="x", padx=0)
        tk.Frame(win, bg=C["border"], height=1).pack(fill="x")

        inner = tk.Frame(win, bg=C["bg"]); inner.pack(fill="both", expand=True, padx=16, pady=10)

        # ── Campo: Mensaje general ────────────────────────────────────────────
        tk.Label(inner, text="💬  Mensaje general al traductor:",
                 bg=C["bg"], fg=C["text"],
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0,3))
        tk.Label(inner,
                 text="  Ej: «Muy bien en general, solo tené en cuenta estos detalles...»",
                 bg=C["bg"], fg=C["dim"],
                 font=("Segoe UI", 7)).pack(anchor="w")
        txt_mensaje = tk.Text(inner, bg=C["bg4"], fg=C["text"],
                              insertbackground=C["teal"], relief="flat",
                              font=("Segoe UI", 9), height=5, wrap="word",
                              padx=6, pady=4)
        txt_mensaje.pack(fill="x", pady=(3, 10))

        # ── Campo: Aprobado ───────────────────────────────────────────────────
        tk.Frame(inner, bg=C["border"], height=1).pack(fill="x", pady=(0,8))
        row_ap = tk.Frame(inner, bg=C["bg"]); row_ap.pack(fill="x")
        var_aprobado = tk.BooleanVar(value=False)
        chk = tk.Checkbutton(row_ap, text="✅  Marcar como APROBADO",
                              variable=var_aprobado,
                              bg=C["bg"], fg=C["ok"],
                              selectcolor=C["bg4"],
                              activebackground=C["bg"],
                              activeforeground=C["ok"],
                              font=("Segoe UI", 9, "bold"),
                              cursor="hand2")
        chk.pack(side="left")

        tk.Label(inner, text="   Nota de aprobación (opcional):",
                 bg=C["bg"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(anchor="w", pady=(6,2))
        var_nota_ap = tk.StringVar()
        tk.Entry(inner, textvariable=var_nota_ap,
                 bg=C["bg4"], fg=C["text"],
                 insertbackground=C["teal"],
                 relief="flat", font=("Segoe UI", 9)).pack(fill="x", ipady=3)

        # ── Campo: Nota adicional ─────────────────────────────────────────────
        tk.Frame(inner, bg=C["border"], height=1).pack(fill="x", pady=(10,8))
        tk.Label(inner, text="📌  Nota adicional (opcional):",
                 bg=C["bg"], fg=C["text"],
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0,3))
        txt_nota = tk.Text(inner, bg=C["bg4"], fg=C["text"],
                           insertbackground=C["teal"], relief="flat",
                           font=("Segoe UI", 9), height=3, wrap="word",
                           padx=6, pady=4)
        txt_nota.pack(fill="x")

        # ── Botones ───────────────────────────────────────────────────────────
        tk.Frame(inner, bg=C["border"], height=1).pack(fill="x", pady=(10,0))
        btn_row = tk.Frame(win, bg=C["bg"]); btn_row.pack(fill="x", padx=16, pady=8)

        def _continuar():
            resultado["cancelado"]   = False
            resultado["mensaje"]     = txt_mensaje.get("1.0", "end-1c").strip()
            resultado["aprobado"]    = var_aprobado.get()
            resultado["nota_ap"]     = var_nota_ap.get().strip()
            resultado["nota_extra"]  = txt_nota.get("1.0", "end-1c").strip()
            win.destroy()

        def _saltar():
            resultado["cancelado"]   = False
            resultado["mensaje"]     = ""
            resultado["aprobado"]    = False
            resultado["nota_ap"]     = ""
            resultado["nota_extra"]  = ""
            win.destroy()

        make_btn(btn_row, "📤  Exportar con feedback", _continuar,
                 bg=C["teal"], fg=C["bg"], hover_bg=C["cyan"],
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "⏭  Exportar sin feedback", _saltar,
                 bg=C["bg4"], fg=C["dim"], hover_bg=C["bg5"], hover_fg=C["text"],
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "✕  Cancelar", win.destroy,
                 bg=C["bg4"], fg=C["err"], hover_bg=C["bg5"],
                 small=True).pack(side="right", padx=4)

        win.wait_window()

        if resultado.get("cancelado", True):
            return None
        return resultado

    def _exportar_word_errores(self):
        """Genera un .docx con el script completo, errores resaltados por color y comentarios."""
        texto = self.editor.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Sin texto", "No hay texto para exportar.")
            return
        # Analizar primero si no hay errores cargados
        if not self._errores_por_linea and not self._errores_tira:
            self._actualizar_vista(texto)

        feedback = self._dialogo_feedback()
        if feedback is None:
            return  # usuario canceló

        path = filedialog.asksaveasfilename(
            title="Guardar reporte Word",
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
            initialfile="reporte_errores.docx")
        if not path:
            return

        self._lbl_estado.config(text="⏳ Generando Word…", fg=C["warn"])
        self.root.update_idletasks()

        try:
            self._generar_docx_errores(texto, path, feedback=feedback)
            self._lbl_estado.config(text=f"✅ Word exportado: {os.path.basename(path)}", fg=C["ok"])
            if messagebox.askyesno("Word exportado", f"¿Abrir el archivo ahora?\n{path}"):
                os.startfile(path) if os.name == "nt" else os.system(f"xdg-open '{path}'")
        except Exception as ex:
            self._lbl_estado.config(text="❌ Error al exportar", fg=C["err"])
            messagebox.showerror("Error", f"No se pudo generar el Word:\n{ex}")

    def _generar_docx_errores(self, texto: str, path: str, feedback: dict = None):
        """Genera el .docx con errores como comentarios reales de Word y capturas embebidas."""
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
        except ImportError:
            import subprocess, sys
            subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"],
                           capture_output=True)
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

        # ── Helpers de color ──────────────────────────────────────────────────
        # Colores de fondo para resaltado (como tuplas hex string para shading XML)
        FILL_MAP = {
            "puntuacion": "FF8080",
            "mayuscula":  "8080FF",
            "simbolo":    "FFCC60",
            "tilde":      "80FF80",
            "estructura": "CC80FF",
            "duplicado":  "CC80FF",
            "largo":      "80C0FF",
            "nombre":     "8080FF",
            "tira":       "FFCC60",
            "ortografia": "FF80CC",
            "mezcla":     "FFAA40",
            "manual":     "FF60A0",
            "comentario": "60C0FF",
        }
        # Colores de texto (RGBColor) para el número de línea y encabezados
        def rgb(r, g, b): return RGBColor(r, g, b)

        def _add_shading(paragraph, hex_fill: str):
            """Aplica color de fondo al párrafo completo."""
            pPr = paragraph._p.get_or_add_pPr()
            # Quitar shading previo si existe
            for old in pPr.findall(qn("w:shd")):
                pPr.remove(old)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  hex_fill)
            pPr.append(shd)

        # ── Comentarios reales de Word ────────────────────────────────────────
        # Necesitamos manipular el XML de comentarios directamente
        import zipfile, shutil, uuid
        from lxml import etree as ET
        from datetime import datetime

        WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        WP_NS   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        A_NS    = "http://schemas.openxmlformats.org/drawingml/2006/main"
        PIC_NS  = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        REL_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        DATE    = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        AUTHOR  = "BloomScans Corrector"

        def _qn(tag):
            ns, local = tag.split(":")
            NS = {"w": WORD_NS, "wp": WP_NS, "a": A_NS, "pic": PIC_NS, "r": REL_NS}
            return f"{{{NS[ns]}}}{local}"

        # ── Construir documento base con python-docx ──────────────────────────
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name  = "Consolas"
        style.font.size  = Pt(10)

        # Título
        h = doc.add_heading("Reporte de Errores — BloomScans", 0)
        if h.runs:
            h.runs[0].font.color.rgb = rgb(0x00, 0xCF, 0xB5)

        # Leyenda
        doc.add_heading("Leyenda de colores", 1)
        ICONOS_TIPO = {
            "puntuacion": "⚠ Puntuación",  "mayuscula":  "Aa Mayúscula",
            "simbolo":    "? Símbolo",       "tilde":      "´ Tilde faltante",
            "estructura": "⚙ Estructura",    "duplicado":  "⊡ Duplicado",
            "largo":      "↔ Línea larga",   "nombre":     "👤 Nombre",
            "tira":       "# Tira",          "ortografia": "abc Ortografía",
            "mezcla":     "⚡ Mezcla",        "manual":     "⚑ Error manual",
            "comentario": "📸 Captura",
        }
        for tipo, fill in FILL_MAP.items():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            _add_shading(p, fill)
            r = p.add_run(f"  {ICONOS_TIPO.get(tipo, tipo.upper())}  ")
            r.bold = True
            r.font.size = Pt(9)

        # ── Sección de feedback ──────────────────────────────────────────────
        if feedback:
            _fb_msg    = feedback.get("mensaje", "")
            _fb_ap     = feedback.get("aprobado", False)
            _fb_nota_ap= feedback.get("nota_ap", "")
            _fb_nota   = feedback.get("nota_extra", "")

            if _fb_ap or _fb_msg or _fb_nota:
                doc.add_heading("Feedback del corrector", 1)

            if _fb_ap:
                p_ap = doc.add_paragraph()
                r_ap = p_ap.add_run(
                    f"✅  APROBADO{(': ' + _fb_nota_ap) if _fb_nota_ap else ''}")
                r_ap.bold = True
                r_ap.font.size = Pt(11)
                r_ap.font.color.rgb = rgb(0x22, 0xC5, 0x5E)
                _add_shading(p_ap, "0d2a1a")

            if _fb_msg:
                doc.add_heading("Mensaje al traductor", 2)
                p_msg = doc.add_paragraph(_fb_msg)
                p_msg.paragraph_format.space_after = Pt(6)

            if _fb_nota:
                doc.add_heading("Nota adicional", 2)
                p_nota = doc.add_paragraph(_fb_nota)
                p_nota.paragraph_format.space_after = Pt(6)

        doc.add_heading("Script completo", 1)

        # ── Recolectar datos de errores ───────────────────────────────────────
        lineas = texto.splitlines()
        tipo_por_linea   = {}
        msgs_por_linea   = {}   # {num: [msg, ...]}

        for num, errs in self._errores_por_linea.items():
            tipo_por_linea[num] = errs[0][0]
            msgs_por_linea[num] = [m for _, m, *_ in errs]
        for num, desc in getattr(self, "_errores_manuales", {}).items():
            tipo_por_linea[num] = "manual"
            msgs_por_linea.setdefault(num, []).append(f"⚑ {desc}")
        for num in getattr(self, "_comentarios_img", {}):
            if num not in tipo_por_linea:
                tipo_por_linea[num] = "comentario"
            msgs_por_linea.setdefault(num, []).append("📸 Ver captura más abajo")

        comentarios_img = getattr(self, "_comentarios_img", {})

        # ── Agregar párrafos con número de línea y resaltado ──────────────────
        comment_id = 0
        comment_data = []   # [(id, linea_num, texto_comentario)]

        for i, linea in enumerate(lineas):
            num = i + 1
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

            # Número de línea gris
            rn = p.add_run(f"L{num:4d}  ")
            rn.font.size = Pt(8)
            rn.font.color.rgb = rgb(0x50, 0x70, 0x90)

            # Texto
            rt = p.add_run(linea or " ")
            rt.font.size = Pt(10)

            if es_encabezado_tira(linea.strip()):
                rt.bold = True
                rt.font.color.rgb = rgb(0x00, 0xCF, 0xB5)
            elif num in tipo_por_linea:
                tipo  = tipo_por_linea[num]
                fill  = FILL_MAP.get(tipo, "FFFF00")
                _add_shading(p, fill)
                # Preparar comentario para este párrafo
                msgs = msgs_por_linea.get(num, [])
                if msgs:
                    comment_data.append((comment_id, num, "\n".join(msgs)))
                    comment_id += 1

        # ── Sección de capturas de imagen ─────────────────────────────────────
        if comentarios_img:
            doc.add_heading("Capturas de imagen", 1)
            for num_linea in sorted(comentarios_img.keys()):
                dat = comentarios_img[num_linea]
                p_c = doc.add_paragraph()
                rh = p_c.add_run(f"L{num_linea}  ")
                rh.bold = True
                rh.font.color.rgb = rgb(0x60, 0xC0, 0xFF)
                p_c.add_run(dat.get("texto", ""))

                if dat.get("img_b64") and PIL_OK:
                    try:
                        raw = base64.b64decode(dat["img_b64"])
                        doc.add_picture(io.BytesIO(raw), width=Inches(4.5))
                    except Exception:
                        p_err = doc.add_paragraph()
                        p_err.add_run("[imagen no disponible]").italic = True

        # ── Guardar docx temporal ─────────────────────────────────────────────
        import tempfile
        tmp_path = path + ".tmp.docx"
        doc.save(tmp_path)

        # ── Inyectar comentarios reales en el XML ─────────────────────────────
        if not comment_data:
            shutil.move(tmp_path, path)
            return

        # Abrir el ZIP del docx
        with zipfile.ZipFile(tmp_path, "r") as zin:
            names = zin.namelist()
            file_contents = {n: zin.read(n) for n in names}

        # ── Construir comments.xml ────────────────────────────────────────────
        W = WORD_NS
        comments_root = ET.Element(f"{{{W}}}comments",
            nsmap={"w": W, "r": REL_NS})
        comments_root.set(
            "{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable",
            "w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14")

        for cid, num_linea, msg_text in comment_data:
            cmt = ET.SubElement(comments_root, f"{{{W}}}comment")
            cmt.set(f"{{{W}}}id",     str(cid))
            cmt.set(f"{{{W}}}author", AUTHOR)
            cmt.set(f"{{{W}}}date",   DATE)
            cmt.set(f"{{{W}}}initials", "BC")
            cp = ET.SubElement(cmt, f"{{{W}}}p")
            cr = ET.SubElement(cp, f"{{{W}}}r")
            ct = ET.SubElement(cr, f"{{{W}}}t")
            ct.text = msg_text

        comments_xml = ET.tostring(comments_root, xml_declaration=True,
                                   encoding="UTF-8", standalone=True)
        file_contents["word/comments.xml"] = comments_xml

        # ── Agregar relación en document.xml.rels ─────────────────────────────
        rels_key = "word/_rels/document.xml.rels"
        rels_tree = ET.fromstring(file_contents[rels_key])
        REL_TYPE_CMT = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                        "relationships/comments")
        # Solo agregar si no existe ya
        existing_rels = {r.get("Type") for r in rels_tree}
        if REL_TYPE_CMT not in existing_rels:
            new_rel = ET.SubElement(rels_tree,
                "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship")
            new_rel.set("Id",     "rIdComments")
            new_rel.set("Type",   REL_TYPE_CMT)
            new_rel.set("Target", "comments.xml")
            file_contents[rels_key] = ET.tostring(rels_tree,
                xml_declaration=True, encoding="UTF-8", standalone=True)

        # ── Registrar en [Content_Types].xml ──────────────────────────────────
        ct_key  = "[Content_Types].xml"
        ct_tree = ET.fromstring(file_contents[ct_key])
        CT_CMT  = ("application/vnd.openxmlformats-officedocument"
                   ".wordprocessingml.comments+xml")
        parts = {o.get("PartName") for o in ct_tree
                 if o.tag.endswith("}Override")}
        if "/word/comments.xml" not in parts:
            ov = ET.SubElement(ct_tree,
                "{http://schemas.openxmlformats.org/package/2006/content-types}Override")
            ov.set("PartName",    "/word/comments.xml")
            ov.set("ContentType", CT_CMT)
            file_contents[ct_key] = ET.tostring(ct_tree,
                xml_declaration=True, encoding="UTF-8", standalone=True)

        # ── Inyectar marcas de comentario en document.xml ─────────────────────
        doc_tree = ET.fromstring(file_contents["word/document.xml"])
        # Obtener todos los párrafos del body
        body = doc_tree.find(f"{{{W}}}body")
        all_paras = body.findall(f".//{{{W}}}p") if body is not None else []

        # Mapear comment_data por orden de aparición en el doc
        # Los párrafos de líneas están después de los de leyenda/título
        # Buscamos los párrafos que contienen el texto "L   N  " (número de línea)
        cmt_idx = 0
        for para in all_paras:
            if cmt_idx >= len(comment_data):
                break
            cid, num_linea, _ = comment_data[cmt_idx]
            # Buscar si este párrafo tiene el run con el número de línea
            runs_text = "".join(
                (t.text or "") for t in para.findall(f".//{{{W}}}t"))
            expected = f"L{num_linea:4d}"
            if expected in runs_text:
                # Insertar commentRangeStart antes del primer run
                runs = para.findall(f"{{{W}}}r")
                if runs:
                    cs = ET.Element(f"{{{W}}}commentRangeStart")
                    cs.set(f"{{{W}}}id", str(cid))
                    runs[0].addprevious(cs)
                    ce = ET.Element(f"{{{W}}}commentRangeEnd")
                    ce.set(f"{{{W}}}id", str(cid))
                    runs[-1].addnext(ce)
                    # commentReference run
                    ref_run = ET.Element(f"{{{W}}}r")
                    rpr = ET.SubElement(ref_run, f"{{{W}}}rPr")
                    rst = ET.SubElement(rpr, f"{{{W}}}rStyle")
                    rst.set(f"{{{W}}}val", "CommentReference")
                    cref = ET.SubElement(ref_run, f"{{{W}}}commentReference")
                    cref.set(f"{{{W}}}id", str(cid))
                    ce.addnext(ref_run)
                cmt_idx += 1

        file_contents["word/document.xml"] = ET.tostring(
            doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

        # ── Escribir ZIP final ────────────────────────────────────────────────
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in file_contents.items():
                zout.writestr(name, data)

        try:
            import os as _os
            _os.remove(tmp_path)
        except Exception:
            pass

    # ── Exportar Google Docs con imágenes inline ──────────────────────────────
    def _exportar_gdocs_errores(self):
        """
        Flujo: feedback → PDF en tempdir → subir a Drive → hacer público → link para copiar.
        Si no hay Drive conectado, ofrece exportar localmente como fallback.
        """
        texto = self.editor.get("1.0", "end-1c")
        if not texto.strip():
            messagebox.showwarning("Sin texto", "No hay texto para exportar.")
            return
        if not self._errores_por_linea and not self._errores_tira:
            self._actualizar_vista(texto)

        # ── Verificar conexión Drive ──────────────────────────────────────────
        drive = self._get_drive_conectado()
        if drive is None:
            resp = messagebox.askyesno(
                "Sin conexión a Drive",
                "No estás conectado a Google Drive.\n\n"
                "¿Querés exportar el reporte localmente (PDF/HTML) sin compartir link?")
            if resp:
                self._exportar_reporte_local(texto)
            return

        # ── Feedback ─────────────────────────────────────────────────────────
        feedback = self._dialogo_feedback()
        if feedback is None:
            return

        # ── Generar PDF en carpeta temporal y subir ───────────────────────────
        self._lbl_estado.config(text="⏳ Generando PDF…", fg=C["warn"])
        self.root.update_idletasks()

        import tempfile, threading

        nombre_archivo = "reporte_errores_bloomscans.pdf"

        def _worker():
            try:
                tmp_dir  = tempfile.mkdtemp()
                html_tmp = os.path.join(tmp_dir, "reporte.html")
                pdf_tmp  = os.path.join(tmp_dir, nombre_archivo)

                self._generar_html_gdocs(texto, html_tmp, feedback=feedback)

                pdf_ok = self._html_a_pdf(html_tmp, pdf_tmp)
                upload_path = pdf_tmp if pdf_ok else html_tmp
                upload_name = nombre_archivo if pdf_ok else "reporte_errores_bloomscans.html"

                self.root.after(0, lambda: self._lbl_estado.config(
                    text="⏳ Subiendo a Drive…", fg=C["warn"]))

                file_meta = drive.subir_archivo(upload_path, nombre=upload_name)
                file_id   = file_meta.get("id", "")

                # Hacer público — cualquiera con el link puede ver, sin cuenta Google
                drive.service.permissions().create(
                    fileId=file_id,
                    body={"type": "anyone", "role": "reader"},
                    fields="id"
                ).execute()

                share_link = (
                    f"https://drive.google.com/file/d/{file_id}/view?usp=sharing"
                    if file_id else "")

                self.root.after(0, lambda: self._mostrar_link_compartir(
                    share_link, upload_name, pdf_ok))

            except Exception as ex:
                err = str(ex)
                self.root.after(0, lambda e=err: self._error_exportar_drive(e, texto, feedback))

        threading.Thread(target=_worker, daemon=True).start()

    def _get_drive_conectado(self):
        """Devuelve el DriveClient conectado o None."""
        try:
            if self._panel_drive and self._panel_drive.winfo_exists():
                d = self._panel_drive._drive
                if getattr(d, "service", None):
                    return d
        except Exception:
            pass
        return None

    def _mostrar_link_compartir(self, link: str, nombre: str, es_pdf: bool):
        """Ventana con el link público listo para copiar y compartir."""
        self._lbl_estado.config(text="✅ Subido a Drive · link listo", fg=C["ok"])

        win = tk.Toplevel(self.root)
        win.title("☁ Reporte compartido")
        win.configure(bg=C["bg"])
        win.geometry("500x260")
        win.grab_set()
        win.resizable(False, False)

        tk.Label(win, text="☁  Reporte subido a Drive",
                 bg=C["bg2"], fg=C["ok"],
                 font=("Segoe UI", 11, "bold")).pack(fill="x", pady=8)
        tk.Frame(win, bg=C["border"], height=1).pack(fill="x")

        inner = tk.Frame(win, bg=C["bg"]); inner.pack(fill="both", expand=True, padx=16, pady=12)

        tk.Label(inner,
                 text=("📄 PDF" if es_pdf else "🌐 HTML") + f"  ·  {nombre}",
                 bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 8)).pack(anchor="w")

        tk.Label(inner, text="Link para compartir:",
                 bg=C["bg"], fg=C["text"],
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(10, 3))

        link_frame = tk.Frame(inner, bg=C["bg4"],
                              highlightbackground=C["teal"], highlightthickness=1)
        link_frame.pack(fill="x")
        link_var = tk.StringVar(value=link)
        link_entry = tk.Entry(link_frame, textvariable=link_var,
                              bg=C["bg4"], fg=C["cyan"],
                              font=("Consolas", 8), relief="flat",
                              readonlybackground=C["bg4"], state="readonly")
        link_entry.pack(fill="x", padx=8, pady=6)

        tk.Label(inner,
                 text="✅ Cualquiera con este link puede verlo (sin necesidad de cuenta Google)",
                 bg=C["bg"], fg=C["ok"],
                 font=("Segoe UI", 7)).pack(anchor="w", pady=(4, 0))

        tk.Frame(inner, bg=C["border"], height=1).pack(fill="x", pady=(10, 0))
        btn_row = tk.Frame(win, bg=C["bg"]); btn_row.pack(fill="x", padx=16, pady=8)

        lbl_copiado = tk.Label(btn_row, text="", bg=C["bg"], font=("Segoe UI", 8))

        def _copiar():
            self.root.clipboard_clear()
            self.root.clipboard_append(link)
            lbl_copiado.config(text="✅ ¡Copiado!", fg=C["ok"])
            win.after(2000, lambda: lbl_copiado.config(text=""))

        def _abrir_browser():
            import webbrowser
            webbrowser.open(link)

        make_btn(btn_row, "📋  Copiar link", _copiar,
                 bg=C["teal"], fg=C["bg"], hover_bg=C["cyan"],
                 small=True).pack(side="left", padx=4)
        make_btn(btn_row, "🌐  Abrir en navegador", _abrir_browser,
                 bg=C["bg4"], fg=C["text"], hover_bg=C["bg5"], hover_fg=C["bright"],
                 small=True).pack(side="left", padx=4)
        lbl_copiado.pack(side="left", padx=8)
        make_btn(btn_row, "Cerrar", win.destroy,
                 bg=C["bg4"], fg=C["dim"], hover_bg=C["bg5"],
                 small=True).pack(side="right", padx=4)

    def _error_exportar_drive(self, err: str, texto: str, feedback: dict):
        """Fallback cuando falla la subida a Drive."""
        self._lbl_estado.config(text="❌ Error al subir a Drive", fg=C["err"])
        resp = messagebox.askyesno(
            "Error al subir a Drive",
            f"No se pudo subir a Drive:\n{err}\n\n"
            "¿Querés guardar el reporte localmente igual?")
        if resp:
            self._exportar_reporte_local(texto, feedback=feedback)

    def _exportar_reporte_local(self, texto: str, feedback: dict = None):
        """Fallback: exportar PDF o HTML localmente sin Drive."""
        if feedback is None:
            feedback = self._dialogo_feedback()
            if feedback is None:
                return

        fmt = self._dialogo_formato_reporte()
        if fmt is None:
            return

        ext = ".pdf" if fmt == "pdf" else ".html"
        path = filedialog.asksaveasfilename(
            title="Guardar reporte localmente",
            defaultextension=ext,
            filetypes=[("PDF", "*.pdf"), ("HTML", "*.html"), ("Todos", "*.*")],
            initialfile="reporte_errores" + ext)
        if not path:
            return

        base = path
        for e in (".pdf", ".html", ".htm"):
            if base.lower().endswith(e):
                base = base[:-len(e)]; break

        self._lbl_estado.config(text="⏳ Generando…", fg=C["warn"])
        self.root.update_idletasks()

        try:
            html_p = base + ".html"
            pdf_p  = base + ".pdf"
            self._generar_html_gdocs(texto, html_p, feedback=feedback)

            pdf_ok = False
            if fmt in ("pdf", "ambos"):
                pdf_ok = self._html_a_pdf(html_p, pdf_p)
                if fmt == "pdf" and pdf_ok:
                    try: os.remove(html_p)
                    except: pass

            self._lbl_estado.config(text="✅ Guardado localmente", fg=C["ok"])
            abrir = pdf_p if (fmt in ("pdf", "ambos") and pdf_ok) else html_p
            if messagebox.askyesno("Guardado", f"¿Abrir el archivo?\n{abrir}"):
                os.startfile(abrir) if os.name == "nt" else os.system(f"xdg-open \'{abrir}\'")
        except Exception as ex:
            self._lbl_estado.config(text="❌ Error", fg=C["err"])
            messagebox.showerror("Error", str(ex))

    def _subir_reporte_drive(self, path: str):
        """Sube el reporte a Google Drive (uso legacy)."""
        try:
            drive = self._get_drive_conectado()
            if not drive:
                raise RuntimeError("No conectado a Drive")
            f = drive.subir_archivo(path)
            messagebox.showinfo("Drive", f"✅ Subido:\n{f.get('name', path)}")
        except Exception as ex:
            messagebox.showerror("Error Drive", str(ex))

    def _dialogo_formato_reporte(self) -> str | None:
        """Ventana pequeña para elegir HTML, PDF o ambos. Devuelve 'html'/'pdf'/'ambos'/None."""
        resultado = [None]

        win = tk.Toplevel(self.root)
        win.title("Formato del reporte")
        win.configure(bg=C["bg"])
        win.geometry("360x220")
        win.grab_set()
        win.resizable(False, False)

        tk.Label(win, text="¿En qué formato querés exportar?",
                 bg=C["bg2"], fg=C["teal"],
                 font=("Segoe UI", 10, "bold")).pack(fill="x", pady=10, padx=0)
        tk.Frame(win, bg=C["border"], height=1).pack(fill="x")

        inner = tk.Frame(win, bg=C["bg"]); inner.pack(fill="both", expand=True, padx=20, pady=14)

        desc = {
            "pdf":   ("📄  PDF",   "Lo abre cualquiera, se ve idéntico,\nno genera desconfianza."),
            "html":  ("🌐  HTML",  "Para Google Docs o previsualizar\nen navegador."),
            "ambos": ("📄🌐  PDF + HTML", "Genera los dos archivos."),
        }

        for fmt, (label, sub) in desc.items():
            f = tk.Frame(inner, bg=C["bg4"], cursor="hand2",
                         highlightbackground=C["border2"], highlightthickness=1)
            f.pack(fill="x", pady=3)
            tk.Label(f, text=label, bg=C["bg4"], fg=C["text"],
                     font=("Segoe UI", 9, "bold"), anchor="w").pack(side="left", padx=10, pady=4)
            tk.Label(f, text=sub, bg=C["bg4"], fg=C["dim"],
                     font=("Segoe UI", 7), anchor="w", justify="left").pack(side="left", padx=4)
            val = fmt
            f.bind("<Button-1>", lambda e, v=val: [resultado.__setitem__(0, v), win.destroy()])
            for child in f.winfo_children():
                child.bind("<Button-1>", lambda e, v=val: [resultado.__setitem__(0, v), win.destroy()])
            f.bind("<Enter>", lambda e, w=f: w.config(bg=C["bg5"]))
            f.bind("<Leave>", lambda e, w=f: w.config(bg=C["bg4"]))

        make_btn(win, "✕  Cancelar", win.destroy,
                 bg=C["bg"], fg=C["dim"], hover_bg=C["bg4"],
                 small=True).pack(pady=4)

        win.wait_window()
        return resultado[0]

    def _html_a_pdf(self, html_path: str, pdf_path: str) -> bool:
        """
        Convierte HTML a PDF probando motores en orden:
          1. Chrome headless  (sin instalar nada, ya en el sistema)
          2. Edge headless    (idem, muy comun en Windows)
          3. pdfkit           (pip install pdfkit)
          4. weasyprint       (fallback, falla en Windows sin GTK)
        """
        import subprocess, sys
        abs_html = os.path.abspath(html_path)
        abs_pdf  = os.path.abspath(pdf_path)

        # 1 & 2: Chrome / Edge headless
        if os.name == "nt":
            chrome_candidates = [
                r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
                os.path.expandvars(r"%LOCALAPPDATA%\Google\Chrome\Application\chrome.exe"),
                r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
                r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            ]
        else:
            chrome_candidates = [
                "google-chrome", "google-chrome-stable",
                "chromium", "chromium-browser", "microsoft-edge",
            ]

        for exe in chrome_candidates:
            try:
                cmd = [
                    exe,
                    "--headless=new",
                    "--disable-gpu",
                    "--no-sandbox",
                    "--disable-dev-shm-usage",
                    f"--print-to-pdf={abs_pdf}",
                    "--print-to-pdf-no-header",
                    f"file:///{abs_html}" if os.name == "nt" else f"file://{abs_html}",
                ]
                result = subprocess.run(cmd, capture_output=True, timeout=60)
                if (result.returncode == 0
                        and os.path.exists(abs_pdf)
                        and os.path.getsize(abs_pdf) > 1024):
                    return True
            except Exception:
                continue

        # 3: pdfkit
        try:
            import pdfkit
            pdfkit.from_file(abs_html, abs_pdf)
            if os.path.exists(abs_pdf) and os.path.getsize(abs_pdf) > 1024:
                return True
        except Exception:
            pass

        # 4: weasyprint
        try:
            import weasyprint
            weasyprint.HTML(filename=abs_html).write_pdf(abs_pdf)
            if os.path.exists(abs_pdf) and os.path.getsize(abs_pdf) > 1024:
                return True
        except ImportError:
            try:
                subprocess.run([sys.executable, "-m", "pip", "install", "weasyprint"],
                               capture_output=True, timeout=120)
                import weasyprint
                weasyprint.HTML(filename=abs_html).write_pdf(abs_pdf)
                if os.path.exists(abs_pdf) and os.path.getsize(abs_pdf) > 1024:
                    return True
            except Exception:
                pass
        except Exception:
            pass

        return False

    def _subir_reporte_drive(self, path: str):
        """Sube el reporte (PDF o HTML) a Google Drive usando la conexión existente."""
        try:
            drive = self._panel_drive._drive
            f = drive.subir_archivo(path)
            nombre = f.get("name", os.path.basename(path))
            messagebox.showinfo("Drive", f"✅ Subido correctamente:\n{nombre}")
        except Exception as ex:
            messagebox.showerror("Error Drive", f"No se pudo subir:\n{ex}")

    def _generar_html_gdocs(self, texto: str, path: str, feedback: dict = None):
        """
        Genera HTML con tabla de 2 columnas: línea | comentario/imagen.
        Sin position:absolute — compatible con Chrome/Edge headless al imprimir a PDF.
        """
        lineas = texto.splitlines()

        # ── Mapa de errores ───────────────────────────────────────────────────
        tipo_por_linea = {}
        msgs_por_linea = {}
        for num, errs in self._errores_por_linea.items():
            tipo_por_linea[num] = errs[0][0]
            msgs_por_linea[num] = [m for _, m, *_ in errs]
        for num, desc in getattr(self, "_errores_manuales", {}).items():
            tipo_por_linea[num] = "manual"
            msgs_por_linea.setdefault(num, []).append(f"\u2691 {desc}")
        for num, dat in getattr(self, "_comentarios_img", {}).items():
            if num not in tipo_por_linea:
                tipo_por_linea[num] = "comentario"

        comentarios_img = getattr(self, "_comentarios_img", {})

        COLOR_MAP = {
            "puntuacion": ("#FF8080", "#5a0000"),
            "mayuscula":  ("#9090FF", "#10103a"),
            "simbolo":    ("#FFCC60", "#4a3000"),
            "tilde":      ("#80FF80", "#003a00"),
            "estructura": ("#CC80FF", "#300050"),
            "duplicado":  ("#CC80FF", "#300050"),
            "largo":      ("#80C0FF", "#001040"),
            "nombre":     ("#9090FF", "#10103a"),
            "tira":       ("#FFCC60", "#4a3000"),
            "ortografia": ("#FF80CC", "#4a0030"),
            "mezcla":     ("#FFAA40", "#4a2000"),
            "manual":     ("#FF60A0", "#4a0020"),
            "comentario": ("#60C0FF", "#003050"),
        }
        ICONOS_TIPO = {
            "puntuacion": "\u26a0 Puntuaci\u00f3n", "mayuscula":  "Aa May\u00fascula",
            "simbolo":    "? S\u00edmbolo",           "tilde":      "\u00b4 Tilde faltante",
            "estructura": "\u2699 Estructura",         "duplicado":  "\u229f Duplicado",
            "largo":      "\u2194 L\u00ednea larga",  "nombre":     "\U0001f464 Nombre",
            "tira":       "# Tira",                     "ortografia": "abc Ortograf\u00eda",
            "mezcla":     "\u26a1 Mezcla",             "manual":     "\u2691 Error manual",
            "comentario": "\U0001f4f8 Captura",
        }

        def _esc(s):
            return (s.replace("&", "&amp;")
                     .replace("<", "&lt;")
                     .replace(">", "&gt;"))

        # ── Filas de la tabla ─────────────────────────────────────────────────
        filas = []
        for i, linea in enumerate(lineas):
            num = i + 1
            linea_esc = _esc(linea).replace(" ", "\u00a0") or "\u00a0"

            # Encabezado de tira
            if es_encabezado_tira(linea.strip()):
                filas.append(
                    f'<tr class="row-hdr">' 
                    f'<td class="ln">{num}</td>'
                    f'<td class="code hdr-txt" colspan="2">{linea_esc}</td>'
                    f'</tr>')
                continue

            if num in tipo_por_linea:
                tipo = tipo_por_linea[num]
                bg, fg = COLOR_MAP.get(tipo, ("#FFFF00", "#000000"))
                msgs   = msgs_por_linea.get(num, [])
                icono  = ICONOS_TIPO.get(tipo, tipo.upper())

                # Celda izquierda: la línea del script
                celda_codigo = (
                    f'<td class="code" style="background:{bg};color:{fg}">' +
                    linea_esc + '</td>')

                # Celda derecha: comentario con errores + imagen
                partes_cmt = []

                # Header del comentario
                partes_cmt.append(
                    f'<div class="cb-hdr">\U0001f4ac BloomScans \u00b7 L{num}</div>')

                # Errores de texto
                for m in msgs:
                    partes_cmt.append(
                        f'<div class="cb-item">' 
                        f'<span class="badge" style="background:{fg};color:{bg}">{icono}</span> ' 
                        f'{_esc(m)}</div>')

                # Imagen de captura (si existe)
                if num in comentarios_img:
                    dat     = comentarios_img[num]
                    img_b64 = dat.get("img_b64", "")
                    desc    = _esc(dat.get("texto", "(sin descripci\u00f3n)"))
                    partes_cmt.append(f'<div class="cb-cap-lbl">\U0001f4f8 {desc}</div>')
                    if img_b64:
                        partes_cmt.append(
                            f'<img src="data:image/png;base64,{img_b64}" ' 
                            f'class="cb-img" />')

                celda_cmt = (
                    '<td class="cmt-cell">' +
                    '<div class="cmt-bubble" style="border-color:' + fg + '">' +
                    "".join(partes_cmt) +
                    '</div></td>')

                filas.append(
                    f'<tr>' 
                    f'<td class="ln" style="background:{bg};color:{fg}">{num}</td>' +
                    celda_codigo + celda_cmt +
                    '</tr>')
            else:
                filas.append(
                    f'<tr>' 
                    f'<td class="ln">{num}</td>' 
                    f'<td class="code">{linea_esc}</td>' 
                    f'<td class="cmt-cell"></td>' 
                    f'</tr>')

        filas_str = "\n".join(filas)

        # ── Leyenda ───────────────────────────────────────────────────────────
        leyenda_items = []
        for tipo, (bg, fg) in COLOR_MAP.items():
            icono = ICONOS_TIPO.get(tipo, tipo)
            leyenda_items.append(
                f'<span style="background:{bg};color:{fg};padding:2px 8px;' 
                f'border-radius:3px;margin:2px;display:inline-block;' 
                f'font-size:11px;">{icono}</span>')
        leyenda_str = " ".join(leyenda_items)

        total_errs     = (len(self._errores_por_linea) +
                          len(getattr(self, "_errores_manuales", {})))
        total_capturas = len(comentarios_img)

        # ── Bloque feedback ───────────────────────────────────────────────────
        feedback_html = ""
        if feedback:
            fb_msg     = feedback.get("mensaje", "")
            fb_ap      = feedback.get("aprobado", False)
            fb_nota_ap = feedback.get("nota_ap", "")
            fb_nota    = feedback.get("nota_extra", "")
            parts_fb   = []

            if fb_ap or fb_msg or fb_nota:
                parts_fb.append('<hr style="border-color:#1e2d47;margin-top:32px;">')
                parts_fb.append('<h2 style="color:#00cfb5">\U0001f4dd Feedback del corrector</h2>')

            if fb_ap:
                nota_ap_txt = (" \u2014 " + _esc(fb_nota_ap)) if fb_nota_ap else ""
                parts_fb.append(
                    '<div style="background:#0d2a1a;border-left:4px solid #22c55e;' 
                    'padding:10px 16px;border-radius:6px;margin-bottom:12px;">' 
                    '<span style="color:#22c55e;font-size:14px;font-weight:bold;">' 
                    + "\u2705 APROBADO" + nota_ap_txt + "</span></div>")

            if fb_msg:
                parts_fb.append('<h3 style="color:#1adeff;margin-bottom:4px;">\U0001f4ac Mensaje al traductor</h3>')
                parts_fb.append(
                    '<div style="background:#111827;padding:10px 14px;' 
                    'border-radius:6px;border-left:3px solid #1adeff;' 
                    'white-space:pre-wrap;color:#ccd9ee;">' + _esc(fb_msg) + "</div>")

            if fb_nota:
                parts_fb.append('<h3 style="color:#f59e0b;margin-bottom:4px;">\U0001f4cc Nota adicional</h3>')
                parts_fb.append(
                    '<div style="background:#111827;padding:10px 14px;' 
                    'border-radius:6px;border-left:3px solid #f59e0b;' 
                    'white-space:pre-wrap;color:#ccd9ee;">' + _esc(fb_nota) + "</div>")

            feedback_html = "\n".join(parts_fb)

        # ── HTML final ────────────────────────────────────────────────────────
        html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Reporte de Errores \u2014 BloomScans</title>
<style>
  @media print {{
    @page {{ margin: 0; }}
    body {{
      background: #0d1120 !important;
      -webkit-print-color-adjust: exact;
      print-color-adjust: exact;
      margin: 0 !important;
      padding: 12px !important;
    }}
  }}
  body {{
    font-family: 'Consolas', 'Courier New', monospace;
    background: #0d1120;
    color: #ccd9ee;
    margin: 0;
    padding: 20px;
  }}
  h1  {{ color: #00cfb5; border-bottom: 2px solid #00cfb5; padding-bottom:8px; }}
  h2  {{ color: #1adeff; margin-top:24px; }}
  .meta    {{ color:#7a9cbf; font-size:12px; margin-bottom:16px; }}
  .leyenda {{ background:#111827; padding:12px; border-radius:6px; margin-bottom:20px; }}
  .resumen {{ background:#111827; border-radius:6px; padding:12px 20px;
              margin-bottom:20px; border-left:4px solid #00cfb5; }}

  /* Tabla principal */
  table {{
    width: 100%;
    border-collapse: collapse;
    table-layout: fixed;
  }}
  td {{ vertical-align: top; padding: 3px 6px; }}
  tr:hover td {{ filter: brightness(1.06); }}
  tr {{ border-bottom: 1px solid #1e2d4718; }}

  /* Columna n\u00famero */
  .ln {{
    width: 36px;
    text-align: right;
    font-size: 11px;
    color: #506070;
    border-right: 1px solid #1e2d47;
    padding-right: 6px;
    white-space: nowrap;
    user-select: none;
  }}

  /* Columna c\u00f3digo */
  .code {{
    width: 55%;
    font-size: 12px;
    white-space: pre-wrap;
    word-break: break-word;
    padding-left: 8px;
  }}

  /* Encabezado de tira */
  .row-hdr td {{ background: #111827; }}
  .hdr-txt {{ color: #00cfb5; font-weight: bold; }}

  /* Columna comentario */
  .cmt-cell {{
    width: 42%;
    padding: 3px 8px;
  }}
  .cmt-bubble {{
    background: #1a2640;
    border: 1.5px solid #60C0FF;
    border-radius: 6px;
    padding: 6px 9px;
    font-size: 11px;
    color: #ccd9ee;
  }}
  .cb-hdr {{
    font-weight: bold;
    color: #60C0FF;
    font-size: 10px;
    border-bottom: 1px solid #2a4060;
    padding-bottom: 3px;
    margin-bottom: 5px;
  }}
  .cb-item {{
    margin: 2px 0;
    line-height: 1.4;
    word-break: break-word;
  }}
  .badge {{
    display: inline-block;
    padding: 1px 4px;
    border-radius: 3px;
    font-size: 9px;
    font-weight: bold;
  }}
  .cb-cap-lbl {{
    color: #60C0FF;
    font-style: italic;
    font-size: 10px;
    margin-top: 5px;
    display: block;
    border-top: 1px solid #2a4060;
    padding-top: 4px;
  }}
  .cb-img {{
    max-width: 100%;
    max-height: 110px;
    display: block;
    margin-top: 4px;
    border-radius: 3px;
    border: 1px solid #60C0FF;
  }}
</style>
</head>
<body>
<h1>\U0001f4cb Reporte de Errores \u2014 BloomScans</h1>
<div class="meta">Corrector de Traducciones v7.0 \u00b7 BloomScans</div>

<div class="resumen">
  <strong style="color:#00cfb5">Resumen:</strong>
  &nbsp; {total_errs} error(es) de texto
  &nbsp;\u00b7&nbsp; {total_capturas} captura(s) con imagen
</div>

<h2>Leyenda</h2>
<div class="leyenda">{leyenda_str}</div>

<h2>Script completo</h2>
<table>
<colgroup>
  <col style="width:36px">
  <col style="width:55%">
  <col style="width:42%">
</colgroup>
<thead>
  <tr style="background:#182030">
    <th class="ln">#</th>
    <th style="text-align:left;padding:4px 8px;color:#7a9cbf;font-size:11px">L\u00ednea</th>
    <th style="text-align:left;padding:4px 8px;color:#7a9cbf;font-size:11px">Comentario / Captura</th>
  </tr>
</thead>
<tbody>
{filas_str}
</tbody>
</table>

{feedback_html}
<p style="color:#354d6a;font-size:11px;margin-top:24px;">
  Para abrir en Google Docs: Drive \u2192 clic derecho \u2192 Abrir con \u2192 Google Docs
</p>
</body>
</html>"""

        with open(path, "w", encoding="utf-8") as f:
            f.write(html)


    def _set_welcome(self):
        txt = (
            "\\## Corrector de Traducciones v6.0\n\n"
            "N: 🤖 IA Gemini — análisis visual con Gemini Vision (¡NUEVO!)\n"
            "N: 📂 Abrir — .docx, .txt o .md (también desde Drive)\n"
            "N: 🔍 Analizar — detecta todos los errores\n"
            "N: ✨ Auto-corregir — aplica fixes mecánicos seguros\n"
            "N: ´ Auto-tilde — preview con checkbox antes de aplicar\n"
            "N: 📊 Stats — estadísticas por símbolo\n"
            "N: 🔀 Comparar — original vs. corregido lado a lado\n"
            "N: 📋 Reporte — exporta lista de errores a .txt\n"
            "N: 🖼 Visor — imágenes del manwha (local/ZIP/Drive), bidireccional\n"
            "N: ☁ Drive — abrir .docx, subir corregido, cargar raws\n\n"
            "\\## IA Gemini v7.0  (Google AI Studio)\n\n"
            "N: Cargá las imágenes RAW en el visor primero\n"
            "N: Conseguí tu API Key GRATIS en: aistudio.google.com/apikey\n"
            "G: MODO DETECTAR — lista: globos sin traducir, orden, simbología, onomatopeyas\n"
            "G: MODO CORREGIR — genera script corregido listo para aplicar\n"
            "N: Podés enviar imagen actual o TODAS las imágenes del capítulo\n"
            "N: Ajustá calidad JPEG y altura máxima para reducir costo de tokens\n\n"
            "\\## Panel de símbolos\n\n"
            "N: Usá los botones arriba del editor para insertar símbolos\n"
            "N: El botón inserta en la posición del cursor\n\n"
            "\\## Detecciones automáticas\n\n"
            "G: ¿? y ¡! sin par de apertura/cierre\n"
            "T2: Tildes: tambien→también, entrare→entraré (Auto-tilde con preview)\n"
            "D2: Mayúsculas internas: 'Tambien', 'Gracias A Ti'\n"
            "P2: Guiones múltiples ta--- → ta-\n"
            "DF: Símbolos con typos d2: → D2:\n"
            "N: Líneas duplicadas por tira\n"
            "(): Línea que termina con coma\n"
            "N: Líneas muy largas (configurable)\n"
            "N: Tiras faltantes o repetidas en la numeración\n"
            "N: Variantes de escritura de nombres propios\n"
            "abc: Typos ortográficos obvios (requiere: pip install pyspellchecker)\n"
            "⚡: Mezcla de sistemas BL/+18 y +15 en el mismo archivo\n"
        )
        self.editor.insert("1.0", txt)
        self._actualizar_vista(txt)


    # MEJORA v6.0: Panel IA — integrado en el layout principal ──────────────────
    def _abrir_panel_ia(self):
        if self._ia_visible:
            # Ocultar — sacar del PanedWindow
            self._h_paned.forget(self._panel_ia)
            self._ia_visible = False
            if hasattr(self, "_btn_ia"):
                self._btn_ia.config(bg=C["bg5"], fg=C["purple"])
        else:
            # Mostrar — agregar como tercer panel
            self._h_paned.add(self._panel_ia, minsize=420, stretch="never")
            self._ia_visible = True
            if hasattr(self, "_btn_ia"):
                self._btn_ia.config(bg=C["purple"], fg=C["bright"])

    def _ia_get_imagen_actual(self):
        """Devuelve la fuente de imagen actualmente visible en el visor."""
        if not self._visor_imagenes:
            return None
        nombre, fuente = self._visor_imagenes[self._visor_idx]
        return fuente

    def _ia_get_todas_imagenes(self) -> list:
        """Devuelve todas las fuentes de imagen cargadas en el visor."""
        return [fuente for _, fuente in self._visor_imagenes]

    def _ia_get_tira_actual(self) -> str:
        """Devuelve el nombre/número de la tira visible."""
        if self._visor_imagenes:
            return self._visor_imagenes[self._visor_idx][0]
        return "Tira actual"

    @staticmethod
    def _es_encabezado_tira(linea: str) -> bool:
        """
        Detecta cualquier variante de encabezado de tira:
          ## Tira 2 | ##Tira 2 | \\## Tira 2 | ### ##Tira 1 | ##tira2 | etc.
        """
        l = linea.strip()
        # Quitar prefijos basura como "### " antes del ##
        l = re.sub(r'^[#\s]+(?=##)', '', l)
        # Ahora debe empezar con ## o \##
        return bool(re.match(r'^\\?##', l))

    @staticmethod
    def _numero_tira(linea: str) -> int | None:
        """Extrae el número de tira de cualquier variante de encabezado."""
        m = re.search(r'\d+', linea)
        return int(m.group()) if m else None

    def _ia_get_script_tira(self) -> str:
        """
        Extrae las líneas de la tira actualmente visible del editor.
        Añade marcadores explícitos de inicio/fin para que Gemini sepa
        exactamente qué tira corresponde a la imagen enviada.
        """
        texto = self.editor.get("1.0", "end-1c")
        if not self._visor_imagenes:
            return texto

        idx = self._visor_idx + 1  # tiras empiezan en 1
        lineas = texto.splitlines()
        resultado = []
        en_tira = False

        for linea in lineas:
            if self._es_encabezado_tira(linea):
                num = self._numero_tira(linea)
                if num == idx:
                    en_tira = True
                    resultado.append(linea)
                elif en_tira:
                    break
            elif en_tira:
                resultado.append(linea)

        if not resultado:
            # Fallback por posición
            encabezados = [(i, self._numero_tira(l))
                           for i, l in enumerate(lineas)
                           if self._es_encabezado_tira(l)]
            encabezados_ord = sorted(encabezados, key=lambda x: x[0])
            if idx - 1 < len(encabezados_ord):
                inicio_idx = encabezados_ord[idx - 1][0]
                fin_idx = (encabezados_ord[idx][0]
                           if idx < len(encabezados_ord) else len(lineas))
                resultado = lineas[inicio_idx:fin_idx]
            else:
                return texto

        # Añadir marcadores explícitos para Gemini
        tira_nombre = self._visor_imagenes[self._visor_idx][0] if self._visor_imagenes else f"Tira {idx}"
        header = f"[INICIO TIRA {idx} — imagen: {tira_nombre}]"
        footer = f"[FIN TIRA {idx}]"
        return header + "\n" + "\n".join(resultado) + "\n" + footer

    def _ia_aplicar_correccion(self, texto_corregido: str):
        """
        Reemplaza las líneas de la tira actual con el texto corregido por la IA.
        Robusto frente a variantes de encabezado.
        """
        texto = self.editor.get("1.0", "end-1c")
        if not self._visor_imagenes:
            self.editor.delete("1.0", "end")
            self.editor.insert("1.0", texto_corregido)
            self._texto_original = self._texto_original or texto
            return

        idx = self._visor_idx + 1
        lineas = texto.splitlines()
        inicio = fin = None

        for i, linea in enumerate(lineas):
            if self._es_encabezado_tira(linea):
                num = self._numero_tira(linea)
                if num == idx:
                    inicio = i
                elif inicio is not None:
                    fin = i
                    break

        if inicio is None:
            # Fallback por posición
            encabezados = [(i, self._numero_tira(l))
                           for i, l in enumerate(lineas)
                           if self._es_encabezado_tira(l)]
            encabezados_ord = sorted(encabezados, key=lambda x: x[0])
            if idx - 1 < len(encabezados_ord):
                inicio = encabezados_ord[idx - 1][0]
                fin = (encabezados_ord[idx][0]
                       if idx < len(encabezados_ord) else len(lineas))
            else:
                return  # No hay dónde aplicar

        if fin is None:
            fin = len(lineas)

        nuevas = texto_corregido.splitlines()
        lineas_nuevas = lineas[:inicio] + nuevas + lineas[fin:]
        texto_final = "\n".join(lineas_nuevas)
        self._texto_original = self._texto_original or texto
        self.editor.delete("1.0", "end")
        self.editor.insert("1.0", texto_final)
        self._actualizar_vista(texto_final)
        messagebox.showinfo("IA Gemini", "✅ Corrección de la IA aplicada al editor.")



# ═══════════════════════════════════════════════════════════════════════════════
#  PANEL IA GEMINI
# ═══════════════════════════════════════════════════════════════════════════════
class PanelGeminiIA(tk.Frame):
    """
    Panel integrado para corregir traducciones con Gemini Vision (Google AI Studio).
    Modo 1: Detectar errores  |  Modo 2: Corregir todo
    Soporta envío de imagen actual o TODAS las imágenes del capítulo.
    Controles de calidad/tamaño para reducir costo de tokens.
    """
    def __init__(self, parent,
                 get_texto_fn, get_script_tira_fn,
                 get_visor_fn, get_todas_imagenes_fn,
                 get_tira_fn, aplicar_correccion_fn,
                 get_nombres_propios_fn=None):
        super().__init__(parent, bg=C["bg"])

        self._get_texto             = get_texto_fn
        self._get_script_tira       = get_script_tira_fn
        self._get_imagen            = get_visor_fn
        self._get_todas_imagenes    = get_todas_imagenes_fn
        self._get_tira              = get_tira_fn
        self._aplicar               = aplicar_correccion_fn
        self._get_nombres_propios   = get_nombres_propios_fn or (lambda: set())
        self._correc_pendiente      = None
        self._en_proceso            = False

        self._build()
        key = _cargar_gemini_key()
        if key:
            self._var_key.set(key)

    # ── Construcción UI ───────────────────────────────────────────────────────
    def _build(self):
        # Header (sin botón ✕ — es panel integrado, se cierra con el botón de la barra principal)
        hdr = tk.Frame(self, bg=C["bg2"]); hdr.pack(fill="x")
        tk.Label(hdr, text="🤖  Gemini Vision — Corrector IA",
                 bg=C["bg2"], fg=C["purple"],
                 font=("Segoe UI", 11, "bold")).pack(side="left", padx=14, pady=9)
        tk.Label(hdr, text="Google AI Studio",
                 bg=C["bg2"], fg=C["dim"],
                 font=("Segoe UI", 8)).pack(side="left")
        sep(self)

        # Scroll canvas para todo el contenido
        outer = tk.Frame(self, bg=C["bg"]); outer.pack(fill="both", expand=True)
        canvas = tk.Canvas(outer, bg=C["bg"], highlightthickness=0)
        vsb = ttk.Scrollbar(outer, orient="vertical", style="Vertical.TScrollbar",
                            command=canvas.yview)
        canvas.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)
        content = tk.Frame(canvas, bg=C["bg"])
        cw = canvas.create_window((0, 0), window=content, anchor="nw")

        def _on_configure(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def _on_canvas_resize(e):
            canvas.itemconfig(cw, width=e.width)
        content.bind("<Configure>", _on_configure)
        canvas.bind("<Configure>", _on_canvas_resize)

        def _card(titulo, color=None):
            f = tk.Frame(content, bg=C["bg3"],
                         highlightbackground=C["border2"], highlightthickness=1)
            f.pack(fill="x", padx=10, pady=(0, 8))
            tk.Label(f, text=titulo, bg=C["bg3"], fg=color or C["teal"],
                     font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=(6,2))
            inner = tk.Frame(f, bg=C["bg3"]); inner.pack(fill="x", padx=10, pady=(0, 8))
            return inner

        # ── API Key ───────────────────────────────────────────────────────────
        kinner = _card("🔑  API Key — Google AI Studio", C["purple"])
        tk.Label(kinner,
                 text="Conseguí tu key GRATIS en: aistudio.google.com/apikey",
                 bg=C["bg3"], fg=C["dim"], font=("Segoe UI", 7)).pack(anchor="w", pady=(0,4))
        key_row = tk.Frame(kinner, bg=C["bg3"]); key_row.pack(fill="x")
        self._var_key = tk.StringVar()
        self._entry_key = tk.Entry(key_row, textvariable=self._var_key,
                                   bg=C["bg4"], fg=C["text"],
                                   insertbackground=C["purple"],
                                   show="•", relief="flat", font=("Consolas", 9))
        self._entry_key.pack(side="left", fill="x", expand=True, ipady=4)
        self._var_ver_key = tk.BooleanVar(value=False)
        def _toggle_key():
            self._entry_key.config(show="" if self._var_ver_key.get() else "•")
        tk.Checkbutton(key_row, text="Ver", variable=self._var_ver_key,
                       command=_toggle_key,
                       bg=C["bg3"], fg=C["dim"], selectcolor=C["bg4"],
                       activebackground=C["bg3"],
                       font=("Segoe UI", 8)).pack(side="left", padx=6)
        make_btn(kinner, "💾 Guardar key", self._guardar_key,
                 bg=C["bg3"], fg=C["dim"], hover_bg=C["bg4"], hover_fg=C["purple"],
                 small=True).pack(anchor="w", pady=(4, 0))

        # ── Modelo ────────────────────────────────────────────────────────────
        minner = _card("🧠  Modelo Gemini")
        self._var_modelo = tk.StringVar(value=GEMINI_MODELO_DEFAULT)
        for modelo_id, modelo_desc in GEMINI_MODELOS.items():
            tk.Radiobutton(minner, text=modelo_desc,
                           variable=self._var_modelo, value=modelo_id,
                           bg=C["bg3"], fg=C["text2"], selectcolor=C["bg4"],
                           activebackground=C["bg3"], font=("Segoe UI", 8)
                           ).pack(anchor="w")

        # ── Modo de análisis ──────────────────────────────────────────────────
        ainner = _card("⚙  Modo de análisis")
        self._var_modo = tk.StringVar(value="detectar")
        modo_row = tk.Frame(ainner, bg=C["bg3"]); modo_row.pack(fill="x")
        for val, ico, titulo, desc in [
            ("detectar", "🔍", "Detectar errores",
             "Visual (Gemini) + mecánico combinados:\nglobos, orden, simbología, tildes, puntuación"),
            ("corregir", "✨", "Corregir todo",
             "Genera el script corregido completo\nlisto para aplicar al editor"),
        ]:
            rb_frame = tk.Frame(modo_row, bg=C["bg4"],
                                highlightbackground=C["border"], highlightthickness=1)
            rb_frame.pack(side="left", fill="both", expand=True, padx=(0, 4))
            ri = tk.Frame(rb_frame, bg=C["bg4"]); ri.pack(fill="x", padx=8, pady=6)
            tk.Radiobutton(ri, text=f"{ico}  {titulo}",
                           variable=self._var_modo, value=val,
                           bg=C["bg4"], fg=C["text"], selectcolor=C["bg5"],
                           activebackground=C["bg4"], activeforeground=C["teal"],
                           font=("Segoe UI", 9, "bold")).pack(anchor="w")
            tk.Label(ri, text=desc, bg=C["bg4"], fg=C["dim"],
                     font=("Segoe UI", 7), justify="left").pack(anchor="w")

        # ── Imágenes a enviar ─────────────────────────────────────────────────
        iinner = _card("🖼  Imágenes a enviar a Gemini")
        self._var_imagenes = tk.StringVar(value="actual")
        for val, txt, hint in [
            ("actual",  "Solo la imagen actual del visor",
             "Más rápido y económico — ideal para revisar tira a tira"),
            ("todas",   "Todas las imágenes del capítulo",
             "Gemini ve todo el contexto — más preciso pero más costo"),
        ]:
            row = tk.Frame(iinner, bg=C["bg3"]); row.pack(fill="x", pady=1)
            tk.Radiobutton(row, text=txt,
                           variable=self._var_imagenes, value=val,
                           bg=C["bg3"], fg=C["text"], selectcolor=C["bg4"],
                           activebackground=C["bg3"],
                           font=("Segoe UI", 9, "bold")).pack(side="left")
            tk.Label(row, text=f"  {hint}", bg=C["bg3"], fg=C["dim"],
                     font=("Segoe UI", 7)).pack(side="left")

        # ── Calidad de imagen (ahorro de tokens) ──────────────────────────────
        cinner = _card("💰  Optimización de imágenes  (reduce costo de tokens)")

        # Calidad JPEG
        cal_row = tk.Frame(cinner, bg=C["bg3"]); cal_row.pack(fill="x", pady=2)
        tk.Label(cal_row, text="Calidad JPEG:", bg=C["bg3"], fg=C["text2"],
                 font=("Segoe UI", 8), width=16, anchor="w").pack(side="left")
        self._var_calidad = tk.IntVar(value=60)
        self._lbl_calidad = tk.Label(cal_row, text="60 %", bg=C["bg3"], fg=C["teal"],
                                     font=("Segoe UI", 8, "bold"), width=5)
        self._lbl_calidad.pack(side="right")
        def _on_calidad(v):
            self._lbl_calidad.config(text=f"{int(float(v))} %")
        tk.Scale(cal_row, variable=self._var_calidad, from_=20, to=95,
                 orient="horizontal", command=_on_calidad,
                 bg=C["bg3"], fg=C["text2"], troughcolor=C["bg4"],
                 highlightthickness=0, sliderrelief="flat",
                 length=300).pack(side="left", fill="x", expand=True, padx=4)
        tk.Label(cinner,
                 text="60 = buen equilibrio calidad/costo  |  30 = máximo ahorro  |  85 = máxima calidad",
                 bg=C["bg3"], fg=C["dim"], font=("Segoe UI", 7)).pack(anchor="w")

        # Altura máxima
        alt_row = tk.Frame(cinner, bg=C["bg3"]); alt_row.pack(fill="x", pady=(6,2))
        tk.Label(alt_row, text="Altura máxima (px):", bg=C["bg3"], fg=C["text2"],
                 font=("Segoe UI", 8), width=16, anchor="w").pack(side="left")
        self._var_max_alto = tk.IntVar(value=1200)
        self._lbl_alto = tk.Label(alt_row, text="1200 px", bg=C["bg3"], fg=C["teal"],
                                  font=("Segoe UI", 8, "bold"), width=7)
        self._lbl_alto.pack(side="right")
        def _on_alto(v):
            self._lbl_alto.config(text=f"{int(float(v))} px")
        tk.Scale(alt_row, variable=self._var_max_alto, from_=400, to=3000,
                 resolution=100, orient="horizontal", command=_on_alto,
                 bg=C["bg3"], fg=C["text2"], troughcolor=C["bg4"],
                 highlightthickness=0, sliderrelief="flat",
                 length=300).pack(side="left", fill="x", expand=True, padx=4)
        tk.Label(cinner,
                 text="800–1200 px = suficiente para leer texto  |  Reducir ahorra tokens significativamente",
                 bg=C["bg3"], fg=C["dim"], font=("Segoe UI", 7)).pack(anchor="w")

        # Alcance del script
        sinner = _card("📄  Alcance del script")
        self._var_alcance = tk.StringVar(value="tira")
        for val, txt in [
            ("tira", "Solo la tira visible en el visor"),
            ("todo", "Todo el script del capítulo"),
        ]:
            tk.Radiobutton(sinner, text=txt,
                           variable=self._var_alcance, value=val,
                           bg=C["bg3"], fg=C["text2"], selectcolor=C["bg4"],
                           activebackground=C["bg3"],
                           font=("Segoe UI", 8)).pack(anchor="w")

        # ── Botón principal ───────────────────────────────────────────────────
        btn_row = tk.Frame(content, bg=C["bg"]); btn_row.pack(fill="x", padx=10, pady=(0, 6))
        self._btn_analizar = make_btn(btn_row, "🤖  Analizar con Gemini",
                                      self._ejecutar,
                                      bg=C["purple"], fg=C["bright"],
                                      hover_bg=C["cyan"], hover_fg=C["bg"])
        self._btn_analizar.pack(side="left", padx=(0, 10))
        self._lbl_estado = tk.Label(btn_row, text="", bg=C["bg"],
                                    fg=C["dim"], font=("Segoe UI", 8))
        self._lbl_estado.pack(side="left")

        # ── Resultado ─────────────────────────────────────────────────────────
        tk.Label(content, text="Resultado:", bg=C["bg"], fg=C["text2"],
                 font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=10, pady=(0, 2))
        res_outer = tk.Frame(content, bg=C["bg"]); res_outer.pack(fill="both", expand=True, padx=10)
        res_sb = ttk.Scrollbar(res_outer, orient="vertical", style="Vertical.TScrollbar")
        self._txt_resultado = tk.Text(
            res_outer, bg=C["bg2"], fg=C["text"],
            font=("Consolas", 9), wrap="word", relief="flat",
            bd=0, padx=10, pady=8, state="disabled",
            height=10, yscrollcommand=res_sb.set)
        res_sb.config(command=self._txt_resultado.yview)
        res_sb.pack(side="right", fill="y")
        self._txt_resultado.pack(side="left", fill="both", expand=True)
        self._txt_resultado.tag_config("ok",   foreground=C["ok"])
        self._txt_resultado.tag_config("err",  foreground=C["err"])
        self._txt_resultado.tag_config("warn", foreground=C["warn"])
        self._txt_resultado.tag_config("code", foreground=C["teal"],
                                        font=("Consolas", 9))

        # ── Botón aplicar corrección ──────────────────────────────────────────
        self._btn_aplicar = make_btn(content, "✅  Aplicar corrección al editor",
                                     self._aplicar_resultado,
                                     bg=C["ok"], fg=C["bg"],
                                     hover_bg=C["teal"], hover_fg=C["bg"])
        self._btn_aplicar.pack(padx=10, pady=(6, 10))
        self._btn_aplicar.pack_forget()

    # ── Lógica ────────────────────────────────────────────────────────────────
    def _guardar_key(self):
        key = self._var_key.get().strip()
        if not key:
            messagebox.showwarning("API Key", "Ingresá tu API Key primero.", parent=self)
            return
        _guardar_gemini_key(key)
        self._lbl_estado.config(text="✅ Key guardada", fg=C["ok"])

    def _set_resultado(self, texto: str, tag: str = ""):
        self._txt_resultado.config(state="normal")
        self._txt_resultado.delete("1.0", "end")
        if tag:
            self._txt_resultado.insert("1.0", texto, tag)
        else:
            self._txt_resultado.insert("1.0", texto)
        self._txt_resultado.config(state="disabled")

    def _ejecutar(self):
        if self._en_proceso:
            return

        api_key = self._var_key.get().strip()
        if not api_key:
            messagebox.showwarning("API Key",
                "Ingresá tu API Key de Google AI Studio.\n"
                "Es GRATUITA en: aistudio.google.com/apikey",
                parent=self)
            return

        if not REQUESTS_OK:
            if messagebox.askyesno("requests no instalado",
                    "La librería 'requests' no está instalada.\n¿Instalar ahora?",
                    parent=self):
                import subprocess, sys
                subprocess.run([sys.executable, "-m", "pip", "install", "requests"],
                               capture_output=True)
                messagebox.showinfo("requests",
                    "Instalado. Reiniciá la aplicación.", parent=self)
            return

        # Determinar imágenes a enviar
        usar_todas = self._var_imagenes.get() == "todas"
        if usar_todas:
            imagenes = self._get_todas_imagenes()
            if not imagenes:
                messagebox.showwarning("Sin imágenes",
                    "No hay imágenes cargadas en el visor.", parent=self)
                return
        else:
            img = self._get_imagen()
            if img is None:
                messagebox.showwarning("Sin imagen",
                    "No hay imagen en el visor.\n"
                    "Abrí las imágenes RAW en el visor primero.", parent=self)
                return
            imagenes = [img]

        # Script
        alcance = self._var_alcance.get()
        texto = self._get_script_tira() if alcance == "tira" else self._get_texto()
        tira  = self._get_tira()
        modelo   = self._var_modelo.get()
        calidad  = self._var_calidad.get()
        max_alto = self._var_max_alto.get()
        modo     = self._var_modo.get()

        # ── Advertencia si imagen y script no están sincronizados ────────────
        # Detectar número de tira en el script enviado vs imagen del visor
        if not usar_todas and alcance == "tira":
            num_en_img = None
            tira_lower = tira.lower()
            m = re.search(r'\d+', tira_lower)
            if m:
                num_en_img = int(m.group())
            num_en_script = None
            for linea in texto.splitlines():
                if es_encabezado_tira(linea):
                    num_en_script = numero_tira_de_linea(linea)
                    break
            if (num_en_img is not None and num_en_script is not None
                    and num_en_img != num_en_script):
                if not messagebox.askyesno(
                        "⚠ Desincronización detectada",
                        f"La imagen visible es de la Tira {num_en_img},\n"
                        f"pero el script enviado corresponde a la Tira {num_en_script}.\n\n"
                        f"Gemini podría reportar errores falsos por comparar imágenes incorrectas.\n\n"
                        f"Recomendación: navegá al panel del visor y seleccioná la imagen "
                        f"correcta antes de analizar.\n\n"
                        f"¿Continuar de todas formas?",
                        parent=self):
                    return

        n_imgs = len(imagenes)
        self._en_proceso = True
        self._correc_pendiente = None
        self._btn_aplicar.pack_forget()
        self._btn_analizar.config(state="disabled", text="⏳  Analizando…")

        modo_txt = "Detectando errores" if modo == "detectar" else "Generando corrección"
        self._lbl_estado.config(
            text=f"🔄 {modo_txt} — {n_imgs} img…", fg=C["warn"])
        self._set_resultado(
            f"⏳ {modo_txt}…\n\n"
            f"Imagen(es): {n_imgs}  |  Tira: {tira}\n"
            f"Modelo: {modelo}\n"
            f"Calidad: {calidad}%  |  Alto máx: {max_alto}px\n\n"
            f"ℹ️  Modo DETECTAR solo reporta errores visuales objetivos:\n"
            f"   globos faltantes, orden, simbología, onomatopeyas fuera de globo.\n"
            f"   NO reporta sugerencias de estilo ni correcciones de redacción.\n\n"
            f"(Puede tardar 10–60 segundos…)",
            "warn")

        def _worker():
            try:
                respuesta = llamar_gemini(
                    api_key=api_key,
                    modo=modo,
                    texto_script=texto,
                    imagenes=imagenes,
                    tira_nombre=tira,
                    modelo=modelo,
                    calidad=calidad,
                    max_alto=max_alto,
                )
                self.after(0, lambda: self._mostrar_resultado(respuesta, modo))
            except Exception as ex:
                err_msg = str(ex)
                self.after(0, lambda m=err_msg: self._mostrar_error(m))

        threading.Thread(target=_worker, daemon=True).start()

    def _mostrar_resultado(self, respuesta: str, modo: str):
        self._en_proceso = False
        self._btn_analizar.config(state="normal", text="🤖  Analizar con Gemini")
        self._lbl_estado.config(text="✅ Análisis completado", fg=C["ok"])

        if modo == "corregir":
            self._correc_pendiente = respuesta
            self._set_resultado(respuesta, "code")
            self._btn_aplicar.pack(padx=10, pady=(6, 10))
            return

        # ── Modo detectar: combinar Gemini + detector mecánico ────────────────
        # 1. Obtener errores mecánicos del script de esta tira
        errores_mecanicos = self._obtener_errores_mecanicos()

        # 2. Construir texto combinado
        lineas_salida = []
        tiene_errores = False

        # Sección Gemini (visual)
        lineas_salida.append("═" * 50)
        lineas_salida.append("🤖  GEMINI — Análisis visual (imagen vs. script)")
        lineas_salida.append("═" * 50)
        gemini_vacio = respuesta.strip().startswith("✓") or not respuesta.strip()
        if gemini_vacio:
            lineas_salida.append("✓ Sin errores visuales detectados.")
        else:
            tiene_errores = True
            lineas_salida.append(respuesta.strip())

        # Sección mecánica
        lineas_salida.append("")
        lineas_salida.append("═" * 50)
        lineas_salida.append("⚙  DETECTOR AUTOMÁTICO — Errores de formato/texto")
        lineas_salida.append("═" * 50)
        if errores_mecanicos:
            tiene_errores = True
            lineas_salida.extend(errores_mecanicos)
        else:
            lineas_salida.append("✓ Sin errores mecánicos detectados.")

        # Resumen
        lineas_salida.append("")
        lineas_salida.append("─" * 50)
        if tiene_errores:
            lineas_salida.append("⚠  Se encontraron errores. Revisá los detalles arriba.")
        else:
            lineas_salida.append("✓  Todo OK — sin errores detectados.")

        texto_final = "\n".join(lineas_salida)
        self._set_resultado(texto_final)

        # Colorear líneas
        self._txt_resultado.config(state="normal")
        for i, linea in enumerate(texto_final.splitlines()):
            n = i + 1
            l = linea.strip()
            if l.startswith("✓"):
                self._txt_resultado.tag_add("ok", f"{n}.0", f"{n}.end")
            elif l.startswith("═") or l.startswith("─"):
                self._txt_resultado.tag_add("info", f"{n}.0", f"{n}.end")
            elif l.startswith("🤖") or l.startswith("⚙"):
                self._txt_resultado.tag_add("info", f"{n}.0", f"{n}.end")
            elif re.match(r'\[(GLOBO FALTANTE|GLOBO EXTRA|ORDEN|SIMBOLOGÍA|ONOMATOPEYA)', l, re.IGNORECASE):
                self._txt_resultado.tag_add("err", f"{n}.0", f"{n}.end")
            elif re.match(r'\[', l) or l.startswith("⚠") or re.match(r'\s+L\d+', linea):
                self._txt_resultado.tag_add("warn", f"{n}.0", f"{n}.end")
        self._txt_resultado.config(state="disabled")

    def _obtener_errores_mecanicos(self) -> list:
        """
        Corre el detector mecánico sobre las líneas de la tira actual
        o de todo el script según el alcance elegido, y devuelve lista de strings.
        """
        try:
            # Respetar el alcance elegido (tira vs todo)
            alcance = self._var_alcance.get() if hasattr(self, "_var_alcance") else "tira"
            if alcance == "todo":
                script = self._get_texto()
            else:
                script = self._get_script_tira()
            lineas_limpias = [l for l in script.splitlines()
                              if not l.startswith("[INICIO TIRA") and
                                 not l.startswith("[FIN TIRA")]
            if not lineas_limpias:
                return []

            nombres = self._get_nombres_propios() if self._get_nombres_propios else set()
            errores_out = []

            for i, linea in enumerate(lineas_limpias):
                num = i + 1
                stripped = linea.strip()
                if not stripped or es_encabezado_tira(stripped):
                    continue
                errs = analizar_linea(linea, nombres)
                for tipo, msg, *_ in errs:
                    ICONOS = {
                        "puntuacion": "⚠ PUNTUACIÓN",
                        "mayuscula":  "Aa MAYÚSCULA",
                        "simbolo":    "? SÍMBOLO",
                        "tilde":      "´ TILDE",
                        "estructura": "⚙ ESTRUCTURA",
                        "duplicado":  "⊡ DUPLICADO",
                        "largo":      "↔ LARGO",
                        "ortografia": "abc ORTOGRAFÍA",
                    }
                    label = ICONOS.get(tipo, f"• {tipo.upper()}")
                    errores_out.append(f"  L{num:3d}  [{label}]  {msg}")
                    errores_out.append(f"         → {stripped[:70]}")

            # Errores de tira (salto de numeración)
            errs_tira = analizar_tiras(lineas_limpias)
            for _, msg in errs_tira:
                errores_out.append(f"  [# ESTRUCTURA TIRA]  {msg}")

            # Mezcla de sistemas
            texto_tira = "\n".join(lineas_limpias)
            for _, _, msg in detectar_mezcla_sistemas(texto_tira):
                errores_out.append(f"  [⚡ MEZCLA SISTEMAS]  {msg}")

            return errores_out
        except Exception as ex:
            return [f"  [Error al correr detector mecánico: {ex}]"]

    def _mostrar_error(self, msg: str):
        self._en_proceso = False
        self._btn_analizar.config(state="normal", text="🤖  Analizar con Gemini")
        self._lbl_estado.config(text="❌ Error", fg=C["err"])
        self._set_resultado(
            f"❌ Error al llamar a Gemini:\n\n{msg}\n\n"
            "Verificá:\n"
            "• Tu API Key es correcta (aistudio.google.com/apikey)\n"
            "• Tenés conexión a internet\n"
            "• El modelo seleccionado está disponible en tu región",
            "err")

    def _aplicar_resultado(self):
        if not self._correc_pendiente:
            return
        if messagebox.askyesno(
                "Aplicar corrección",
                "¿Reemplazar las líneas de la tira con la versión corregida por Gemini?\n\n"
                "Podés deshacer con Ctrl+Z.",
                parent=self):
            self._aplicar(self._correc_pendiente)
            self._btn_aplicar.pack_forget()
            self._correc_pendiente = None


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    try:
        root = tk.Tk()
        root.configure(bg=C["bg"])
        try:
            root.iconbitmap("icon.ico")
        except Exception:
            pass
        CorrectorApp(root)
        root.mainloop()
    except Exception as e:
        import traceback
        try:
            import tkinter.messagebox as mb
            mb.showerror("Error al iniciar", traceback.format_exc())
        except Exception:
            print(traceback.format_exc())

if __name__ == "__main__":
    main()

