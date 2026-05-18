from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\corba\Downloads\Solo Leveling Completo Continuo")


def insert_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(text)
    return new_para


def replace_with_lines(doc: Document, match_text: str, new_lines: list[str]) -> bool:
    for p in doc.paragraphs:
        if p.text.strip() == match_text:
            p.text = new_lines[0]
            anchor = p
            for line in new_lines[1:]:
                anchor = insert_after(anchor, line)
            return True
    return False


def fix_file(path: Path, first_heading: str | None = None, replacements: list[tuple[str, list[str]]] | None = None) -> None:
    doc = Document(path)
    if first_heading:
        for p in doc.paragraphs:
            if p.text.strip():
                p.text = first_heading
                break
    for old, new in replacements or []:
        replace_with_lines(doc, old, new)
    doc.save(path)


def main() -> None:
    fixes: dict[str, dict] = {
        r"Volumen 2\Capitulo 12 - Trato.docx": {
            "heading": "Capitulo 12 - Trato",
            "replacements": [
                (
                    "Recompensa 2: Todas las estadísticas +3 Recompensa 3: Seleccione una de las dos opciones disponibles.",
                    [
                        "Recompensa 2: Todas las estadísticas +3",
                        "Recompensa 3: Seleccione una de las dos opciones disponibles.",
                    ],
                ),
                ("Has recibo varias recompensas.", ["Has recibido varias recompensas."]),
            ],
        },
        r"Volumen 10\Capitulo 97 - Proyectando sombras.docx": {
            "replacements": [
                ("Aquel que… (2 más) Fatiga:", ["Aquel que… (2 más)", "Fatiga:"]),
                ("Inteligencia: 310 Vitalidad:", ["Inteligencia: 310", "Vitalidad:"]),
                (
                    "Percepción: 277 Reduc de daño físico: 65% Puntos sin asignar:",
                    ["Percepción: 277", "Reduc de daño físico: 65%", "Puntos sin asignar:"],
                ),
                ("Habilidades pasivas [Desconocido]: Nivel máximo.", ["Habilidades pasivas", "[Desconocido]: Nivel máximo."]),
                ("Habilidades activas [Correr a toda velocidad]: Nivel máximo.", ["Habilidades activas", "[Correr a toda velocidad]: Nivel máximo."]),
            ]
        },
        r"Volumen 11\Capitulo 114 - Nuevo Presidente de la Asociación de Cazadores.docx": {
            "replacements": [
                ("Aquel que… (2 más) Fatiga:", ["Aquel que… (2 más)", "Fatiga:"]),
                ("Inteligencia: 321 Vitalidad:", ["Inteligencia: 321", "Vitalidad:"]),
                (
                    "Percepción: 298 Reducción de daño físico: 65% Puntos sin asignar:",
                    ["Percepción: 298", "Reducción de daño físico: 65%", "Puntos sin asignar:"],
                ),
            ]
        },
        r"Volumen 2\Capitulo 16 - Equipo.docx": {
            "replacements": [
                ("Asesino de lobos Fatiga:", ["Asesino de lobos", "Fatiga:"]),
                ("Inteligencia: 39 Vitalidad:", ["Inteligencia: 39", "Vitalidad:"]),
                (
                    "Percepción: 69 Reduc de daño físico: 20% Puntos sin asignar:",
                    ["Percepción: 69", "Reduc de daño físico: 20%", "Puntos sin asignar:"],
                ),
                ("Habilidades pasivas [Desconocido]: Nivel máximo.", ["Habilidades pasivas", "[Desconocido]: Nivel máximo."]),
                ("Habilidades activas [Correr a toda velocidad]: Nivel 1.", ["Habilidades activas", "[Correr a toda velocidad]: Nivel 1."]),
                (
                    "Equipamiento Collar del guardián de la puerta (A) En poco tiempo casi había alcanzado el nivel 30. Al principio de la incursión se había sentido decepcionado, pero tras derribar al jefe de los goblins y a sus subordinados, había ganado inesperadamente dos niveles. Fue una suerte que todas las bestias mágicas fuertes se hubiesen reunido en la sala del jefe.",
                    [
                        "Equipamiento",
                        "Collar del guardián de la puerta (A)",
                        "En poco tiempo casi había alcanzado el nivel 30. Al principio de la incursión se había sentido decepcionado, pero tras derribar al jefe de los goblins y a sus subordinados, había ganado inesperadamente dos niveles. Fue una suerte que todas las bestias mágicas fuertes se hubiesen reunido en la sala del jefe.",
                    ],
                ),
                (
                    "Reduc de daño físico: 20 % Puntos sin asignar: 0 ‘Sí. Los múltiples de cinco son una buena cosa.’ – pensó mientras sonreía con satisfacción.",
                    [
                        "Reduc de daño físico: 20 %",
                        "Puntos sin asignar: 0",
                        "‘Sí. Los múltiples de cinco son una buena cosa.’ – pensó mientras sonreía con satisfacción.",
                    ],
                ),
            ]
        },
        r"Volumen 2\Capitulo 18 - Presintiendo el éxito.docx": {
            "replacements": [
                ("Asesino de lobos Fatiga:", ["Asesino de lobos", "Fatiga:"]),
                ("Inteligencia: 51 Vitalidad:", ["Inteligencia: 51", "Vitalidad:"]),
                (
                    "Percepción: 81 Reduc de daño físico: 20% Puntos sin asignar:",
                    ["Percepción: 81", "Reduc de daño físico: 20%", "Puntos sin asignar:"],
                ),
                ("Habilidades pasivas [Desconocido]: Nivel máximo.", ["Habilidades pasivas", "[Desconocido]: Nivel máximo."]),
                ("Habilidades activas [Correr a toda velocidad]: Nivel 2.", ["Habilidades activas", "[Correr a toda velocidad]: Nivel 2."]),
            ]
        },
        r"Volumen 2\Capitulo 19 - Misión de cambio de clase.docx": {
            "replacements": [
                (
                    "40 Raza: Humano Clase: Ninguna La tercera línea de su pantalla de estado, ‘Clase’.",
                    ["Raza: Humano", "Clase: Ninguna", "La tercera línea de su pantalla de estado, ‘Clase’."],
                ),
            ]
        },
        r"Volumen 3\Capitulo 20 - El funcionamiento de la mazmorra.docx": {
            "replacements": [
                ("Asesino de lobos Fatiga: 43 Salud:", ["Asesino de lobos", "Fatiga: 43", "Salud:"]),
                (
                    "4.511/8.330 Maná: 660/790 ‘Tengo que ganar con esto.’ Después de realizar una verificación final de su estado, Jinwoo agarró el pomo de la puerta.",
                    ["4.511/8.330", "Maná: 660/790", "‘Tengo que ganar con esto.’ Después de realizar una verificación final de su estado, Jinwoo agarró el pomo de la puerta."],
                ),
            ]
        },
        r"Volumen 3\Capitulo 22 - La verdadera prueba.docx": {
            "replacements": [
                ("Inteligencia: 66 Vitalidad:", ["Inteligencia: 66", "Vitalidad:"]),
                (
                    "87 (+20) Percepción: 89 Reduc de daño físico: 46% (+15%) Puntos sin asignar:",
                    ["87 (+20)", "Percepción: 89", "Reduc de daño físico: 46% (+15%)", "Puntos sin asignar:"],
                ),
                ("Fatiga: 61 Salud:", ["Fatiga: 61", "Salud:"]),
                ("4.161/10.270 Maná: 390/850 ‘La [Fatiga] es un poco alta, pero aún es manejable.’", ["4.161/10.270", "Maná: 390/850", "‘La [Fatiga] es un poco alta, pero aún es manejable.’"]),
                ("Salud: 1.036 / 10.270 ‘¿Cuánto tiempo llevo…?’", ["Salud: 1.036 / 10.270", "‘¿Cuánto tiempo llevo…?’"]),
                ("Fatiga: 91 Salud:", ["Fatiga: 91", "Salud:"]),
                ("104/10.270 Maná: 202/850 Se preguntaba por qué le era tan complicado moverse, pero ahora sabía el motivo: su [Fatiga] había pasado los 90 puntos.", ["104/10.270", "Maná: 202/850", "Se preguntaba por qué le era tan complicado moverse, pero ahora sabía el motivo: su [Fatiga] había pasado los 90 puntos."]),
                ("Fatiga: 0 Salud:", ["Fatiga: 0", "Salud:"]),
                ("106/10.270 Maná: 204/850 Aunque había vaciado toda la botella, su [Salud] se mantenía inmutable. Era drásticamente diferente a la [Fatiga] que se había reducido a cero.", ["106/10.270", "Maná: 204/850", "Aunque había vaciado toda la botella, su [Salud] se mantenía inmutable. Era drásticamente diferente a la [Fatiga] que se había reducido a cero."]),
            ]
        },
        r"Volumen 3\Capitulo 25 - Monarca de las sombras.docx": {
            "replacements": [
                ("Asesino de lobos (1 más) Fatiga:", ["Asesino de lobos (1 más)", "Fatiga:"]),
                ("Inteligencia: 70 Vitalidad:", ["Inteligencia: 70", "Vitalidad:"]),
                (
                    "Percepción: 93 Reduc de daño físico: 46% Puntos sin asignar:",
                    ["Percepción: 93", "Reduc de daño físico: 46%", "Puntos sin asignar:"],
                ),
                ("Habilidades pasivas [Desconocido]: Nivel máximo.", ["Habilidades pasivas", "[Desconocido]: Nivel máximo."]),
                ("Habilidades activas [Correr a toda velocidad]: Nivel 2.", ["Habilidades activas", "[Correr a toda velocidad]: Nivel 2."]),
            ]
        },
        r"Volumen 4\Capitulo 32 - Bestia o humano.docx": {
            "replacements": [
                ("El que triunfa… (1 más) Fatiga:", ["El que triunfa… (1 más)", "Fatiga:"]),
                ("Inteligencia: 99 Vitalidad:", ["Inteligencia: 99", "Vitalidad:"]),
                (
                    "Percepción: 103 Reduc de daño físico: 46% Puntos sin asignar:",
                    ["Percepción: 103", "Reduc de daño físico: 46%", "Puntos sin asignar:"],
                ),
            ]
        },
        r"Volumen 5\Capitulo 50 - Segunda incursión en el castillo demoniaco.docx": {
            "replacements": [
                ("Aquel que… (1 más) Fatiga:", ["Aquel que… (1 más)", "Fatiga:"]),
                ("Inteligencia: 189 Vitalidad:", ["Inteligencia: 189", "Vitalidad:"]),
                (
                    "Percepción: 126 Reduc de daño físico: 46% Puntos sin asignar:",
                    ["Percepción: 126", "Reduc de daño físico: 46%", "Puntos sin asignar:"],
                ),
                ("Habilidades pasivas [Desconocido]: Nivel máximo.", ["Habilidades pasivas", "[Desconocido]: Nivel máximo."]),
                ("Habilidades activas [Correr a toda velocidad]: Nivel Máximo.", ["Habilidades activas", "[Correr a toda velocidad]: Nivel Máximo."]),
            ]
        },
    }

    for rel, config in fixes.items():
        fix_file(ROOT / rel, config.get("heading"), config.get("replacements"))
        print(ROOT / rel)


if __name__ == "__main__":
    main()
